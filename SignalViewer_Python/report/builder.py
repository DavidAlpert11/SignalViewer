"""
Signal Viewer Pro - Report Builder
===================================
Build and export HTML/DOCX reports with plots and metadata.
"""

import os
import io
import json
from datetime import datetime
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, field

from core.models import Run, ViewState, SubplotConfig
from compare.engine import CompareResult

# Optional imports for DOCX export
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[WARN] python-docx not installed - DOCX export disabled", flush=True)


@dataclass
class ReportSection:
    """A section in the report"""
    title: str
    content: str = ""
    plot_data: Optional[str] = None  # Base64 PNG or Plotly JSON
    signals: List[str] = field(default_factory=list)


@dataclass
class Report:
    """Complete report structure"""
    title: str = "Signal Viewer Report"
    introduction: str = ""
    conclusion: str = ""
    
    # Metadata
    created_at: str = field(default_factory=lambda: datetime.now().isoformat())
    runs: List[str] = field(default_factory=list)  # Run names/paths
    
    # Sections
    subplot_sections: List[ReportSection] = field(default_factory=list)
    compare_sections: List[ReportSection] = field(default_factory=list)


def build_report(
    runs: List[Run],
    view_state: ViewState,
    compare_results: List[CompareResult],
    intro_text: str = "",
    conclusion_text: str = "",
) -> Report:
    """
    Build a report from current state.
    
    Args:
        runs: List of loaded runs
        view_state: Current view state
        compare_results: List of compare results
        intro_text: Introduction text
        conclusion_text: Conclusion text
        
    Returns:
        Report object
    """
    report = Report(
        introduction=intro_text,
        conclusion=conclusion_text,
        runs=[r.csv_display_name for r in runs],
    )
    
    # Add subplot sections
    for sp_config in view_state.subplots:
        if not sp_config.include_in_report:
            continue
        
        section = ReportSection(
            title=sp_config.caption or f"Subplot {sp_config.index + 1}",
            content=sp_config.description,
            signals=sp_config.assigned_signals.copy(),
        )
        report.subplot_sections.append(section)
    
    # Add compare sections
    for result in compare_results:
        section = ReportSection(
            title=f"Compare: {result.signal_name}",
            content=_format_compare_metrics(result),
            signals=[result.signal_name],
        )
        report.compare_sections.append(section)
    
    return report


def _format_compare_metrics(result: CompareResult) -> str:
    """Format compare metrics as text"""
    lines = [
        f"Baseline: {result.baseline_name}",
        f"Compare To: {result.compare_to_name}",
        "",
        "Metrics:",
        f"  Max Absolute Difference: {result.max_abs_diff:.6g}",
        f"  RMS Difference: {result.rms_diff:.6g}",
        f"  Mean Difference: {result.mean_diff:.6g}",
        f"  Correlation: {result.correlation:.4f}",
    ]
    
    if result.tolerance_violations > 0:
        lines.extend([
            "",
            f"Tolerance Violations: {result.tolerance_violations} ({result.tolerance_violation_pct:.1f}%)",
        ])
    
    return "\n".join(lines)


def export_html(
    report: Report,
    filepath: str,
    figure_json: Optional[str] = None,
) -> bool:
    """
    Export report to offline HTML.
    
    Args:
        report: Report to export
        filepath: Output file path
        figure_json: Optional Plotly figure JSON for embedding
        
    Returns:
        True if successful
    """
    try:
        html = _generate_html(report, figure_json)
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html)
        
        print(f"[REPORT] Exported to {filepath}")
        return True
        
    except Exception as e:
        print(f"[ERROR] Failed to export report: {e}")
        return False


def _generate_html(report: Report, figure_json: Optional[str]) -> str:
    """Generate HTML content"""
    
    # Minimal inline Plotly (for offline viewing)
    plotly_script = """
    <script>
    // Minimal Plotly renderer - requires figure JSON
    function renderFigure(containerId, figureJson) {
        if (typeof Plotly !== 'undefined' && figureJson) {
            var fig = JSON.parse(figureJson);
            Plotly.newPlot(containerId, fig.data, fig.layout);
        }
    }
    </script>
    """
    
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{report.title}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            background: #f5f5f5;
            color: #333;
        }}
        h1 {{ color: #2E86AB; border-bottom: 2px solid #2E86AB; padding-bottom: 10px; }}
        h2 {{ color: #444; margin-top: 30px; }}
        .section {{ background: white; padding: 20px; margin: 20px 0; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }}
        .metadata {{ color: #666; font-size: 0.9em; }}
        .metrics {{ background: #f8f9fa; padding: 15px; border-radius: 4px; font-family: monospace; white-space: pre-line; }}
        .signals {{ color: #666; font-size: 0.85em; margin-top: 10px; }}
        .plot-container {{ width: 100%; height: 400px; }}
        details {{ margin: 10px 0; }}
        summary {{ cursor: pointer; color: #2E86AB; }}
    </style>
    <script src="https://cdn.plot.ly/plotly-2.27.0.min.js"></script>
    {plotly_script}
</head>
<body>
    <h1>{report.title}</h1>
    
    <div class="section metadata">
        <p><strong>Generated:</strong> {report.created_at}</p>
        <details>
            <summary>Data Sources ({len(report.runs)} files)</summary>
            <ul>
                {''.join(f'<li>{r}</li>' for r in report.runs)}
            </ul>
        </details>
    </div>
"""
    
    if report.introduction:
        html += f"""
    <div class="section">
        <h2>Introduction</h2>
        <p>{report.introduction}</p>
    </div>
"""
    
    # Subplot sections
    if report.subplot_sections:
        html += """
    <h2>Plot Sections</h2>
"""
        for i, section in enumerate(report.subplot_sections):
            html += f"""
    <div class="section">
        <h3>{section.title}</h3>
        {f'<p>{section.content}</p>' if section.content else ''}
        <div id="plot-{i}" class="plot-container"></div>
        <div class="signals">Signals: {', '.join(section.signals) if section.signals else 'None'}</div>
    </div>
"""
    
    # Compare sections
    if report.compare_sections:
        html += """
    <h2>Comparison Results</h2>
"""
        for section in report.compare_sections:
            html += f"""
    <div class="section">
        <h3>{section.title}</h3>
        <div class="metrics">{section.content}</div>
    </div>
"""
    
    if report.conclusion:
        html += f"""
    <div class="section">
        <h2>Conclusion</h2>
        <p>{report.conclusion}</p>
    </div>
"""
    
    # Embed figure if provided
    if figure_json:
        html += f"""
    <script>
        document.addEventListener('DOMContentLoaded', function() {{
            renderFigure('plot-0', '{figure_json.replace("'", "\\'")}');
        }});
    </script>
"""
    
    html += """
</body>
</html>
"""
    
    return html


def export_docx(
    report: Report,
    filepath: str,
    figure: Optional[Any] = None,
    rtl: bool = False,
) -> bool:
    """
    Export report to Word document (.docx).
    
    Args:
        report: Report to export
        filepath: Output file path
        figure: Optional Plotly figure for embedding as image
        rtl: Enable RTL/Hebrew text direction
        
    Returns:
        True if successful
    """
    if not DOCX_AVAILABLE:
        print("[ERROR] python-docx not installed - cannot export DOCX", flush=True)
        return False
    
    try:
        doc = Document()
        
        # Title
        title = doc.add_heading(report.title, level=0)
        if rtl:
            title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Metadata
        meta = doc.add_paragraph()
        meta.add_run(f"Generated: {report.created_at}\n").italic = True
        if report.runs:
            meta.add_run(f"Data Sources: {', '.join(report.runs)}\n").italic = True
        if rtl:
            meta.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Introduction
        if report.introduction:
            h = doc.add_heading("Introduction", level=1)
            if rtl:
                h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p = doc.add_paragraph(report.introduction)
            if rtl:
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Subplot sections
        if report.subplot_sections:
            h = doc.add_heading("Plot Sections", level=1)
            if rtl:
                h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            for section in report.subplot_sections:
                sh = doc.add_heading(section.title, level=2)
                if rtl:
                    sh.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                if section.content:
                    p = doc.add_paragraph(section.content)
                    if rtl:
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                if section.signals:
                    sig_p = doc.add_paragraph()
                    sig_p.add_run("Signals: ").bold = True
                    sig_p.add_run(", ".join(section.signals))
                    if rtl:
                        sig_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Embed figure as image if provided
        if figure is not None:
            try:
                import kaleido
                img_bytes = figure.to_image(format="png", width=1200, height=600, scale=2)
                doc.add_picture(io.BytesIO(img_bytes), width=Inches(6.5))
            except Exception as e:
                print(f"[WARN] Could not embed figure: {e}", flush=True)
        
        # Compare sections
        if report.compare_sections:
            h = doc.add_heading("Comparison Results", level=1)
            if rtl:
                h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            for section in report.compare_sections:
                sh = doc.add_heading(section.title, level=2)
                if rtl:
                    sh.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                if section.content:
                    p = doc.add_paragraph(section.content)
                    p.style = 'Quote'
                    if rtl:
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Conclusion
        if report.conclusion:
            h = doc.add_heading("Conclusion", level=1)
            if rtl:
                h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p = doc.add_paragraph(report.conclusion)
            if rtl:
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Save
        doc.save(filepath)
        print(f"[REPORT] Exported DOCX to {filepath}", flush=True)
        return True
        
    except Exception as e:
        print(f"[ERROR] Failed to export DOCX: {e}", flush=True)
        return False

