"""
Signal Viewer Pro - Professional Signal Visualization Application
==================================================================

A modern, feature-rich signal visualization tool for analyzing time-series
and correlation data from CSV files.

Core Features:
- Multi-CSV loading with automatic duplicate handling
- Multi-tab, multi-subplot layouts (up to 4x4 grid per tab)
- Interactive time cursor with synchronized value display across subplots
- Signal customization (color, scale, line width, display name)

Analysis Features:
- X-Y plot mode for signal correlation analysis
- Derived signals (derivative, integral, custom math operations)
- Multi-signal operations (average, sum, difference, etc.)
- Custom time column selection per CSV

Data Management:
- Session save/load with full state persistence
- Template save/load for layout reuse across sessions
- Resizable panels using Split.js

Export Features:
- HTML report export with all tabs/subplots
- CSV export for signal data
- Subplot metadata (title, caption, description) included in reports

Author: Signal Viewer Team
Version: 2.4
"""

import dash
from dash import (
    dcc,
    html,
    Input,
    Output,
    State,
    callback_context,
    ALL,
    no_update,
    Patch,
)
import dash_bootstrap_components as dbc
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import pandas as pd
import numpy as np
from collections import Counter  # For efficient duplicate detection
import os
import base64
import json
import logging
from datetime import datetime
from typing import Optional, List, Dict, Tuple, Any
# from layout import create_layout  # Using self.create_layout() instead
from data_manager import DataManager
from signal_operations import SignalOperationsManager
from linking_manager import LinkingManager
from config import (
    SIGNAL_COLORS,
    get_theme_dict,
    APP_TITLE,
    APP_HOST,
    APP_PORT,
)
# Import CSV loading support
from flexible_csv_loader import FlexibleCSVLoader

from helpers import (
    get_csv_display_name,
    get_csv_short_name,
    get_signal_label,
    make_signal_key,
    parse_signal_key,
    interpolate_value_at_x,
    safe_json_parse,
    calculate_derived_signal,
    calculate_multi_signal_operation,
    clamp,
    get_text_direction_style,
    get_text_direction_attr,
    compare_signals,
    compare_csv_signals,
)
from callback_helpers import clear_performance_cache
import hashlib

# Configure logging
logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger("SignalViewer")


# ============================================================================
# FIX #1: Safe ID hashing for pattern-matching callbacks
# Prevents "TypeError: unhashable type: 'dict'" errors
# ============================================================================

def safe_id_hash(id_dict: Dict) -> str:
    """
    Convert dict ID to hashable string for safe comparison.
    Fixes: TypeError: unhashable type: 'dict'
    
    Args:
        id_dict: Pattern-matching callback ID dictionary
        
    Returns:
        Stable hash string
    """
    if not isinstance(id_dict, dict):
        return str(id_dict)
    
    # Sort keys for stable hash
    items = sorted(id_dict.items())
    id_str = json.dumps(items, sort_keys=True)
    return hashlib.md5(id_str.encode()).hexdigest()


def get_triggered_id_info() -> Tuple[Optional[str], Optional[Dict]]:
    """
    Get triggered callback ID and info safely without unhashable dict errors.
    
    Returns:
        (prop_id_string, id_dict or None)
    """
    ctx = callback_context
    
    if not ctx.triggered:
        return None, None
    
    triggered = ctx.triggered[0]
    prop_id = triggered.get("prop_id", "")
    
    # Try to extract dict ID from prop_id
    try:
        if "." in prop_id:
            id_part = prop_id.split(".")[0]
            # Check if it's a dict (pattern-matching)
            if id_part.startswith("{"):
                id_dict = json.loads(id_part)
                return prop_id, id_dict
    except:
        pass
    
    return prop_id, None


# Legacy compatibility - keep THEMES dict for existing code
THEMES = {
    "dark": get_theme_dict("dark"),
    "light": get_theme_dict("light"),
}


class SignalViewerApp:
    def __init__(self):
        self.app = dash.Dash(
            __name__,
            external_stylesheets=[
                "/assets/bootstrap-cyborg.min.css",
                "/assets/font-awesome.min.css",
            ],
            external_scripts=[
                "/assets/split.min.js",
                "/assets/collapse.js",
            ],
            suppress_callback_exceptions=True,
        )
        self.app.title = "Signal Viewer Pro"

        # Clean up cache on startup
        self._cleanup_cache_on_startup()

        self.data_manager = DataManager(self)
        self.signal_operations = SignalOperationsManager(self)
        self.linking_manager = LinkingManager(self)

        self.derived_signals = {}
        self.signal_properties = {}
        
        # Store original file paths for streaming/refresh
        self.original_file_paths = {}  # {csv_path: original_source_path}
        
        # Native file dialog support
        self._pending_file_paths = []  # Paths selected in file dialog
        self._file_dialog_active = False  # Flag for dialog state

        # Performance settings
        self._signal_cache = {}  # Raw data cache: {(csv_idx, signal_name, time_col): (x_data, y_data)}
        self._cache_valid = True
        self._last_figure_hash = None  # Track if figure needs rebuild
        self._last_figure = None  # Cache last figure for fast cursor updates
        self._signal_list_cache = {}  # Cache for signal names per CSV: {csv_idx: [signal_names]}
        self._signal_list_version = 0  # Increment when CSVs change
        self._last_tree_state = {}  # Cache for signal tree updates
        self._numpy_array_cache = {}  # Pre-converted numpy arrays: {(csv_idx, signal): ndarray}

        # PERFORMANCE BOOST: Enhanced figure caching to eliminate white flash
        self._figure_cache = {}  # {(tab_idx, state_hash): go.Figure}
        self._state_hash_cache = {}  # {tab_idx: last_state_hash}
        self._last_render_time = 0  # Track render performance
        self._render_count = 0  # Count renders for monitoring
        self._cache_hit_count = 0  # Track cache effectiveness
        self._cache_miss_count = 0  # Track cache misses

        # Display settings - PERFORMANCE via WebGL (no downsampling - show all raw points)
        self.WEBGL_THRESHOLD = 0  # ALWAYS use WebGL for hardware acceleration
        self.HOVER_THRESHOLD = 5000  # Disable hover above this many points for performance
        self.LARGE_DATA_THRESHOLD = 50000  # Above this, use thinner lines

        # PERFORMANCE: Pre-compute thresholds for adaptive rendering
        self.ULTRA_LARGE_THRESHOLD = 500000  # Above this, use minimal rendering options

        self.app.layout = self.create_layout()
        self.setup_callbacks()
    
    def _cleanup_cache_on_startup(self):
        """Clean up all cache files on app startup for fresh state"""
        import shutil
        
        print("[STARTUP] ========================================")
        print("[STARTUP] Signal Viewer Pro - Initializing...")
        print("[STARTUP] ========================================")
        print("[CLEANUP] Cleaning up all caches on startup...")
        
        # Clean entire uploads folder (CSVs and cache)
        uploads_dir = os.path.join(os.path.dirname(__file__), "uploads")
        if os.path.exists(uploads_dir):
            try:
                shutil.rmtree(uploads_dir)
                print(f"   [OK] Cleared uploads folder")
            except Exception as e:
                print(f"   [WARN] Could not clear uploads: {e}")
        
        # Recreate empty uploads folder
        os.makedirs(uploads_dir, exist_ok=True)
        
        # Clear in-memory caches (important for fresh start)
        self._signal_cache = {}
        self._signal_list_cache = {}
        self._figure_cache = {}
        self._state_hash_cache = {}
        self._numpy_array_cache = {}
        self._last_tree_state = {}
        self._last_figure_hash = None
        self._last_figure = None
        
        print("   [OK] Cleared in-memory caches")
        print("[CLEANUP] Cleanup complete - starting with fresh state")
        print("[STARTUP] ========================================")
    
    def _perf_timer(self, name: str):
        """Context manager for performance timing - outputs to console"""
        import time
        
        class PerfTimer:
            def __init__(self, timer_name):
                self.name = timer_name
                self.start = None
            
            def __enter__(self):
                self.start = time.perf_counter()
                return self
            
            def __exit__(self, *args):
                elapsed_ms = (time.perf_counter() - self.start) * 1000
                # Color code based on performance
                if elapsed_ms < 50:
                    status = "⚡"  # Fast
                elif elapsed_ms < 200:
                    status = "⏱️"   # OK
                elif elapsed_ms < 500:
                    status = "🐢"  # Slow
                else:
                    status = "🔴"  # Very slow
                print(f"[PERF] {status} {self.name}: {elapsed_ms:.1f}ms")
        
        return PerfTimer(name)

    def invalidate_caches(self):
        """Invalidate all caches when data changes"""
        self._signal_cache.clear()
        self._signal_list_cache.clear()
        self._last_figure_hash = None

        # PERFORMANCE: Clear enhanced figure cache
        self._figure_cache.clear()
        self._state_hash_cache.clear()

        # Also clear the callback helper caches (signal tree, highlighted signals)
        clear_performance_cache()
        self._cache_valid = True
        self._signal_list_version += 1

        # Log cache statistics for monitoring
        if self._cache_hit_count + self._cache_miss_count > 0:
            hit_rate = (self._cache_hit_count / (self._cache_hit_count + self._cache_miss_count)) * 100
            print(f"[CACHE] Stats: {self._cache_hit_count} hits, {self._cache_miss_count} misses ({hit_rate:.1f}% hit rate)")
        self._cache_hit_count = 0
        self._cache_miss_count = 0
    
    def _compute_figure_hash(self, assignments, tab_key, layouts, props,
                              display_options, cursor_x, link_axes, time_columns,
                              x_axis_signals=None, subplot_modes=None) -> str:
        """Compute a hash of figure-affecting state to detect if rebuild is needed.

        This allows skipping expensive figure rebuilds when only non-data things change.
        Includes X-axis signals and subplot modes for X-Y mode support.
        """
        import hashlib

        # Get relevant state
        tab_assignments = assignments.get(tab_key, {})
        tab_layout = layouts.get(tab_key, {"rows": 1, "cols": 1})

        # Create a string representation of state that affects traces
        # NOTE: cursor_x is NOT included because cursor line can be updated via Patch
        # without rebuilding the entire figure (MAJOR PERFORMANCE OPTIMIZATION)
        state_parts = [
            f"tab:{tab_key}",
            f"layout:{tab_layout}",
            f"link:{link_axes}",
            # REMOVED: f"cursor:{cursor_x}", - cursor updates use Patch now
            f"timecols:{time_columns}",
        ]

        # Add subplot modes (affects X-Y vs time mode)
        if subplot_modes:
            tab_modes = subplot_modes.get(tab_key, {})
            for sp_key, mode in sorted(tab_modes.items()):
                state_parts.append(f"mode:{sp_key}:{mode}")

        # Add X-axis signals (critical for X-Y mode)
        if x_axis_signals:
            tab_x_axis = x_axis_signals.get(tab_key, {})
            for sp_key, x_sig in sorted(tab_x_axis.items()):
                state_parts.append(f"xaxis:{sp_key}:{x_sig}")

        # Add assignments (signal data)
        for sp_key, signals in sorted(tab_assignments.items()):
            if isinstance(signals, list):
                for s in signals:
                    state_parts.append(f"sig:{sp_key}:{s.get('csv_idx')}:{s.get('signal')}")

        # Add relevant props
        for sig_key, p in sorted(props.items()):
            state_parts.append(f"prop:{sig_key}:{p.get('scale',1)}:{p.get('color','')}:{p.get('width',1.5)}")

        # Add display options
        tab_display = display_options.get(tab_key, {})
        for sp_key, opts in sorted(tab_display.items()):
            state_parts.append(f"disp:{sp_key}:{opts}")

        state_str = "|".join(state_parts)
        return hashlib.md5(state_str.encode()).hexdigest()[:16]

    def get_signal_names_cached(self, csv_idx: int) -> list[str]:
        """Get cached list of signal names for a CSV (excluding Time column)"""
        if csv_idx in self._signal_list_cache:
            return self._signal_list_cache[csv_idx]
        
        if csv_idx < 0 or csv_idx >= len(self.data_manager.data_tables):
            return []
            
        df = self.data_manager.data_tables[csv_idx]
        if df is None or df.empty:
            return []
            
        # Get all columns except Time-like columns
        signals = [c for c in df.columns if c.lower() != "time"]
        self._signal_list_cache[csv_idx] = signals
        return signals
        
    def get_signal_data_cached(
        self, csv_idx: int, signal_name: str, time_col: str = "Time"
    ) -> tuple[np.ndarray, np.ndarray]:
        """Get RAW signal data with memory caching (no decimation)"""
        cache_key = (csv_idx, signal_name, time_col)

        # Check cache first
        if self._cache_valid and cache_key in self._signal_cache:
            return self._signal_cache[cache_key]

        # Cache miss - read from DataFrame
        if csv_idx < 0 or csv_idx >= len(self.data_manager.data_tables):
            return np.array([]), np.array([])

        df = self.data_manager.data_tables[csv_idx]
        if df is None or signal_name not in df.columns:
            return np.array([]), np.array([])

        # Get time column
        if time_col not in df.columns:
            if "Time" in df.columns:
                time_col = "Time"
            else:
                time_col = df.columns[0]

        x_data = df[time_col].values
        y_data = df[signal_name].values

        # Store in cache
        self._signal_cache[cache_key] = (x_data, y_data)

        return x_data, y_data

    def create_layout(self):
        return html.Div(
            id="app-container",
            style={
                "backgroundColor": "#1a1a2e",
                "minHeight": "100vh",
                "padding": "10px",
            },
            children=[
                # Stores
                dcc.Store(id="store-csv-files", data=[]),
                dcc.Store(id="store-assignments", data={"0": {"0": []}}),
                dcc.Store(id="store-layouts", data={"0": {"rows": 1, "cols": 1}}),
                dcc.Store(id="store-theme", data="dark"),
                dcc.Store(id="store-selected-subplot", data=0),
                dcc.Store(id="store-selected-tab", data=0),
                dcc.Store(id="store-num-tabs", data=1),
                dcc.Store(id="store-derived", data={}),
                dcc.Store(id="store-signal-props", data={}),
                dcc.Store(id="store-links", data=[]),
                dcc.Store(id="store-context-signal", data=None),
                dcc.Store(id="store-link-axes", data={}),
                dcc.Store(id="store-highlighted", data=[]),
                dcc.Store(id="store-refresh-trigger", data=0),
                dcc.Store(id="store-streaming-active", data=False),
                dcc.Interval(id='interval-streaming', interval=200, disabled=True),
                dcc.Store(id="store-search-filters", data=[]),
                dcc.Store(id="store-cursor-x", data={"x": None, "initialized": False}),
                dcc.Store(id="store-cursor-slider", data=50),  # Slider position (0-100) to break circular dependency
                dcc.Store(
                    id="store-subplot-modes", data={}
                ),  # {tab: {subplot: "time"|"xy"}}
                dcc.Store(
                    id="store-time-columns", data={}
                ),  # {csv_idx: column_name} - which column is time
                dcc.Store(
                    id="store-x-axis-signal", data={}
                ),  # {tab: {subplot: signal_key}} - custom X axis for xy mode
                dcc.Store(
                    id="store-document-text",
                    data={"introduction": "", "conclusion": ""},
                ),  # Persistent intro/conclusion
                dcc.Store(
                    id="store-subplot-metadata", data={}
                ),  # {tab: {subplot: {title, caption, description}}}
                # NOTE: store-collapsed-csvs removed - collapse handled clientside via CSS
                dcc.Store(
                    id="store-time-offsets", data={}
                ),  # {csv_idx: offset} or {"csv_idx:signal_name": offset}
                dcc.Store(
                    id="store-csv-settings", data={}
                ),  # {csv_idx: {header_row: int, has_header: bool}}
                dcc.Store(
                    id="store-annotations", data={}
                ),  # {tab: {subplot: [{x, y, text, color, fontsize, arrow}]}}
                dcc.Store(
                    id="store-display-options", data={}
                ),  # {tab: {subplot: {markers: bool, normalize: bool}}}
                # PERFORMANCE: Cache figure to avoid rebuilds when only layout changes
                dcc.Store(id="store-figure-cache", data=None),
                dcc.Store(id="store-figure-hash", data=""),
                # PERFORMANCE: Track visible X range for viewport-based updates
                dcc.Store(id="store-visible-range", data={}),  # {tab: {xmin, xmax}}
                # PRESETS: Saved layout presets
                dcc.Store(id="store-presets", data={}, storage_type="local"),  # Persistent presets saved to localStorage
                # COMPARE MODE: Side-by-side plot comparison
                dcc.Store(id="store-compare-mode", data=False),  # Compare mode enabled/disabled
                dcc.Store(id="store-compare-tabs", data=[]),  # Tabs to compare
                # SIGNAL REORDER: Track drag & drop signal reorder events
                dcc.Store(id="store-signal-reorder", data=None),  # {fromIdx, toIdx}
                dcc.Download(id="download-session"),
                dcc.Download(id="download-template"),
                dcc.Download(id="download-csv-export"),
                dcc.Download(id="download-word-export"),
                dbc.Container(
                    [
                        # Header
                        dbc.Row(
                            [
                                dbc.Col(
                                    html.H4(
                                        "📊 Signal Viewer Pro",
                                        className="text-primary fw-bold mb-0",
                                    ),
                                    width=6,
                                ),
                                dbc.Col(
                                    [
                                        html.Div(
                                            [
                                                html.Span("☀️", className="me-1"),
                                                dbc.Switch(
                                                    id="theme-switch",
                                                    value=True,
                                                    className="d-inline",
                                                ),
                                                html.Span("🌙", className="ms-1"),
                                                dbc.Badge(
                                                    id="status-badge",
                                                    children="Ready",
                                                    color="info",
                                                    className="ms-3",
                                                ),
                                            ],
                                            className="d-flex align-items-center justify-content-end",
                                        )
                                    ],
                                    width=6,
                                ),
                            ],
                            className="py-2 mb-2 border-bottom",
                        ),
                        # Split trigger store and interval
                        dcc.Store(id="store-split-init", data=False),
                        dcc.Interval(
                            id="interval-split-init", interval=500, max_intervals=5
                        ),
                        # Main content - split layout
                        html.Div(
                            [
                                # Sidebar
                                html.Div(
                                    [
                                        # Panel 1: CSV Upload
                                        html.Div(
                                            [
                                                dbc.Card(
                                                    [
                                                        dbc.CardHeader(
                                                            [
                                                                html.Small(
                                                                    "📁 Data Sources",
                                                                    className="fw-bold",
                                                                ),
                                                                html.Div(
                                                                    [
                                                                        dbc.Button(
                                                                            "⏱",
                                                                            id="btn-time-cols",
                                                                            size="sm",
                                                                            color="info",
                                                                            outline=True,
                                                                            className="py-0 me-1",
                                                                            style={
                                                                                "fontSize": "10px"
                                                                            },
                                                                            title="Select time column",
                                                                        ),
                                                                        dbc.Button(
                                                                            "Clear",
                                                                            id="btn-clear-csv",
                                                                            size="sm",
                                                                            color="danger",
                                                                            outline=True,
                                                                            className="py-0",
                                                                            style={
                                                                                "fontSize": "10px"
                                                                            },
                                                                        ),
                                                                    ],
                                                                    className="float-end",
                                                                ),
                                                            ],
                                                            id="card-header-csv",
                                                            className="py-2",
                                                        ),
                                                        dbc.CardBody(
                                                            [
                                                                # Native file browser button
                                                                dbc.Button(
                                                                    [
                                                                        html.I(className="fas fa-folder-open me-2"),
                                                                        "Browse Files...",
                                                                    ],
                                                                    id="btn-browse-files",
                                                                    color="primary",
                                                                    className="w-100",
                                                                    size="sm",
                                                                ),
                                                                # Hidden upload for compatibility
                                                                html.Div([
                                                                    dcc.Upload(id="upload-csv", style={"display": "none"}),
                                                                ], style={"display": "none"}),
                                                                # CSV file list
                                                                html.Div(
                                                                    id="csv-list",
                                                                    className="mt-2",
                                                                    style={
                                                                        "maxHeight": "100px",
                                                                        "overflowY": "auto",
                                                                    },
                                                                ),
                                                                # Store for pending file paths from native dialog
                                                                dcc.Store(id="store-pending-paths", data=[]),
                                                                # Interval to poll for file dialog results
                                                                dcc.Interval(id="interval-file-check", interval=500, disabled=True),
                                                                # Hidden elements for removed features
                                                                html.Div([
                                                                    dbc.Input(id="input-csv-path", style={"display": "none"}),
                                                                    html.Button(id="btn-load-path", style={"display": "none"}),
                                                                ], style={"display": "none"}),
                                                            ],
                                                            id="card-body-csv",
                                                            className="py-2",
                                                        ),
                                                    ],
                                                    id="card-csv",
                                                ),
                                            ],
                                            id="split-panel-1",
                                            className="split-panel",
                                        ),
                                        # Panel 2: Signal Browser
                                        html.Div(
                                            [
                                                dbc.Card(
                                                    [
                                                        dbc.CardHeader(
                                                            [
                                                                html.Small(
                                                                    "📶 Signals",
                                                                    className="fw-bold",
                                                                ),
                                                                dbc.Button(
                                                                    "🔗",
                                                                    id="btn-link",
                                                                    size="sm",
                                                                    color="info",
                                                                    outline=True,
                                                                    title="Link CSVs",
                                                                    className="float-end ms-1",
                                                                ),
                                                                dbc.Button(
                                                                    "📊",
                                                                    id="btn-compare-csvs",
                                                                    size="sm",
                                                                    color="warning",
                                                                    outline=True,
                                                                    title="Compare CSVs",
                                                                    className="float-end",
                                                                ),
                                                            ],
                                                            id="card-header-signals",
                                                            className="py-2",
                                                        ),
                                                        dbc.CardBody(
                                                            [
                                                                # Search with autocomplete and filter list
                                                                html.Div(
                                                                    [
                                                                        dbc.InputGroup(
                                                                            [
                                                                                dbc.Input(
                                                                                    id="search-input",
                                                                                    placeholder="Search signals... (type to see suggestions)",
                                                                                    size="sm",
                                                                                    debounce=False,  # Instant for autocomplete
                                                                                    autocomplete="off",
                                                                                ),
                                                                                dbc.Button(
                                                                                    "+",
                                                                                    id="btn-add-filter",
                                                                                    size="sm",
                                                                                    color="success",
                                                                                    outline=True,
                                                                                    title="Add to filter list",
                                                                                ),
                                                                            ],
                                                                            size="sm",
                                                                        ),
                                                                        # Autocomplete suggestions dropdown
                                                                        html.Div(
                                                                            id="search-autocomplete-dropdown",
                                                                            style={
                                                                                "display": "none",
                                                                                "position": "absolute",
                                                                                "zIndex": 1000,
                                                                                "backgroundColor": "#1a1a2e",
                                                                                "border": "1px solid #4ea8de",
                                                                                "borderRadius": "4px",
                                                                                "maxHeight": "200px",
                                                                                "overflowY": "auto",
                                                                                "width": "100%",
                                                                                "boxShadow": "0 4px 12px rgba(0,0,0,0.3)",
                                                                            },
                                                                        ),
                                                                        # Active filters display
                                                                        html.Div(
                                                                            id="filter-list-display",
                                                                            className="mt-1",
                                                                            style={
                                                                                "fontSize": "10px"
                                                                            },
                                                                        ),
                                                                    ],
                                                                    className="mb-2",
                                                                ),
                                                                html.Div(
                                                                    [
                                                                        html.Small(
                                                                            "Assign → ",
                                                                            className="text-muted",
                                                                        ),
                                                                        html.Span(
                                                                            id="target-info",
                                                                            className="text-info fw-bold small",
                                                                        ),
                                                                    ],
                                                                    className="mb-1 p-1 rounded",
                                                                    id="target-box",
                                                                ),
                                                                html.Div(
                                                                    [
                                                                        html.Small(
                                                                            "Selected for ops: ",
                                                                            className="text-muted",
                                                                        ),
                                                                        html.Span(
                                                                            id="highlight-count",
                                                                            children="0",
                                                                            className="text-warning fw-bold",
                                                                        ),
                                                                        dbc.Button(
                                                                            "⚙️ Operate",
                                                                            id="btn-operate-selected",
                                                                            size="sm",
                                                                            color="warning",
                                                                            outline=True,
                                                                            className="ms-2",
                                                                            style={
                                                                                "fontSize": "10px"
                                                                            },
                                                                        ),
                                                                        dbc.Button(
                                                                            "Clear",
                                                                            id="btn-clear-highlight",
                                                                            size="sm",
                                                                            color="secondary",
                                                                            outline=True,
                                                                            className="ms-1",
                                                                            style={
                                                                                "fontSize": "10px"
                                                                            },
                                                                        ),
                                                                    ],
                                                                    className="mb-2",
                                                                ),
                                                                html.Div(
                                                                    id="signal-tree",
                                                                    children=[
                                                                        html.Span(
                                                                            "Loading...",
                                                                            className="text-muted small",
                                                                        )
                                                                    ],
                                                                    style={
                                                                        "flex": "1",
                                                                        "overflowY": "auto",
                                                                    },
                                                                ),
                                                            ],
                                                            id="card-body-signals",
                                                            className="py-2",
                                                        ),
                                                    ],
                                                    id="card-signals",
                                                ),
                                            ],
                                            id="split-panel-2",
                                            className="split-panel",
                                        ),
                                        # Hidden elements for callback compatibility
                                        html.Div(
                                            [
                                                html.Div(
                                                    id="derived-list",
                                                    style={"display": "none"},
                                                ),
                                                html.Div(
                                                    id="btn-clear-derived",
                                                    style={"display": "none"},
                                                ),
                                            ],
                                            style={"display": "none"},
                                        ),
                                        # Panel 3: Assigned
                                        html.Div(
                                            [
                                                dbc.Card(
                                                    [
                                                            dbc.CardHeader(
                                                                [
                                                                    html.Span("📋", className="me-2", style={"fontSize": "14px"}),
                                                                    html.Small(
                                                                        "Assigned",
                                                                        className="fw-bold",
                                                                        style={"fontSize": "12px", "letterSpacing": "0.3px"},
                                                                    ),
                                                                ],
                                                                id="card-header-assigned",
                                                                className="py-2",
                                                            ),
                                                        dbc.CardBody(
                                                            [
                                                                # Mode toggle: Time vs X-Y
                                                                html.Div(
                                                                    [
                                                                        dbc.RadioItems(
                                                                            id="subplot-mode-toggle",
                                                                            options=[
                                                                                {
                                                                                    "label": "📈 Time",
                                                                                    "value": "time",
                                                                                },
                                                                                {
                                                                                    "label": "⚡ X-Y",
                                                                                    "value": "xy",
                                                                                },
                                                                            ],
                                                                            value="time",
                                                                            inline=True,
                                                                            className="small",
                                                                        ),
                                                                    ],
                                                                    className="mb-2 text-center",
                                                                ),
                                                                # X-axis selector (shown in xy mode - select X signal, then add Y signals normally)
                                                                html.Div(
                                                                    [
                                                                        html.Small(
                                                                            "X-Axis: ",
                                                                            className="text-info fw-bold me-1",
                                                                        ),
                                                                        dbc.Select(
                                                                            id="xy-x-select",
                                                                            size="sm",
                                                                            options=[
                                                                                {
                                                                                    "label": "⏱ Time (default)",
                                                                                    "value": "time",
                                                                                }
                                                                            ],
                                                                            value="time",
                                                                            style={
                                                                                "fontSize": "10px",
                                                                                "flex": "1",
                                                                                "minWidth": "150px",
                                                                                "maxWidth": "250px",
                                                                            },
                                                                        ),
                                                                        # Hidden Y select for callback compatibility
                                                                        dbc.Select(
                                                                            id="xy-y-select",
                                                                            style={
                                                                                "display": "none"
                                                                            },
                                                                        ),
                                                                    ],
                                                                    id="xy-controls",
                                                                    style={
                                                                        "display": "none"
                                                                    },
                                                                    className="mb-2 d-flex align-items-center",
                                                                ),
                                                                # Hidden elements for callback compatibility
                                                                html.Div(
                                                                    id="xy-x-signal",
                                                                    style={
                                                                        "display": "none"
                                                                    },
                                                                ),
                                                                html.Div(
                                                                    id="xy-y-signal",
                                                                    style={
                                                                        "display": "none"
                                                                    },
                                                                ),
                                                                # Subplot metadata (title, caption, description)
                                                                html.Div(
                                                                    [
                                                                        dbc.Input(
                                                                            id="subplot-title-input",
                                                                            placeholder="Plot title (replaces 'Subplot X')...",
                                                                            size="sm",
                                                                            className="mb-1",
                                                                            style={
                                                                                "fontSize": "10px"
                                                                            },
                                                                        ),
                                                                        dcc.Textarea(
                                                                            id="subplot-caption-input",
                                                                            placeholder="Caption (under Fig #)...",
                                                                            className="form-control mb-1 rtl-textarea",
                                                                            style={
                                                                                "fontSize": "10px",
                                                                                "height": "45px",
                                                                                "resize": "vertical",
                                                                                "direction": "auto",
                                                                                "unicodeBidi": "plaintext",
                                                                            },
                                                                        ),
                                                                        dcc.Textarea(
                                                                            id="subplot-description-input",
                                                                            placeholder="Description (detailed text after plot)...",
                                                                            className="form-control rtl-textarea",
                                                                            style={
                                                                                "fontSize": "10px",
                                                                                "height": "80px",
                                                                                "resize": "vertical",
                                                                                "direction": "auto",
                                                                                "unicodeBidi": "plaintext",
                                                                            },
                                                                        ),
                                                                    ],
                                                                    className="mb-2 border-bottom pb-2",
                                                                ),
                                                                html.Div(
                                                                    id="assigned-list",
                                                                    style={
                                                                        "overflowY": "auto",
                                                                        "flex": "1",
                                                                    },
                                                                ),
                                                                # Quick Statistics Panel (collapsible)
                                                                dbc.Collapse(
                                                                    dbc.Card(
                                                                        dbc.CardBody(
                                                                            html.Div(
                                                                                id="stats-panel-content",
                                                                                className="small",
                                                                            ),
                                                                            className="p-2",
                                                                        ),
                                                                        className="mb-2 border-info",
                                                                        style={"fontSize": "10px"},
                                                                    ),
                                                                    id="stats-panel-collapse",
                                                                    is_open=False,
                                                                ),
                                                                # Signal Display Options
                                                                html.Div(
                                                                    [
                                                                        dbc.Checklist(
                                                                            id="signal-display-options",
                                                                            options=[
                                                                                {"label": "📊 Stats", "value": "stats"},
                                                                                {"label": "⚫ Markers", "value": "markers"},
                                                                                {"label": "📏 Normalize", "value": "normalize"},
                                                                            ],
                                                                            value=[],
                                                                            inline=True,
                                                                            switch=True,
                                                                            className="small",
                                                                            style={"fontSize": "10px"},
                                                                        ),
                                                                    ],
                                                                    className="mb-2 border-top pt-2",
                                                                ),
                                                                dbc.Button(
                                                                    "🗑️ Remove",
                                                                    id="btn-remove",
                                                                    size="sm",
                                                                    color="danger",
                                                                    className="w-100 mt-2",
                                                                ),
                                                            ],
                                                            id="card-body-assigned",
                                                            className="py-2",
                                                        ),
                                                    ],
                                                    id="card-assigned",
                                                ),
                                                html.Div(
                                                    id="status-text",
                                                    className="text-center small mt-1",
                                                    style={"fontSize": "10px"},
                                                ),
                                            ],
                                            id="split-panel-3",
                                            className="split-panel",
                                        ),
                                    ],
                                    id="split-sidebar",
                                    className="split-sidebar",
                                ),
                                # Plot Area
                                html.Div(
                                    [
                                        dbc.Card(
                                            [
                                                dbc.CardHeader(
                                                    [
                                                        dbc.Row(
                                                            [
                                                                dbc.Col(
                                                                    [
                                                                        dbc.ButtonGroup(
                                                                            [
                                                                                # Session management
                                                                                dbc.Button(
                                                                                    "💾",
                                                                                    id="btn-save",
                                                                                    size="sm",
                                                                                    color="success",
                                                                                    title="Save session",
                                                                                    className="px-2",
                                                                                ),
                                                                                dcc.Upload(
                                                                                    id="upload-session",
                                                                                    children=dbc.Button(
                                                                                        "📂",
                                                                                        size="sm",
                                                                                        color="info",
                                                                                        title="Load session",
                                                                                        className="px-2",
                                                                                    ),
                                                                                    accept=".json",
                                                                                ),
                                                                                # Preset Templates dropdown
                                                                                dbc.DropdownMenu(
                                                                                    [
                                                                                        dbc.DropdownMenuItem("📋 Save as Preset", id="btn-save-preset"),
                                                                                        dbc.DropdownMenuItem(divider=True),
                                                                                        dbc.DropdownMenuItem("Empty Presets", header=True, id="preset-list-header"),
                                                                                        # Dynamic presets will be added here via callback
                                                                                    ],
                                                                                    id="dropdown-presets",
                                                                                    label="📑",
                                                                                    size="sm",
                                                                                    color="secondary",
                                                                                    className="d-inline-block",
                                                                                    toggle_style={"padding": "4px 8px"},
                                                                                ),
                                                                                # Compare Mode toggle (disabled - needs layout fix)
                                                                                # dbc.Button(
                                                                                #     "⚖️",
                                                                                #     id="btn-compare-mode",
                                                                                #     size="sm",
                                                                                #     color="secondary",
                                                                                #     outline=True,
                                                                                #     title="Compare Mode (side-by-side plots)",
                                                                                #     className="px-2",
                                                                                # ),
                                                                                html.Span(id="btn-compare-mode", style={"display": "none"}),  # Placeholder
                                                                                # Export buttons
                                                                                dbc.Button(
                                                                                    "📊",
                                                                                    id="btn-export-csv",
                                                                                    size="sm",
                                                                                    color="secondary",
                                                                                    outline=True,
                                                                                    title="Export signals to CSV",
                                                                                    className="px-2",
                                                                                ),
                                                                                dbc.Button(
                                                                                    "📑",
                                                                                    id="btn-export-pdf",
                                                                                    size="sm",
                                                                                    color="secondary",
                                                                                    outline=True,
                                                                                    title="Export to HTML",
                                                                                    className="px-2",
                                                                                ),
                                                                                dbc.Button(
                                                                                    "📝",
                                                                                    id="btn-export-word",
                                                                                    size="sm",
                                                                                    color="secondary",
                                                                                    outline=True,
                                                                                    title="Export to Word",
                                                                                    className="px-2",
                                                                                ),
                                                                                dbc.Button(
                                                                                    "📌",
                                                                                    id="btn-open-annotation",
                                                                                    size="sm",
                                                                                    color="info",
                                                                                    outline=True,
                                                                                    title="Add annotation",
                                                                                    className="px-2",
                                                                                ),
                                                                                # Hidden elements for callback compatibility
                                                                                html.Div(
                                                                                    id="btn-load",
                                                                                    style={
                                                                                        "display": "none"
                                                                                    },
                                                                                ),
                                                                                html.Div(
                                                                                    id="btn-del-tab",
                                                                                    style={
                                                                                        "display": "none"
                                                                                    },
                                                                                ),
                                                                                dbc.Button(
                                                                                    "🔄",
                                                                                    id="btn-refresh-csv",
                                                                                    size="sm",
                                                                                    color="secondary",
                                                                                    title="Refresh CSVs",
                                                                                    className="px-2",
                                                                                ),
                                                                                dbc.Button(
                                                                                    "▶ Stream",
                                                                                    id="btn-stream-csv",
                                                                                    size="sm",
                                                                                    color="success",
                                                                                    outline=True,
                                                                                    title="Stream CSV updates",
                                                                                    className="px-2",
                                                                                    style={'whiteSpace': 'nowrap'},
                                                                                ),
                                                                            ],
                                                                            size="sm",
                                                                        )
                                                                    ],
                                                                    width=4,
                                                                ),
                                                                dbc.Col(
                                                                    [
                                                                        html.Div(
                                                                            [
                                                                                dbc.Input(
                                                                                    id="rows-input",
                                                                                    type="number",
                                                                                    value=1,
                                                                                    min=1,
                                                                                    max=4,
                                                                                    size="sm",
                                                                                    style={
                                                                                        "width": "40px",
                                                                                        "display": "inline-block",
                                                                                    },
                                                                                ),
                                                                                html.Span(
                                                                                    "×",
                                                                                    className="mx-1",
                                                                                ),
                                                                                dbc.Input(
                                                                                    id="cols-input",
                                                                                    type="number",
                                                                                    value=1,
                                                                                    min=1,
                                                                                    max=4,
                                                                                    size="sm",
                                                                                    style={
                                                                                        "width": "40px",
                                                                                        "display": "inline-block",
                                                                                    },
                                                                                ),
                                                                                # Hidden subplot selector (needed for callbacks)
                                                                                dbc.Select(
                                                                                    id="subplot-select",
                                                                                    style={
                                                                                        "display": "none"
                                                                                    },
                                                                                ),
                                                                                dbc.Input(
                                                                                    id="subplot-input",
                                                                                    type="number",
                                                                                    style={
                                                                                        "display": "none"
                                                                                    },
                                                                                ),
                                                                                dbc.Checkbox(
                                                                                    id="link-axes-check",
                                                                                    label="Link",
                                                                                    value=False,
                                                                                    className="ms-2 small",
                                                                                ),
                                                                                    dbc.Checkbox(
                                                                                    id="time-cursor-check",
                                                                                    label="Cursor",
                                                                                    value=True,
                                                                                    className="ms-2 small",
                                                                                ),
                                                                                # Cursor playback controls (beside layout buttons)
                                                                                html.Div(
                                                                                    [
                                                                                        dbc.Button(
                                                                                            id="btn-cursor-play-header",
                                                                                            children="▶",
                                                                                            size="sm",
                                                                                            color="success",
                                                                                            outline=True,
                                                                                            className="ms-2 px-2",
                                                                                            title="Play cursor animation",
                                                                                            style={"fontFamily": "system-ui"},
                                                                                        ),
                                                                                        dbc.Button(
                                                                                            id="btn-cursor-stop-header",
                                                                                            children="⏹",
                                                                                            size="sm",
                                                                                            color="danger",
                                                                                            outline=True,
                                                                                            className="px-2",
                                                                                            title="Stop cursor animation",
                                                                                            style={"fontFamily": "system-ui"},
                                                                                        ),
                                                                                        dcc.Slider(
                                                                                            id="cursor-slider-header",
                                                                                            min=0,
                                                                                            max=100,
                                                                                            value=50,
                                                                                            step=0.5,
                                                                                            marks=None,
                                                                                            tooltip={"placement": "bottom", "always_visible": False},
                                                                                            className="cursor-slider-header ms-2",
                                                                                            disabled=False,
                                                                                        ),
                                                                                    ],
                                                                                    id="cursor-controls-header",
                                                                                    className="d-flex align-items-center",
                                                                                    style={"display": "none", "minWidth": "150px"},  # Hidden by default
                                                                                ),
                                                                            ],
                                                                            className="d-flex align-items-center justify-content-end",
                                                                        )
                                                                    ],
                                                                    width=8,
                                                                ),
                                                            ]
                                                        )
                                                    ],
                                                    id="card-header-plot",
                                                    className="py-2",
                                                ),
                                                dbc.CardBody(
                                                    [
                                                        # Browser-style tabs with + button
                                                        html.Div(
                                                            [
                                                                html.Div(
                                                                    id="tabs-container",
                                                                    className="d-flex align-items-center",
                                                                    style={
                                                                        "flexWrap": "wrap",
                                                                        "gap": "2px",
                                                                    },
                                                                    children=[
                                                                        # Tab buttons will be dynamically generated
                                                                    ],
                                                                ),
                                                                dbc.Button(
                                                                    "+",
                                                                    id="btn-add-tab",
                                                                    size="sm",
                                                                    color="primary",
                                                                    outline=True,
                                                                    className="ms-1 px-2",
                                                                    title="Add new tab",
                                                                    style={
                                                                        "fontSize": "12px",
                                                                        "lineHeight": "1",
                                                                    },
                                                                ),
                                                            ],
                                                            className="d-flex align-items-center mb-2",
                                                        ),
                                                        # Hidden dcc.Tabs for state management
                                                        dcc.Tabs(
                                                            id="tabs",
                                                            value="tab-0",
                                                            children=[
                                                                dcc.Tab(
                                                                    label="Tab 1",
                                                                    value="tab-0",
                                                                )
                                                            ],
                                                            style={"display": "none"},
                                                        ),
                                                        # Direct Graph - no Spinner to avoid white flash
                                                        # Always keep a valid figure structure to prevent blank screen
                                                        dcc.Graph(
                                                            id="plot",
                                                            # PERFORMANCE: Config for fast rendering
                                                            config={
                                                                "displayModeBar": True,
                                                                "displaylogo": False,
                                                                "scrollZoom": True,
                                                                "modeBarButtonsToRemove": ["lasso2d", "select2d"],
                                                                "doubleClick": "reset+autosize",
                                                                "showTips": False,
                                                                # CRITICAL: Preserve figure during updates
                                                                "plotGlPixelRatio": 2,  # Better WebGL quality
                                                            },
                                                            figure={
                                                                "data": [{
                                                                    "x": [0],
                                                                    "y": [0],
                                                                    "mode": "markers",
                                                                    "marker": {"size": 1, "opacity": 0},
                                                                    "showlegend": False,
                                                                    "hoverinfo": "skip",
                                                                }],  # Invisible placeholder to prevent blank
                                                                "layout": {
                                                                    "xaxis": {
                                                                        "title": "Time",
                                                                        "range": [0, 1],
                                                                    },
                                                                    "yaxis": {
                                                                        "title": "Value",
                                                                        "range": [0, 1],
                                                                    },
                                                                    "template": "plotly_dark",
                                                                    "paper_bgcolor": "#16213e",
                                                                    "plot_bgcolor": "#1a1a2e",
                                                                    "font": {
                                                                        "color": "#e8e8e8"
                                                                    },
                                                                    "margin": {
                                                                        "l": 50,
                                                                        "r": 20,
                                                                        "t": 30,
                                                                        "b": 40,
                                                                    },
                                                                    # CRITICAL: Preserve zoom/pan state during updates
                                                                    "uirevision": True,  # Use True instead of "constant" for better state preservation
                                                                },
                                                            },
                                                            style={"height": "100%"},
                                                        ),
                                                        # Modern: Slider cursor control at bottom of plot with play/stop
                                                        html.Div(
                                                            [
                                                                html.Div(
                                                                    [
                                                                        html.Small("Time Cursor: ", className="text-muted me-2"),
                                                                        dbc.Button(
                                                                            id="btn-cursor-play",
                                                                            children="▶",
                                                                            size="sm",
                                                                            style={"fontFamily": "system-ui"},
                                                                            color="primary",
                                                                            outline=True,
                                                                            className="me-2",
                                                                        ),
                                                                        dbc.Button(
                                                                            id="btn-cursor-stop",
                                                                            children="⏹",
                                                                            size="sm",
                                                                            color="secondary",
                                                                            outline=True,
                                                                            className="me-2",
                                                                            style={"fontFamily": "system-ui"},
                                                                        ),
                                                                        html.Small(id="cursor-slider-value", className="text-muted me-2", children="--"),
                                                                    ],
                                                                    className="d-flex align-items-center mb-2",
                                                                ),
                                                                dcc.Slider(
                                                                    id="cursor-slider",
                                                                    min=0,
                                                                    max=100,
                                                                    value=50,
                                                                    step=0.001,  # High resolution stepping
                                                                    marks=None,
                                                                    tooltip={"placement": "bottom", "always_visible": True},
                                                                    disabled=True,  # Enabled when cursor is active
                                                                    className="cursor-slider",
                                                                    updatemode="drag",  # Update on drag for smooth control
                                                                ),
                                                                dcc.Interval(
                                                                    id="interval-cursor-animation",
                                                                    interval=100,  # Update every 100ms
                                                                    disabled=True,
                                                                ),
                                                            ],
                                                            id="cursor-slider-container",
                                                            style={
                                                                "display": "none",  # Hidden by default, shown when cursor enabled
                                                                "padding": "5px 10px",
                                                                "backgroundColor": "rgba(0,0,0,0.05)",
                                                                "borderRadius": "4px",
                                                                "marginTop": "5px",
                                                            },
                                                        ),
                                                    ],
                                                    id="card-body-plot",
                                                    className="p-2",
                                                ),
                                            ],
                                            id="card-plot",
                                        )
                                    ],
                                    id="split-plot",
                                    className="split-plot",
                                ),
                            ],
                            id="split-container",
                            className="split-container",
                        ),
                    ],
                    fluid=True,
                ),
                # Modals
                self.create_link_modal(),
                self.create_props_modal(),
                self.create_ops_modal(),
                self.create_multi_ops_modal(),
                self.create_csv_export_modal(),
                self.create_pdf_export_modal(),
                self.create_word_export_modal(),
                self.create_preset_modal(),
                self.create_time_column_modal(),
                self.create_annotation_modal(),
                self.create_compare_modal(),
                self.create_compare_mode_panel(),
            ],
        )

    def create_link_modal(self):
        return dbc.Modal(
            [
                dbc.ModalHeader("🔗 Link CSV Files"),
                dbc.ModalBody(
                    [
                        html.P(
                            "Linked CSVs: same signal names auto-assign/unassign together",
                            className="small text-muted",
                        ),
                        dbc.Checklist(
                            id="link-checks", options=[], value=[], switch=True
                        ),
                        dbc.Input(
                            id="link-name",
                            placeholder="Group name",
                            size="sm",
                            className="mt-2",
                        ),
                    ]
                ),
                dbc.ModalFooter(
                    [
                        dbc.Button(
                            "Create", id="btn-create-link", color="primary", size="sm"
                        ),
                        dbc.Button(
                            "Close", id="btn-close-link", color="secondary", size="sm"
                        ),
                    ]
                ),
            ],
            id="modal-link",
            is_open=False,
        )

    def create_props_modal(self):
        return dbc.Modal(
            [
                dbc.ModalHeader(id="props-title", children="Signal Properties"),
                dbc.ModalBody(
                    [
                        dbc.Label("Display Name:", className="small"),
                        dbc.Input(id="prop-name", size="sm", className="mb-2"),
                        dbc.Label("Scale Factor:", className="small"),
                        dbc.Input(
                            id="prop-scale",
                            type="number",
                            value=1.0,
                            size="sm",
                            className="mb-2",
                        ),
                        dbc.Label("Color:", className="small"),
                        dbc.Input(
                            id="prop-color",
                            type="color",
                            value="#2E86AB",
                            className="mb-2",
                            style={"height": "35px"},
                        ),
                        dbc.Label("Line Width:", className="small"),
                        dbc.Input(
                            id="prop-width",
                            type="number",
                            value=1.5,
                            min=0.5,
                            max=5,
                            step=0.5,
                            size="sm",
                            className="mb-2",
                        ),
                        dbc.Label("Time Offset (seconds):", className="small"),
                        dbc.InputGroup(
                            [
                                dbc.Input(
                                    id="prop-time-offset",
                                    type="number",
                                    value=0,
                                    step=0.001,
                                    size="sm",
                                    placeholder="0.0",
                                ),
                                dbc.InputGroupText("sec", style={"fontSize": "11px"}),
                            ],
                            size="sm",
                            className="mb-2",
                        ),
                        html.Small(
                            "Positive = shift right, Negative = shift left",
                            className="text-muted d-block mb-2",
                        ),
                        dbc.Checkbox(
                            id="prop-apply-tree",
                            label="Show name in tree",
                            value=True,
                            className="mt-2",
                        ),
                        dbc.Checkbox(
                            id="prop-state-signal",
                            label="State signal (step display)",
                            value=False,
                            className="mt-2",
                        ),
                    ]
                ),
                dbc.ModalFooter(
                    [
                        dbc.Button(
                            "Apply", id="btn-apply-props", color="primary", size="sm"
                        ),
                        dbc.Button(
                            "Close", id="btn-close-props", color="secondary", size="sm"
                        ),
                    ]
                ),
            ],
            id="modal-props",
            is_open=False,
        )

    def create_ops_modal(self):
        return dbc.Modal(
            [
                dbc.ModalHeader(id="ops-title", children="Signal Operation"),
                dbc.ModalBody(
                    [
                        dbc.Label("Operation:", className="small"),
                        dbc.Select(
                            id="op-type",
                            size="sm",
                            className="mb-2",
                            options=[
                                {"label": "∂ Derivative", "value": "derivative"},
                                {"label": "∫ Integral", "value": "integral"},
                                {"label": "|x| Absolute", "value": "abs"},
                                {"label": "√x Square Root", "value": "sqrt"},
                                {"label": "-x Negate", "value": "negate"},
                            ],
                            value="derivative",
                        ),
                        dbc.Label("Result Name:", className="small"),
                        dbc.Input(id="op-result-name", size="sm", className="mb-2"),
                    ]
                ),
                dbc.ModalFooter(
                    [
                        dbc.Button(
                            "Compute", id="btn-compute-op", color="primary", size="sm"
                        ),
                        dbc.Button(
                            "Close", id="btn-close-ops", color="secondary", size="sm"
                        ),
                    ]
                ),
            ],
            id="modal-ops",
            is_open=False,
        )

    def create_multi_ops_modal(self):
        return dbc.Modal(
            [
                dbc.ModalHeader("⚙️ Operations on Selected Signals"),
                dbc.ModalBody(
                    [
                        html.Div(id="selected-signals-info", className="mb-2"),
                        dbc.Label("Operation:", className="small"),
                        dbc.Select(
                            id="multi-op-type",
                            size="sm",
                            className="mb-2",
                            options=[
                                {"label": "A + B", "value": "add"},
                                {"label": "A - B", "value": "sub"},
                                {"label": "A × B", "value": "mul"},
                                {"label": "A ÷ B", "value": "div"},
                                {"label": "||signals|| Norm", "value": "norm"},
                                {"label": "Mean", "value": "mean"},
                            ],
                            value="add",
                        ),
                        dbc.Label("Result Name:", className="small"),
                        dbc.Input(id="multi-op-name", size="sm", placeholder="result"),
                    ]
                ),
                dbc.ModalFooter(
                    [
                        dbc.Button(
                            "Compute",
                            id="btn-compute-multi",
                            color="primary",
                            size="sm",
                        ),
                        dbc.Button(
                            "Close", id="btn-close-multi", color="secondary", size="sm"
                        ),
                    ]
                ),
            ],
            id="modal-multi-ops",
            is_open=False,
        )

    def create_csv_export_modal(self):
        """Modal for exporting signals to CSV."""
        return dbc.Modal(
            [
                dbc.ModalHeader("📊 Export Signals to CSV"),
                dbc.ModalBody(
                    [
                        dbc.Label("Export Scope:", className="small"),
                        dbc.RadioItems(
                            id="export-csv-scope",
                            options=[
                                {
                                    "label": "Current subplot signals",
                                    "value": "subplot",
                                },
                                {"label": "All signals in current tab", "value": "tab"},
                                {"label": "All signals in all tabs", "value": "all"},
                            ],
                            value="subplot",
                            className="mb-2",
                        ),
                        dbc.Checkbox(
                            id="export-csv-include-time",
                            label="Include time column",
                            value=True,
                            className="mb-2",
                        ),
                    ]
                ),
                dbc.ModalFooter(
                    [
                        dbc.Button(
                            "Export",
                            id="btn-do-export-csv",
                            color="primary",
                            size="sm",
                        ),
                        dbc.Button(
                            "Close",
                            id="btn-close-export-csv",
                            color="secondary",
                            size="sm",
                        ),
                    ]
                ),
            ],
            id="modal-export-csv",
            is_open=False,
        )

    def create_pdf_export_modal(self):
        """Modal for exporting plots to PDF."""
        return dbc.Modal(
            [
                dbc.ModalHeader("📑 Export Plots to PDF"),
                dbc.ModalBody(
                    [
                        dbc.Label("Export Scope:", className="small fw-bold"),
                        dbc.RadioItems(
                            id="export-pdf-scope",
                            options=[
                                {"label": "Current subplot", "value": "subplot"},
                                {
                                    "label": "All subplots in current tab",
                                    "value": "tab",
                                },
                                {"label": "All tabs", "value": "all"},
                            ],
                            value="tab",
                            className="mb-3",
                        ),
                        dbc.Label("Report Title:", className="small fw-bold"),
                        dbc.Input(
                            id="export-pdf-title",
                            placeholder="Signal Analysis Report",
                            size="sm",
                            className="mb-3",
                        ),
                        html.Hr(),
                        html.Small(
                            "📝 Document text (stored in app):", className="text-muted"
                        ),
                        dbc.Label("Introduction:", className="small mt-2"),
                        dcc.Textarea(
                            id="export-pdf-intro",
                            placeholder="Introduction text for the report...",
                            className="form-control mb-2 rtl-textarea",
                            style={
                                "height": "100px",
                                "resize": "vertical",
                                "fontFamily": "monospace",
                                "padding": "8px",
                                "direction": "auto",
                                "unicodeBidi": "plaintext",
                            },
                        ),
                        dbc.Label("Conclusion:", className="small"),
                        dcc.Textarea(
                            id="export-pdf-conclusion",
                            placeholder="Conclusion text for the report...",
                            className="form-control mb-2 rtl-textarea",
                            style={
                                "height": "100px",
                                "resize": "vertical",
                                "fontFamily": "monospace",
                                "padding": "8px",
                                "direction": "auto",
                                "unicodeBidi": "plaintext",
                            },
                        ),
                        dbc.Button(
                            "💾 Save Document Text",
                            id="btn-save-doc-text",
                            size="sm",
                            color="info",
                            outline=True,
                            className="mb-2",
                        ),
                        html.Small(
                            "Note: Subplot titles/captions are set in the Assigned panel for each subplot.",
                            className="text-muted d-block mt-2",
                        ),
                    ]
                ),
                dbc.ModalFooter(
                    [
                        dbc.Button(
                            "Export PDF",
                            id="btn-do-export-pdf",
                            color="primary",
                            size="sm",
                        ),
                        dbc.Button(
                            "Close",
                            id="btn-close-export-pdf",
                            color="secondary",
                            size="sm",
                        ),
                    ]
                ),
            ],
            id="modal-export-pdf",
            is_open=False,
            size="lg",
        )

    def create_word_export_modal(self):
        """Modal for exporting plots to Word document."""
        return dbc.Modal(
            [
                dbc.ModalHeader("📝 Export to Word Document"),
                dbc.ModalBody(
                    [
                        dbc.Label("Export Scope:", className="small fw-bold"),
                        dbc.RadioItems(
                            id="export-word-scope",
                            options=[
                                {"label": "Current subplot", "value": "subplot"},
                                {"label": "All subplots in current tab", "value": "tab"},
                                {"label": "All tabs", "value": "all"},
                            ],
                            value="tab",
                            className="mb-3",
                        ),
                        dbc.Label("Report Title:", className="small fw-bold"),
                        dbc.Input(
                            id="export-word-title",
                            placeholder="Signal Analysis Report",
                            size="sm",
                            className="mb-3",
                        ),
                        html.Hr(),
                        dbc.Label("Introduction:", className="small"),
                        dcc.Textarea(
                            id="export-word-intro",
                            placeholder="Introduction text for the report...",
                            className="form-control mb-2",
                            style={"height": "80px", "resize": "vertical"},
                        ),
                        dbc.Label("Conclusion:", className="small"),
                        dcc.Textarea(
                            id="export-word-conclusion",
                            placeholder="Conclusion text for the report...",
                            className="form-control mb-2",
                            style={"height": "80px", "resize": "vertical"},
                        ),
                        html.Small(
                            "Note: Plots will be exported as images embedded in the Word document.",
                            className="text-muted d-block mt-2",
                        ),
                    ]
                ),
                dbc.ModalFooter(
                    [
                        dbc.Button(
                            "Export Word",
                            id="btn-do-export-word",
                            color="primary",
                            size="sm",
                        ),
                        dbc.Button(
                            "Close",
                            id="btn-close-export-word",
                            color="secondary",
                            size="sm",
                        ),
                    ]
                ),
            ],
            id="modal-export-word",
            is_open=False,
            size="lg",
        )

    def create_preset_modal(self):
        """Modal for saving layout presets."""
        return dbc.Modal(
            [
                dbc.ModalHeader("📋 Save Layout Preset"),
                dbc.ModalBody(
                    [
                        dbc.Label("Preset Name:", className="fw-bold"),
                        dbc.Input(
                            id="preset-name-input",
                            placeholder="My Layout Preset",
                            size="sm",
                            className="mb-3",
                        ),
                        dbc.Label("Description (optional):", className="small"),
                        dcc.Textarea(
                            id="preset-description-input",
                            placeholder="Describe this preset...",
                            className="form-control mb-3",
                            style={"height": "60px", "resize": "vertical"},
                        ),
                        dbc.Checkbox(
                            id="preset-include-signals",
                            label="Include assigned signals",
                            value=True,
                            className="mb-2",
                        ),
                        dbc.Checkbox(
                            id="preset-include-layout",
                            label="Include layout (rows/cols)",
                            value=True,
                            className="mb-2",
                        ),
                        html.Small(
                            "Presets are saved locally in your browser.",
                            className="text-muted d-block mt-2",
                        ),
                    ]
                ),
                dbc.ModalFooter(
                    [
                        dbc.Button(
                            "Save Preset",
                            id="btn-do-save-preset",
                            color="primary",
                            size="sm",
                        ),
                        dbc.Button(
                            "Cancel",
                            id="btn-close-preset-modal",
                            color="secondary",
                            size="sm",
                        ),
                    ]
                ),
            ],
            id="modal-preset",
            is_open=False,
            size="md",
        )

    def create_compare_mode_panel(self):
        """Panel for compare mode - shows two plots side by side."""
        return html.Div(
            [
                html.Div(
                    [
                        html.H6("⚖️ Compare Mode", className="text-primary mb-2"),
                        html.P(
                            "Select two tabs to compare side-by-side:",
                            className="small text-muted mb-2",
                        ),
                        dbc.Row(
                            [
                                dbc.Col(
                                    dbc.Select(
                                        id="compare-tab-left",
                                        placeholder="Left plot...",
                                        size="sm",
                                    ),
                                    width=5,
                                ),
                                dbc.Col(
                                    html.Span("vs", className="text-muted"),
                                    width=2,
                                    className="text-center",
                                ),
                                dbc.Col(
                                    dbc.Select(
                                        id="compare-tab-right",
                                        placeholder="Right plot...",
                                        size="sm",
                                    ),
                                    width=5,
                                ),
                            ],
                            className="mb-2",
                        ),
                        dbc.Button(
                            "Exit Compare Mode",
                            id="btn-exit-compare",
                            size="sm",
                            color="secondary",
                            outline=True,
                            className="w-100",
                        ),
                    ],
                    className="p-3 border rounded mb-2",
                    style={"backgroundColor": "rgba(78, 168, 222, 0.1)"},
                ),
                # Side-by-side plots container
                html.Div(
                    [
                        html.Div(
                            dcc.Graph(id="compare-plot-left", style={"height": "100%"}),
                            className="compare-mode-panel active",
                            style={"flex": "1", "minWidth": "0"},
                        ),
                        html.Div(
                            dcc.Graph(id="compare-plot-right", style={"height": "100%"}),
                            className="compare-mode-panel",
                            style={"flex": "1", "minWidth": "0"},
                        ),
                    ],
                    className="compare-mode-container",
                    style={"display": "flex", "height": "calc(100% - 120px)"},
                ),
            ],
            id="compare-mode-panel",
            style={"display": "none", "height": "100%"},
        )

    def create_time_column_modal(self):
        """Modal for selecting time/X-axis column, time offsets, and header settings for CSVs."""
        return dbc.Modal(
            [
                dbc.ModalHeader("⏱️ CSV Settings"),
                dbc.ModalBody(
                    [
                        # Time Column Selection
                        html.H6("📊 Time Column", className="mb-2"),
                        html.P(
                            "Select which column to use as time/X-axis:",
                            className="small text-muted",
                        ),
                        html.Div(id="time-column-selectors"),
                        html.Hr(),
                        # Time Offsets
                        html.H6("⏱️ Time Offsets", className="mt-3"),
                        html.P(
                            "Add time offset (in seconds) to shift signal timing. "
                            "Positive = shift right, Negative = shift left.",
                            className="small text-muted",
                        ),
                        html.Div(id="time-offset-inputs"),
                        html.Hr(),
                        # CSV Header Settings
                        html.H6("📄 Header Settings", className="mt-3"),
                        html.P(
                            "For CSVs without headers or with headers not in the first row. "
                            "Set header row number (0-based) or leave empty for no headers.",
                            className="small text-muted",
                        ),
                        html.Div(id="csv-header-settings"),
                    ]
                ),
                dbc.ModalFooter(
                    [
                        dbc.Button(
                            "Apply & Reload",
                            id="btn-apply-time-cols",
                            color="primary",
                            size="sm",
                        ),
                        dbc.Button(
                            "Close",
                            id="btn-close-time-cols",
                            color="secondary",
                            size="sm",
                        ),
                    ]
                ),
            ],
            id="modal-time-cols",
            is_open=False,
            size="lg",
        )

    def create_annotation_modal(self):
        """Modal for adding text annotations to plots."""
        return dbc.Modal(
            [
                dbc.ModalHeader("📝 Add Annotation"),
                dbc.ModalBody(
                    [
                        html.P(
                            "Add a text note at a specific point on the plot. "
                            "Click on the plot first to set the position.",
                            className="small text-muted",
                        ),
                        dbc.Row(
                            [
                                dbc.Col(
                                    [
                                        dbc.Label("X Position:", className="small"),
                                        dbc.Input(
                                            id="annotation-x",
                                            type="number",
                                            placeholder="X value",
                                            size="sm",
                                        ),
                                    ],
                                    width=6,
                                ),
                                dbc.Col(
                                    [
                                        dbc.Label("Y Position:", className="small"),
                                        dbc.Input(
                                            id="annotation-y",
                                            type="number",
                                            placeholder="Y value",
                                            size="sm",
                                        ),
                                    ],
                                    width=6,
                                ),
                            ],
                            className="mb-2",
                        ),
                        dbc.Label("Text:", className="small"),
                        dbc.Textarea(
                            id="annotation-text",
                            placeholder="Enter annotation text...",
                            size="sm",
                            style={"height": "60px"},
                        ),
                        dbc.Row(
                            [
                                dbc.Col(
                                    [
                                        dbc.Label("Color:", className="small"),
                                        dbc.Input(
                                            id="annotation-color",
                                            type="color",
                                            value="#ffcc00",
                                            style={"width": "50px", "height": "30px"},
                                        ),
                                    ],
                                    width=4,
                                ),
                                dbc.Col(
                                    [
                                        dbc.Label("Font Size:", className="small"),
                                        dbc.Input(
                                            id="annotation-fontsize",
                                            type="number",
                                            value=12,
                                            min=8,
                                            max=24,
                                            size="sm",
                                        ),
                                    ],
                                    width=4,
                                ),
                                dbc.Col(
                                    [
                                        dbc.Label("Arrow:", className="small"),
                                        dbc.Checklist(
                                            id="annotation-arrow",
                                            options=[{"label": "Show", "value": True}],
                                            value=[True],
                                            switch=True,
                                        ),
                                    ],
                                    width=4,
                                ),
                            ],
                            className="mt-2",
                        ),
                    ]
                ),
                dbc.ModalFooter(
                    [
                        dbc.Button(
                            "Add",
                            id="btn-add-annotation",
                            color="primary",
                            size="sm",
                        ),
                        dbc.Button(
                            "Clear All",
                            id="btn-clear-annotations",
                            color="warning",
                            size="sm",
                        ),
                        dbc.Button(
                            "Close",
                            id="btn-close-annotation",
                            color="secondary",
                            size="sm",
                        ),
                    ]
                ),
            ],
            id="modal-annotation",
            is_open=False,
        )

    def create_compare_modal(self):
        """Modal for comparing CSVs and signals (supports 2+ CSVs)."""
        return dbc.Modal(
            [
                dbc.ModalHeader("📊 Compare CSVs"),
                dbc.ModalBody(
                    [
                        # CSV Selection - supports multiple
                        html.P(
                            "Select 2 or more CSVs to compare (first is reference):",
                            className="small text-muted",
                        ),
                        dbc.Row(
                            [
                                dbc.Col(
                                    [
                                        dbc.Label("Reference CSV:", className="small"),
                                        dbc.Select(
                                            id="compare-csv1",
                                            options=[],
                                            size="sm",
                                        ),
                                    ],
                                    width=6,
                                ),
                                dbc.Col(
                                    [
                                        dbc.Label("Compare to:", className="small"),
                                        dbc.Checklist(
                                            id="compare-csv2",
                                            options=[],
                                            value=[],
                                            inline=True,
                                            style={"fontSize": "11px"},
                                        ),
                                    ],
                                    width=6,
                                ),
                            ],
                            className="mb-3",
                        ),
                        html.Hr(),
                        # Signal comparison
                        dbc.Label("Or compare specific signals:", className="small"),
                        dbc.Row(
                            [
                                dbc.Col(
                                    dbc.Select(
                                        id="compare-signal1",
                                        options=[],
                                        placeholder="Signal 1...",
                                        size="sm",
                                    ),
                                    width=6,
                                ),
                                dbc.Col(
                                    dbc.Select(
                                        id="compare-signal2",
                                        options=[],
                                        placeholder="Signal 2...",
                                        size="sm",
                                    ),
                                    width=6,
                                ),
                            ],
                            className="mb-3",
                        ),
                        html.Hr(),
                        html.Div(
                            id="compare-results",
                            style={"maxHeight": "400px", "overflowY": "auto"},
                        ),
                    ]
                ),
                dbc.ModalFooter(
                    [
                        dbc.Button(
                            "Compare CSVs",
                            id="btn-do-compare-csvs",
                            color="primary",
                            size="sm",
                        ),
                        dbc.Button(
                            "Compare Signals",
                            id="btn-do-compare-signals",
                            color="info",
                            size="sm",
                        ),
                        dbc.Button(
                            "Plot Difference",
                            id="btn-plot-diff",
                            color="warning",
                            size="sm",
                        ),
                        dbc.Button(
                            "Close",
                            id="btn-close-compare",
                            color="secondary",
                            size="sm",
                        ),
                    ]
                ),
            ],
            id="modal-compare",
            is_open=False,
            size="lg",
        )

    def create_figure(
        self,
        rows: int,
        cols: int,
        theme: str,
        selected_subplot: int = 0,
        assignments: Optional[Dict] = None,
        tab_key: str = "0",
        link_axes: bool = False,
        time_cursor: bool = True,
        cursor_x: Optional[float] = None,
        subplot_modes: Optional[Dict[str, str]] = None,
        x_axis_signals: Optional[Dict[str, str]] = None,
        time_columns: Optional[Dict[str, str]] = None,
        subplot_metadata: Optional[Dict[str, Dict]] = None,
        time_offsets: Optional[Dict[str, float]] = None,
        display_options: Optional[Dict] = None,
        annotations: Optional[Dict] = None,
    ):
        """
        Create a Plotly figure with subplots for signal visualization.

        Args:
            rows: Number of subplot rows (1-4)
            cols: Number of subplot columns (1-4)
            theme: Color theme ("dark" or "light")
            selected_subplot: Index of highlighted subplot (-1 for no highlight)
            assignments: Dict of signal assignments {tab_key: {subplot_key: [signals]}}
            tab_key: Current tab identifier
            link_axes: Whether to link X axes across subplots
            time_cursor: Whether to show the time cursor
            cursor_x: X position of the time cursor
            subplot_modes: Mode per subplot ("time" or "xy")
            x_axis_signals: X-axis signal for XY mode per subplot
            time_columns: Custom time column per CSV file
            subplot_metadata: Titles/captions/descriptions per subplot

        Returns:
            go.Figure: Configured Plotly figure with subplots
        """
        import time as _time
        _start = _time.perf_counter()
        print(f"[PERF] 🔧 create_figure() started - tab={tab_key}, {rows}x{cols} layout")
        
        colors = THEMES[theme]

        # Calculate optimal spacing based on subplot count
        # More vertical spacing for titles, more horizontal for y-labels
        v_spacing = max(0.12, 0.18 / rows) if rows > 1 else 0.05  # More space for titles
        h_spacing = max(0.10, 0.15 / cols) if cols > 1 else 0.02  # More space for y-labels

        # Build subplot titles from metadata or use defaults
        subplot_metadata = subplot_metadata or {}
        subplot_titles = []
        for i in range(rows * cols):
            sp_meta = subplot_metadata.get(str(i), {})
            title = sp_meta.get("title", "")
            subplot_titles.append(title if title else f"Subplot {i+1}")

        fig = make_subplots(
            rows=rows,
            cols=cols,
            subplot_titles=subplot_titles,
            vertical_spacing=v_spacing,
            horizontal_spacing=h_spacing,
            shared_xaxes=link_axes,
            shared_yaxes=link_axes,
        )

        # Calculate legend positions for each subplot
        legend_configs = {}
        for sp_idx in range(rows * cols):
            r = sp_idx // cols  # 0-indexed row
            c = sp_idx % cols  # 0-indexed col

            # Calculate x,y position for this subplot's legend
            # Each subplot gets a legend in its top-right corner
            x_start = c / cols
            x_end = (c + 1) / cols
            y_start = 1 - (r + 1) / rows
            y_end = 1 - r / rows

            legend_name = f"legend{sp_idx + 1}" if sp_idx > 0 else "legend"
            legend_configs[legend_name] = dict(
                bgcolor=colors["card"],
                bordercolor=colors["border"],
                borderwidth=1,  # ← ADD THIS
                font=dict(size=8),
                x=x_end - 0.02,
                y=y_end - 0.02,
                xanchor="right",
                yanchor="top",
                xref="paper",  # ← ADD THIS (CRITICAL!)
                yref="paper",  # ← ADD THIS (CRITICAL!)
                valign="top",  # ← ADD THIS
                orientation="v",  # ← ADD THIS
                tracegroupgap=1,
                itemclick="toggle",
                itemdoubleclick="toggleothers",
            )

        # PERFORMANCE: uirevision controls zoom/pan state persistence
        # CRITICAL: Use a constant value to preserve zoom/pan state across updates
        # Changing uirevision causes zoom/pan to reset, leading to white pages
        ui_revision = "preserve_zoom_pan"
        
        fig.update_layout(
            paper_bgcolor=colors["paper_bg"],
            plot_bgcolor=colors["plot_bg"],
            font=dict(color=colors["text"], size=10),
            height=max(480, 280 * rows),  # Increased height per row
            showlegend=True,
            margin=dict(l=40, r=20, t=40, b=35),  # Reduced margins for more plot space
            # PERFORMANCE: Use "x unified" for faster hover with many traces
            hovermode="x unified",
            hoverlabel=dict(
                bgcolor="rgba(0,0,0,0.8)",
                font_size=10,
                font_color="#e8e8e8",
                namelength=-1,  # Show full name
            ),
            # PERFORMANCE: Preserve zoom/pan state on updates EXCEPT when link_axes changes
            uirevision=ui_revision,
            # PERFORMANCE: Disable drag mode animations
            dragmode="zoom",
            **legend_configs,
        )

        # Style each subplot - always show grid even if empty
        # ALWAYS highlight selected subplot with visible border
        for i in range(rows * cols):
            r = i // cols + 1
            c = i % cols + 1

            is_selected = (i == selected_subplot)
            border_color = "#4ea8de" if is_selected else colors["grid"]
            border_width = 3 if is_selected else 1  # Thicker border for selected
            
            # Add a more visible background for selected subplot
            subplot_bg = "rgba(78, 168, 222, 0.08)" if is_selected else colors["plot_bg"]

            # Add a small invisible trace to ensure axes are shown
            # Use Scattergl for WebGL acceleration
            fig.add_trace(
                go.Scattergl(
                    x=[0],
                    y=[0],
                    mode="markers",
                    marker=dict(size=1, opacity=0),
                    showlegend=False,
                    hoverinfo="skip",
                ),
                row=r,
                col=c,
            )

            fig.update_xaxes(
                gridcolor=colors["grid"],
                zerolinecolor=colors["border"],
                title_text="Time",
                title_font=dict(size=9),
                title_standoff=10,  # Fixed distance from axis to prevent jumping
                tickfont=dict(size=8),
                showline=True,
                linewidth=border_width,
                linecolor=border_color,
                mirror=True,
                showgrid=True,
                autorange=True,
                # CRITICAL: Fix xlabel position to prevent jumping
                automargin=False,  # Disable automatic margin adjustment
                side="bottom",  # Always at bottom
                showspikes=False,  # Cursor is click-based, not hover-based
                row=r,
                col=c,
            )
            fig.update_yaxes(
                gridcolor=colors["grid"],
                zerolinecolor=colors["border"],
                title_text="Value",
                title_font=dict(size=9),
                tickfont=dict(size=8),
                showline=True,
                linewidth=border_width,
                linecolor=border_color,
                mirror=True,
                showgrid=True,
                autorange=True,
                automargin=True,  # Prevent y-axis labels from overlapping with adjacent subplots
                row=r,
                col=c,
            )

        # Plot signals - each subplot gets its own legend group
        # Also collect signal data for cursor value display
        subplot_signal_data = {}  # {sp_idx: [(x_arr, y_arr, name, color), ...]}
        subplot_modes = subplot_modes or {}
        display_options = display_options or {}
        annotations = annotations or {}

        if assignments and tab_key in assignments:
            # First pass: collect all signal names to check for duplicates
            # PERFORMANCE: Use list comprehension and Counter for faster duplicate detection
            all_signal_names = [
                sig.get("signal", "")
                for sp_idx in range(rows * cols)
                for sig in (assignments.get(tab_key, {}).get(str(sp_idx), []) or [])
                if isinstance(assignments.get(tab_key, {}).get(str(sp_idx), []), list)
            ]
            # PERFORMANCE: Use Counter for O(n) duplicate detection instead of O(n²) count()
            signal_counts = Counter(all_signal_names)
            duplicate_signals = {name for name, count in signal_counts.items() if count > 1}

            # Track colors by signal key for consistency across subplots
            signal_color_map = {}  # {signal_key: color}
            trace_idx = 0  # Unique trace counter for independent legend behavior
            for sp_idx in range(rows * cols):
                subplot_signal_data[sp_idx] = []
                # FIX: Sanitize assignments access to prevent unhashable dict errors, especially for tab "0"
                tab_assignments = assignments.get(tab_key, {})
                if not isinstance(tab_assignments, dict):
                    tab_assignments = {}
                sp_assignment_raw = tab_assignments.get(str(sp_idx), [])
                # Ensure sp_assignment is always a list with proper dict items
                if isinstance(sp_assignment_raw, list):
                    sp_assignment = [
                        s for s in sp_assignment_raw
                        if isinstance(s, dict) and 'csv_idx' in s and 'signal' in s
                    ]
                elif isinstance(sp_assignment_raw, dict):
                    # Old dict format - convert to empty list
                    sp_assignment = []
                else:
                    # Invalid format - use empty list
                    sp_assignment = []
                r = sp_idx // cols + 1
                c = sp_idx % cols + 1

                # Get subplot mode and X-axis signal
                # IMPORTANT: subplot_modes is ALREADY the per-tab dict passed from handle_assignments
                # Structure: {subplot_key: mode} NOT {tab_key: {subplot_key: mode}}
                subplot_mode = "time"  # Default
                if subplot_modes:
                    if isinstance(subplot_modes, dict):
                        # It's a dict - get mode for this subplot
                        subplot_mode = subplot_modes.get(str(sp_idx), "time")
                    elif isinstance(subplot_modes, str):
                        # Corrupted data - it's a mode string directly
                        subplot_mode = subplot_modes
                        print(f"[WARN] subplot_modes is a string '{subplot_modes}', expected dict")
                
                # IMPORTANT: x_axis_signals is ALREADY the per-tab dict passed from handle_assignments
                # Structure: {subplot_key: signal_key} NOT {tab_key: {subplot_key: signal_key}}
                x_axis_choice = "time"  # Default
                if x_axis_signals:
                    if isinstance(x_axis_signals, dict):
                        x_axis_choice = x_axis_signals.get(str(sp_idx), "time")
                    elif isinstance(x_axis_signals, str):
                        # Corrupted data - it's a signal key directly
                        x_axis_choice = x_axis_signals
                        print(f"[WARN] x_axis_signals is a string '{x_axis_signals}', expected dict")
                
                # Debug: Log X-Y mode status
                if subplot_mode == "xy":
                    print(f"[X-Y DEBUG] Subplot {sp_idx}: mode={subplot_mode}, x_axis_choice={x_axis_choice}")
                time_columns = time_columns or {}

                # Handle both list and dict assignment formats (convert dict to empty list for time/xy mode)
                sp_signals = sp_assignment if isinstance(sp_assignment, list) else []

                # Determine X-axis data source
                x_axis_data = None
                x_axis_label = "Time"

                if subplot_mode == "xy" and x_axis_choice != "time":
                    # X-axis is a signal
                    print(f"[X-Y DEBUG] Using signal for X-axis: {x_axis_choice}")
                    try:
                        x_csv_idx, x_signal_name = x_axis_choice.split(":", 1)
                        x_csv_idx = int(x_csv_idx)

                        if x_csv_idx == -1 and x_signal_name in self.derived_signals:
                            ds = self.derived_signals[x_signal_name]
                            x_axis_data = np.array(ds.get("data", []))
                            x_axis_label = f"{x_signal_name} (D)"
                        elif 0 <= x_csv_idx < len(self.data_manager.data_tables):
                            df = self.data_manager.data_tables[x_csv_idx]
                            if df is not None and x_signal_name in df.columns:
                                x_axis_data = df[x_signal_name].values
                                csv_name = (
                                    get_csv_short_name(
                                        self.data_manager.csv_file_paths[x_csv_idx]
                                    )
                                    if x_csv_idx < len(self.data_manager.csv_file_paths)
                                    else f"C{x_csv_idx+1}"
                                )
                                x_axis_label = f"{x_signal_name} ({csv_name})"
                    except:
                        pass

                # Update X-axis label for X-Y mode
                if subplot_mode == "xy" and x_axis_data is not None:
                    print(f"[X-Y DEBUG] Setting X-axis label: {x_axis_label}, data len: {len(x_axis_data)}")
                    fig.update_xaxes(title_text=x_axis_label, row=r, col=c)
                elif subplot_mode == "xy":
                    print(f"[X-Y DEBUG] X-Y mode but x_axis_data is None (x_axis_choice={x_axis_choice})")

                for sig in sp_signals:
                    csv_idx = sig.get("csv_idx", -1)
                    signal_name = sig.get("signal", "")
                    is_state_signal = sig.get("is_state", False)

                    # In X-Y mode, skip plotting the X-axis signal as a Y trace
                    if subplot_mode == "xy" and x_axis_choice != "time":
                        sig_key = f"{csv_idx}:{signal_name}"
                        if sig_key == x_axis_choice:
                            continue  # Don't plot X-axis signal as Y

                    if csv_idx == -1 and signal_name in self.derived_signals:
                        ds = self.derived_signals[signal_name]
                        x_data = np.array(ds.get("time", []))
                        y_data = np.array(ds.get("data", []))
                        csv_label = "D"  # Derived
                    elif csv_idx >= 0 and csv_idx < len(self.data_manager.data_tables):
                        # Determine time column name
                        custom_time_col = time_columns.get(str(csv_idx))
                        if custom_time_col:
                            time_col = custom_time_col
                        else:
                            time_col = "Time"

                        # Get RAW data - no downsampling, all points preserved
                        x_data, y_data = self.get_signal_data_cached(
                            csv_idx, signal_name, time_col
                        )

                        # Check if data was found
                        if len(x_data) == 0 or len(y_data) == 0:
                            continue

                        # Apply time offset if set (avoid array copy when offset is 0)
                        time_offsets = time_offsets or {}
                        sig_offset_key = f"{csv_idx}:{signal_name}"
                        csv_offset_key = str(csv_idx)
                        time_offset = time_offsets.get(sig_offset_key, time_offsets.get(csv_offset_key, 0))
                        if time_offset != 0:
                            x_data = np.asarray(x_data) + time_offset  # asarray avoids copy if already ndarray

                        # Override X-axis if in X-Y mode
                        if subplot_mode == "xy" and x_axis_data is not None:
                            x_data = x_axis_data

                        # Get CSV filename for legend (include folder if duplicate)
                        if csv_idx < len(self.data_manager.csv_file_paths):
                            csv_path = self.data_manager.csv_file_paths[csv_idx]
                            csv_label = get_csv_display_name(csv_path, self.data_manager.csv_file_paths)
                            # Remove .csv extension for cleaner legend
                            csv_label = os.path.splitext(csv_label)[0]
                        else:
                            csv_label = f"C{csv_idx+1}"
                    else:
                        continue

                    prop_key = f"{csv_idx}:{signal_name}"
                    props = self.signal_properties.get(prop_key, {})
                    
                    # Get color: user-defined > previously assigned > deterministic based on signal key
                    if "color" in props:
                        color = props["color"]
                    elif prop_key in signal_color_map:
                        color = signal_color_map[prop_key]
                    else:
                        # Use deterministic color based on signal key hash
                        color_idx = hash(prop_key) % len(SIGNAL_COLORS)
                        color = SIGNAL_COLORS[color_idx]
                        signal_color_map[prop_key] = color
                    
                    width = props.get("width", 1.5)
                    scale = props.get("scale", 1.0)
                    display_name = props.get("display_name", signal_name)
                    is_state_signal = props.get("is_state", is_state_signal)

                    # Build legend name: format as "signal (csv_name)"
                    legend_name = f"{display_name} ({csv_label})"

                    # PERFORMANCE: Avoid array copy when scale is 1.0 (default)
                    # Data from cache is already numpy array
                    x_arr = x_data if isinstance(x_data, np.ndarray) else np.asarray(x_data)
                    if scale == 1.0:
                        y_scaled = y_data if isinstance(y_data, np.ndarray) else np.asarray(y_data)
                    else:
                        y_scaled = np.asarray(y_data) * scale

                    # Collect signal data for cursor value display
                    subplot_signal_data[sp_idx].append(
                        (x_arr, y_scaled, legend_name, color)
                    )

                    # Assign trace to this subplot's legend
                    legend_ref = f"legend{sp_idx + 1}" if sp_idx > 0 else "legend"
                    # Unique legendgroup for each trace so clicking hides only that trace
                    unique_legend_group = f"trace_{trace_idx}"

                    if is_state_signal:
                        # State signal: draw vertical lines at value changes
                        # Find indices where value changes
                        changes = np.where(np.diff(y_scaled) != 0)[0] + 1
                        
                        # PERFORMANCE: Limit number of vlines for large state signals
                        MAX_STATE_LINES = 200
                        if len(changes) > MAX_STATE_LINES:
                            # Sample evenly across the changes
                            step = len(changes) // MAX_STATE_LINES
                            changes = changes[::step]

                        for change_idx in changes:
                            if change_idx < len(x_arr):
                                # Draw vertical line using shape
                                fig.add_vline(
                                    x=x_arr[change_idx],
                                    line=dict(color=color, width=width, dash="solid"),
                                    row=r,
                                    col=c,
                                )

                        # Add invisible trace for legend
                        # Use Scattergl for consistency
                        fig.add_trace(
                            go.Scattergl(
                                x=[None],
                                y=[None],
                                mode="lines",
                                name=f"{legend_name} (state)",
                                line=dict(color=color, width=width),
                                legendgroup=unique_legend_group,
                                showlegend=True,
                                legend=legend_ref,
                            ),
                            row=r,
                            col=c,
                        )
                    else:
                        # Regular signal: ALWAYS use Scattergl for WebGL hardware acceleration
                        # WebGL is critical for plotting large datasets with all raw points
                        trace_type = go.Scattergl
                        
                        # Get display options for this subplot
                        sp_display_opts = display_options.get(tab_key, {}).get(str(sp_idx), {})
                        show_markers = sp_display_opts.get("markers", False)
                        normalize_signals = sp_display_opts.get("normalize", False)
                        
                        # Normalize if enabled (scale to 0-1 range)
                        y_display = y_scaled
                        if normalize_signals and len(y_scaled) > 0:
                            y_min = np.nanmin(y_scaled)
                            y_max = np.nanmax(y_scaled)
                            if y_max != y_min:
                                y_display = (y_scaled - y_min) / (y_max - y_min)
                                legend_name = f"{legend_name} [N]"  # Mark as normalized
                        
                        # PERFORMANCE: Adjust rendering based on data size
                        # All data is shown - no downsampling - but rendering is optimized
                        n_points = len(x_arr)
                        
                        # Ultra-large data optimizations (500k+ points per signal)
                        if n_points > self.ULTRA_LARGE_THRESHOLD:
                            show_markers = False
                            line_width = 0.5  # Minimal line width for speed
                            hover_mode = "skip"  # Critical: disable hover completely
                        # Large data optimizations (50k+ points)
                        elif n_points > self.LARGE_DATA_THRESHOLD:
                            show_markers = False
                            line_width = 1  # Thinner lines render faster
                            hover_mode = "skip"
                        # Medium data (5k-50k points)
                        elif n_points > self.HOVER_THRESHOLD:
                            show_markers = False
                            line_width = width
                            hover_mode = "skip"
                        # Small data - full features
                        else:
                            line_width = width
                            hover_mode = "x+y+name"
                        
                        trace_mode = "lines+markers" if show_markers else "lines"
                        marker_dict = dict(size=3, color=color) if show_markers else None

                        # Trace options optimized for performance
                        trace_opts = dict(
                            x=x_arr,
                            y=y_display,
                            mode=trace_mode,
                            name=legend_name,
                            line=dict(color=color, width=line_width),
                            legendgroup=unique_legend_group,
                            showlegend=True,
                            legend=legend_ref,
                            hoverinfo=hover_mode,
                        )
                        
                        # Only add marker dict if needed
                        if marker_dict:
                            trace_opts["marker"] = marker_dict
                        
                        fig.add_trace(trace_type(**trace_opts), row=r, col=c)
                    trace_idx += 1

        # Add time cursor vertical lines to all subplots with signal values
        if time_cursor and cursor_x is not None:
            for sp_idx in range(rows * cols):
                r = sp_idx // cols + 1
                c = sp_idx % cols + 1

                # Get axis references for this subplot
                xref = f"x{sp_idx + 1}" if sp_idx > 0 else "x"
                yref = f"y{sp_idx + 1} domain" if sp_idx > 0 else "y domain"
                yref_data = f"y{sp_idx + 1}" if sp_idx > 0 else "y"

                fig.add_shape(
                    type="line",
                    x0=cursor_x,
                    x1=cursor_x,
                    y0=0,
                    y1=1,
                    xref=xref,
                    yref=yref,
                    line=dict(color="#ff6b6b", width=2, dash="solid"),
                )

                # Build cursor info text with signal values
                cursor_text_lines = [f"t={cursor_x:.4f}"]

                # Get signal values at cursor position for this subplot
                if sp_idx in subplot_signal_data:
                    for x_arr, y_arr, sig_name, sig_color in subplot_signal_data[
                        sp_idx
                    ]:
                        if len(x_arr) > 0 and len(y_arr) > 0:
                            # Interpolate value at cursor_x
                            try:
                                val = np.interp(cursor_x, x_arr, y_arr)
                                cursor_text_lines.append(f"{sig_name}: {val:.4f}")
                            except:
                                pass

                cursor_text = "<br>".join(cursor_text_lines)

                # Add annotation with time and signal values
                fig.add_annotation(
                    x=cursor_x,
                    y=1,
                    xref=xref,
                    yref=yref,
                    text=cursor_text,
                    showarrow=False,
                    bgcolor="rgba(40, 40, 40, 0.9)",
                    font=dict(color="white", size=9),
                    align="left",
                    yshift=10,
                    bordercolor="#ff6b6b",
                    borderwidth=1,
                    borderpad=4,
                )

        # Add user annotations
        tab_annotations = annotations.get(tab_key, {})
        for sp_idx in range(rows * cols):
            sp_annotations = tab_annotations.get(str(sp_idx), [])
            if not sp_annotations:
                continue
                
            xref = f"x{sp_idx + 1}" if sp_idx > 0 else "x"
            yref = f"y{sp_idx + 1}" if sp_idx > 0 else "y"
            
            for ann in sp_annotations:
                ann_x = ann.get("x")
                ann_y = ann.get("y", 0)
                ann_text = ann.get("text", "")
                ann_color = ann.get("color", "#ffcc00")
                ann_fontsize = ann.get("fontsize", 12)
                ann_arrow = ann.get("arrow", True)
                
                if ann_x is not None and ann_text:
                    fig.add_annotation(
                        x=ann_x,
                        y=ann_y if ann_y is not None else 0.5,
                        xref=xref,
                        yref=yref if ann_y is not None else f"y{sp_idx + 1} domain" if sp_idx > 0 else "y domain",
                        text=ann_text,
                        showarrow=ann_arrow,
                        arrowhead=2,
                        arrowsize=1,
                        arrowwidth=1.5,
                        arrowcolor=ann_color,
                        font=dict(color=ann_color, size=ann_fontsize),
                        bgcolor="rgba(0, 0, 0, 0.7)",
                        bordercolor=ann_color,
                        borderwidth=1,
                        borderpad=3,
                    )

        # Performance timing
        _elapsed = (_time.perf_counter() - _start) * 1000
        if _elapsed < 50:
            status = "⚡"
        elif _elapsed < 200:
            status = "⏱️"
        elif _elapsed < 500:
            status = "🐢"
        else:
            status = "🔴"
        print(f"[PERF] {status} create_figure() completed: {_elapsed:.1f}ms")
        
        return fig

    def setup_callbacks(self):
        logger.info("Setting up callbacks...")
        
        # ================================================================
        # PERFORMANCE: Clientside callbacks for instant UI responses
        # These run in the browser without server round-trip
        # ================================================================
        
        # DISABLED: This clientside callback was causing white flash
        # It conflicts with the server-side callback that rebuilds figures
        # The server-side callback now handles all subplot highlighting
        #
        # # Instant subplot highlight update (no server call needed)
        # self.app.clientside_callback(
        #     """
        #     function(subplot_val, rows, cols, is_dark) {
        #         // This runs instantly in browser for responsive subplot selection
        #         if (subplot_val === undefined || subplot_val === null) {
        #             return window.dash_clientside.no_update;
        #         }
        #
        #         // Calculate total subplots
        #         var total = (rows || 1) * (cols || 1);
        #         var selected = parseInt(subplot_val);
        #
        #         // Update axis borders using Patch-like logic
        #         // We return a Patch object that updates only the borders
        #         var patch = {};
        #         var theme = is_dark ? 'dark' : 'light';
        #         var borderColor = '#4ea8de';
        #         var defaultColor = theme === 'dark' ? '#444' : '#dee2e6';
        #
        #         for (var i = 0; i < Math.min(total, 16); i++) {
        #             var xKey = i === 0 ? 'xaxis' : 'xaxis' + (i + 1);
        #             var yKey = i === 0 ? 'yaxis' : 'yaxis' + (i + 1);
        #
        #             var isSelected = (i === selected);
        #             patch[xKey] = {
        #                 linecolor: isSelected ? borderColor : defaultColor,
        #                 linewidth: isSelected ? 3 : 1
        #             };
        #             patch[yKey] = {
        #                 linecolor: isSelected ? borderColor : defaultColor,
        #                 linewidth: isSelected ? 3 : 1
        #             };
        #         }
        #
        #         return {layout: patch};
        #     }
        #     """,
        #     Output("plot", "figure", allow_duplicate=True),
        #     [
        #         Input("subplot-select", "value"),
        #         Input("rows-input", "value"),
        #         Input("cols-input", "value"),
        #         Input("theme-switch", "value"),
        #     ],
        #     prevent_initial_call=True,
        # )
        
        # Keyboard shortcuts handler - clientside for instant response
        self.app.clientside_callback(
            """
            function(id) {
                // Set up keyboard shortcuts once
                if (!window._signalViewerKeyboardSetup) {
                    window._signalViewerKeyboardSetup = true;
                    
                    document.addEventListener('keydown', function(e) {
                        // Ignore if typing in input fields
                        if (e.target.tagName === 'INPUT' || e.target.tagName === 'TEXTAREA') {
                            return;
                        }
                        
                        // Ctrl+S - Save session
                        if (e.ctrlKey && e.key === 's') {
                            e.preventDefault();
                            var saveBtn = document.getElementById('btn-save');
                            if (saveBtn) saveBtn.click();
                        }
                        
                        // Space - Toggle cursor play/pause
                        if (e.key === ' ' && !e.ctrlKey) {
                            e.preventDefault();
                            var playBtn = document.getElementById('btn-cursor-play');
                            var stopBtn = document.getElementById('btn-cursor-stop');
                            // Toggle based on which is active
                            if (playBtn && playBtn.style.opacity !== '0.5') {
                                playBtn.click();
                            } else if (stopBtn) {
                                stopBtn.click();
                            }
                        }
                        
                        // Number keys 1-9 - Select subplot
                        if (e.key >= '1' && e.key <= '9' && !e.ctrlKey && !e.altKey) {
                            var subplotSelect = document.getElementById('subplot-select');
                            if (subplotSelect) {
                                var idx = parseInt(e.key) - 1;
                                var options = subplotSelect.querySelectorAll('option');
                                if (idx < options.length) {
                                    subplotSelect.value = idx;
                                    subplotSelect.dispatchEvent(new Event('change', { bubbles: true }));
                                }
                            }
                        }
                        
                        // Delete or Backspace - Remove selected signals
                        if (e.key === 'Delete' || e.key === 'Backspace') {
                            // Don't handle if in input
                            if (e.target.tagName !== 'INPUT' && e.target.tagName !== 'TEXTAREA') {
                                // Find checked remove checkboxes and click remove
                                var checks = document.querySelectorAll('[id*="remove-check"]');
                                var hasChecked = false;
                                checks.forEach(function(c) {
                                    if (c.checked) hasChecked = true;
                                });
                                if (hasChecked) {
                                    var removeBtn = document.getElementById('btn-remove-selected');
                                    if (removeBtn) removeBtn.click();
                                }
                            }
                        }
                    });
                    
                    console.log('[SignalViewer] Keyboard shortcuts enabled');
                }
                return window.dash_clientside.no_update;
            }
            """,
            Output("store-refresh-trigger", "data", allow_duplicate=True),
            Input("app-container", "id"),
            prevent_initial_call=True,
        )
        
        # Track visible range for viewport-based optimizations
        self.app.clientside_callback(
            """
            function(relayoutData, currentRange) {
                if (!relayoutData) {
                    return window.dash_clientside.no_update;
                }
                
                // Extract visible X range from relayout data
                var newRange = currentRange || {};
                
                // Check for autosize or range changes
                if (relayoutData['xaxis.range[0]'] !== undefined) {
                    newRange.xmin = relayoutData['xaxis.range[0]'];
                    newRange.xmax = relayoutData['xaxis.range[1]'];
                } else if (relayoutData['xaxis.autorange']) {
                    newRange.xmin = null;
                    newRange.xmax = null;
                }
                
                return newRange;
            }
            """,
            Output("store-visible-range", "data"),
            Input("plot", "relayoutData"),
            State("store-visible-range", "data"),
            prevent_initial_call=True,
        )

        # ================================================================
        # SIGNAL SEARCH AUTOCOMPLETE
        # ================================================================
        @self.app.callback(
            [
                Output("search-autocomplete-dropdown", "children"),
                Output("search-autocomplete-dropdown", "style"),
            ],
            Input("search-input", "value"),
            State("store-csv-files", "data"),
            prevent_initial_call=True,
        )
        def update_autocomplete(search_value, csv_files):
            """Show autocomplete suggestions as user types"""
            if not search_value or len(search_value) < 2:
                return [], {"display": "none"}
            
            suggestions = []
            search_lower = search_value.lower()
            max_suggestions = 10
            
            # Search through all signals
            for csv_idx, df in enumerate(self.data_manager.data_tables):
                if df is None:
                    continue
                csv_name = self.data_manager.csv_file_paths[csv_idx] if csv_idx < len(self.data_manager.csv_file_paths) else f"CSV {csv_idx}"
                csv_short = csv_name.split("/")[-1].split("\\")[-1][:20]
                
                for col in df.columns:
                    if col.lower() == "time":
                        continue
                    if search_lower in col.lower():
                        suggestions.append({
                            "csv_idx": csv_idx,
                            "signal": col,
                            "csv_name": csv_short,
                        })
                        if len(suggestions) >= max_suggestions:
                            break
                if len(suggestions) >= max_suggestions:
                    break
            
            # Also search derived signals
            for sig_name in self.derived_signals.keys():
                if search_lower in sig_name.lower():
                    suggestions.append({
                        "csv_idx": -1,
                        "signal": sig_name,
                        "csv_name": "Derived",
                    })
                    if len(suggestions) >= max_suggestions:
                        break
            
            if not suggestions:
                return [html.Div("No matches found", className="text-muted p-2", style={"fontSize": "11px"})], {
                    "display": "block",
                    "position": "absolute",
                    "zIndex": 1000,
                    "backgroundColor": "#1a1a2e",
                    "border": "1px solid #4ea8de",
                    "borderRadius": "4px",
                    "width": "100%",
                    "boxShadow": "0 4px 12px rgba(0,0,0,0.3)",
                }
            
            # Create suggestion items
            items = []
            for i, s in enumerate(suggestions):
                items.append(
                    html.Div(
                        [
                            html.Span(s["signal"], style={"fontWeight": "bold"}),
                            html.Span(f" ({s['csv_name']})", className="text-muted", style={"fontSize": "10px"}),
                        ],
                        id={"type": "autocomplete-item", "idx": i, "csv_idx": s["csv_idx"], "signal": s["signal"]},
                        style={
                            "padding": "6px 10px",
                            "cursor": "pointer",
                            "fontSize": "11px",
                            "borderBottom": "1px solid rgba(255,255,255,0.1)",
                        },
                        className="autocomplete-item",
                    )
                )
            
            return items, {
                "display": "block",
                "position": "absolute",
                "zIndex": 1000,
                "backgroundColor": "#1a1a2e",
                "border": "1px solid #4ea8de",
                "borderRadius": "4px",
                "maxHeight": "200px",
                "overflowY": "auto",
                "width": "100%",
                "boxShadow": "0 4px 12px rgba(0,0,0,0.3)",
            }

        # Theme toggle
        @self.app.callback(
            [
                Output("app-container", "style"),
                Output("card-csv", "style"),
                Output("card-header-csv", "style"),
                Output("card-body-csv", "style"),
                Output("card-signals", "style"),
                Output("card-header-signals", "style"),
                Output("card-body-signals", "style"),
                Output("card-assigned", "style"),
                Output("card-header-assigned", "style"),
                Output("card-body-assigned", "style"),
                Output("card-plot", "style"),
                Output("card-header-plot", "style"),
                Output("card-body-plot", "style"),
                Output("target-box", "style"),
                Output("search-input", "style"),
                Output("store-theme", "data"),
            ],
            Input("theme-switch", "value"),
        )
        def update_theme(is_dark):
            theme = "dark" if is_dark else "light"
            c = THEMES[theme]
            # Modern card styling with rounded corners and subtle shadows
            card = {
                "backgroundColor": c["card"],
                "borderColor": c["border"],
                "borderRadius": "8px",  # Modern rounded corners
                "borderWidth": "1px",
                "boxShadow": "0 2px 4px rgba(0,0,0,0.1)" if not is_dark else "0 2px 8px rgba(0,0,0,0.3)",
            }
            header = {
                "backgroundColor": c["card_header"],
                "borderColor": c["border"],
                "color": c["text"],
                "borderRadius": "8px 8px 0 0",  # Rounded top corners
                "borderWidth": "0 0 1px 0",
            }
            body = {
                "backgroundColor": c["card"],
                "color": c["text"],
                "borderRadius": "0 0 8px 8px",  # Rounded bottom corners
            }

            # Modern input styling
            input_style = {
                "backgroundColor": c["input_bg"],
                "color": c["text"],
                "borderColor": c["border"],
                "borderRadius": "6px",  # Rounded inputs
                "borderWidth": "1px",
                "padding": "6px 10px",
                "transition": "all 0.2s ease",  # Smooth transitions
            }
            
            return (
                {
                    "backgroundColor": c["bg"],
                    "color": c["text"],
                    "minHeight": "100vh",
                    "padding": "12px",  # Slightly more padding
                    "fontFamily": "'Segoe UI', 'Inter', -apple-system, BlinkMacSystemFont, sans-serif",  # Modern font stack
                },
                card,
                header,
                body,
                card,
                header,
                body,
                card,
                header,
                body,
                card,
                header,
                body,
                input_style,
                input_style,
                theme,
            )

        # =================================================================
        # Native File Browser - Opens Windows file dialog
        # =================================================================
        
        @self.app.callback(
            [
                Output("store-csv-files", "data", allow_duplicate=True),
                Output("csv-list", "children", allow_duplicate=True),
                Output("status-text", "children", allow_duplicate=True),
                Output("store-refresh-trigger", "data", allow_duplicate=True),
            ],
            Input("btn-browse-files", "n_clicks"),
            [
                State("store-csv-files", "data"),
                State("store-refresh-trigger", "data"),
            ],
            prevent_initial_call=True,
        )
        def open_file_browser(n_clicks, current_files, refresh_trigger):
            """Open native file browser and load selected CSV files"""
            if not n_clicks:
                return no_update, no_update, no_update, no_update
            
            current_files = current_files or []
            
            try:
                # Use tkinter for native file dialog (runs on server = local machine)
                import tkinter as tk
                from tkinter import filedialog
                
                # Create hidden root window
                root = tk.Tk()
                root.withdraw()
                root.attributes('-topmost', True)  # Bring dialog to front
                
                # Open file dialog
                file_paths = filedialog.askopenfilenames(
                    title="Select CSV Files",
                    filetypes=[
                        ("CSV files", "*.csv"),
                        ("All files", "*.*")
                    ],
                    initialdir=os.path.expanduser("~")
                )
                
                root.destroy()
                
                if not file_paths:
                    return no_update, no_update, "Cancelled", no_update
                
                # Load selected files using ORIGINAL paths (no copy!)
                new_count = 0
                for file_path in file_paths:
                    if file_path in current_files:
                        print(f"[WARN] Already loaded: {file_path}")
                        continue
                    
                    if not os.path.exists(file_path):
                        print(f"[ERROR] Not found: {file_path}")
                        continue
                    
                    # Add to files list (using original path!)
                    current_files.append(file_path)
                    self.original_file_paths[file_path] = file_path
                    new_count += 1
                    
                    # Update data_manager
                    idx = len(current_files) - 1
                    self.data_manager.csv_file_paths = current_files
                    self.data_manager.original_source_paths[idx] = file_path
                    
                    while len(self.data_manager.data_tables) < len(current_files):
                        self.data_manager.data_tables.append(None)
                        self.data_manager.last_read_rows.append(0)
                        self.data_manager.last_file_mod_times.append(0)
                    
                    # Load the CSV
                    try:
                        self.data_manager.read_initial_data(idx)
                        rows = len(self.data_manager.data_tables[idx]) if self.data_manager.data_tables[idx] is not None else 0
                        parent = os.path.basename(os.path.dirname(file_path))
                        fname = os.path.basename(file_path)
                        print(f"[OK] Loaded: {parent}/{fname} ({rows:,} rows)")
                    except Exception as e:
                        print(f"[ERROR] Error loading {file_path}: {e}")
                
                # Build CSV list UI
                items = [
                    html.Div(
                        [
                            html.I(className="fas fa-file me-1", style={"color": "#f4a261"}),
                            html.Span(
                                get_csv_display_name(f, current_files), 
                                style={"fontSize": "10px"},
                                title=f  # Show full path on hover
                            ),
                            html.A(
                                "×",
                                id={"type": "del-csv", "idx": i},
                                className="float-end text-danger",
                                style={"cursor": "pointer"},
                            ),
                        ],
                        className="py-1",
                    )
                    for i, f in enumerate(current_files)
                ]
                
                total_rows = sum(len(df) for df in self.data_manager.data_tables if df is not None)
                status = f"[OK] Loaded {new_count} file(s), {total_rows:,} rows total"
                
                return current_files, items, status, (refresh_trigger or 0) + 1
                
            except Exception as e:
                print(f"[ERROR] File browser error: {e}")
                import traceback
                traceback.print_exc()
                # UX: Provide helpful error message with suggestion
                error_msg = str(e)
                if "tkinter" in error_msg.lower():
                    return no_update, no_update, "⚠️ File dialog requires display. Try drag-drop or path entry.", no_update
                return no_update, no_update, f"⚠️ Could not open files: {error_msg}", no_update

        # CSV Delete, Clear & Refresh
        @self.app.callback(
            [
                Output("store-csv-files", "data", allow_duplicate=True),
                Output("csv-list", "children"),
                Output("status-text", "children", allow_duplicate=True),
                Output("store-refresh-trigger", "data"),
                Output("store-assignments", "data", allow_duplicate=True),
            ],
            [
                Input("upload-csv", "contents"),
                Input("btn-clear-csv", "n_clicks"),
                Input({"type": "del-csv", "idx": ALL}, "n_clicks"),
                Input("btn-refresh-csv", "n_clicks"),
            ],
            [
                State("upload-csv", "filename"),
                State("store-csv-files", "data"),
                State("store-refresh-trigger", "data"),
                State("store-assignments", "data"),
            ],
            prevent_initial_call=True,
        )
        def handle_csv(
            contents,
            clear_click,
            del_clicks,
            refresh_click,
            filenames,
            files,
            refresh_counter,
            assignments,
        ):
            ctx = callback_context
            trigger = ctx.triggered[0]["prop_id"] if ctx.triggered else ""
            files = files or []
            refresh_counter = refresh_counter or 0
            assignments = assignments or {}

            if "btn-clear-csv" in trigger:
                self.data_manager.csv_file_paths = []
                self.data_manager.data_tables = []
                # CRITICAL: Clear derived signals and all caches
                self.derived_signals.clear()
                self.signal_properties.clear()
                self.original_file_paths.clear()
                # PERFORMANCE: Invalidate all caches when clearing data
                self.invalidate_caches()
                # Clear all assignments when clearing CSV files
                return [], [], "Cleared", refresh_counter + 1, {"0": {"0": []}}

            if "btn-refresh-csv" in trigger:
                # Refresh is handled by refresh_all_csvs callback - just return no_update here
                # This avoids two callbacks fighting over the same button
                return no_update, no_update, no_update, no_update, no_update

            if "del-csv" in trigger:
                for i, n in enumerate(del_clicks or []):
                    if n and i < len(files):
                        removed_path = files.pop(i)
                        
                        # Clear cache for this CSV
                        self.data_manager.invalidate_cache(csv_idx=i, clear_disk_cache=True)
                        
                        # Remove derived signals that depend on this CSV
                        derived_to_remove = []
                        for derived_name, derived_info in self.derived_signals.items():
                            source_signal_key = derived_info.get("source_signal_key")
                            source_csv_idx = derived_info.get("source_csv_idx")
                            
                            # Check if this derived signal depends on the removed CSV
                            if source_csv_idx == i:
                                derived_to_remove.append(derived_name)
                            elif source_signal_key and source_signal_key.startswith(f"{i}:"):
                                derived_to_remove.append(derived_name)
                            # Also check source_signals list for multi-op signals
                            elif "source_signals" in derived_info:
                                for sig_key in derived_info["source_signals"]:
                                    if sig_key.startswith(f"{i}:"):
                                        derived_to_remove.append(derived_name)
                                        break
                        
                        # Remove derived signals
                        for name in derived_to_remove:
                            del self.derived_signals[name]
                            if name in (derived or {}):
                                del derived[name]
                            print(f"[REMOVED] Derived signal '{name}' (source CSV deleted)")
                        
                        # Remove from data tables
                        if i < len(self.data_manager.data_tables):
                            self.data_manager.data_tables.pop(i)
                        if i < len(self.data_manager.last_read_rows):
                            self.data_manager.last_read_rows.pop(i)
                        if i < len(self.data_manager.last_file_mod_times):
                            self.data_manager.last_file_mod_times.pop(i)
                        
                        # Remove from original paths tracking
                        if removed_path in self.original_file_paths:
                            del self.original_file_paths[removed_path]
                        if i in self.data_manager.original_source_paths:
                            del self.data_manager.original_source_paths[i]
                        
                        # Remove assignments that reference this CSV
                        for tab_key in assignments:
                            for sp_key in assignments[tab_key]:
                                sp_assignment = assignments[tab_key][sp_key]
                                if isinstance(sp_assignment, list):
                                    assignments[tab_key][sp_key] = [
                                        s for s in sp_assignment 
                                        if s.get("csv_idx") != i
                                    ]
                        
                        self.data_manager.csv_file_paths = files
                        refresh_counter = refresh_counter + 1
                        print(f"[REMOVED] {os.path.basename(removed_path)}")
                        break

            # Browser upload disabled - using native file dialog instead
            if "upload-csv" in trigger:
                return no_update, no_update, "Use 'Browse Files' button", no_update, no_update

            # Build CSV list UI
            items = [
                html.Div(
                    [
                        html.I(className="fas fa-file me-1", style={"color": "#f4a261"}),
                        html.Span(
                            get_csv_display_name(f, files), 
                            style={"fontSize": "10px"},
                            title=f  # Show full path on hover
                        ),
                        html.A(
                            "×",
                            id={"type": "del-csv", "idx": i},
                            className="float-end text-danger",
                            style={"cursor": "pointer"},
                        ),
                    ],
                    className="py-1",
                )
                for i, f in enumerate(files)
            ]
            return files, items, f"{len(files)} CSV(s)", refresh_counter, dash.no_update

        # Search filter list management
        @self.app.callback(
            [
                Output("store-search-filters", "data"),
                Output("filter-list-display", "children"),
                Output("search-input", "value", allow_duplicate=True),
            ],
            [
                Input("btn-add-filter", "n_clicks"),
                Input({"type": "remove-filter", "idx": ALL}, "n_clicks"),
            ],
            [
                State("search-input", "value"),
                State("store-search-filters", "data"),
            ],
            prevent_initial_call=True,
        )
        def manage_filters(add_click, remove_clicks, search_value, filters):
            from dash import ctx

            filters = filters or []

            if ctx.triggered_id == "btn-add-filter" and add_click:
                # Add current search to filter list
                if (
                    search_value
                    and search_value.strip()
                    and search_value.strip() not in filters
                ):
                    filters.append(search_value.strip())

            elif (
                isinstance(ctx.triggered_id, dict)
                and ctx.triggered_id.get("type") == "remove-filter"
            ):
                # Remove filter by index
                idx = ctx.triggered_id.get("idx")
                if idx is not None and 0 <= idx < len(filters):
                    filters.pop(idx)

            # Build filter display
            if filters:
                filter_badges = [
                    dbc.Badge(
                        [
                            f,
                            html.Span(
                                " ×",
                                id={"type": "remove-filter", "idx": i},
                                style={"cursor": "pointer", "marginLeft": "5px"},
                            ),
                        ],
                        color="info",
                        className="me-1 mb-1",
                        style={"fontSize": "9px"},
                    )
                    for i, f in enumerate(filters)
                ]
                display = html.Div(
                    [
                        html.Small("Filters: ", className="text-muted"),
                        html.Span(filter_badges),
                    ]
                )
            else:
                display = html.Small("No active filters", className="text-muted")

            # Clear search input after adding
            clear_input = "" if ctx.triggered_id == "btn-add-filter" else dash.no_update

            return filters, display, clear_input

        # Signal tree with highlight selection - updates when data changes
        # NOTE: Collapse/expand is handled purely clientside via CSS for instant response
        @self.app.callback(
            [
                Output("signal-tree", "children"),
                Output("target-info", "children"),
                Output("highlight-count", "children"),
            ],
            [
                Input("store-csv-files", "data"),
                Input("search-input", "value"),
                Input("store-selected-tab", "data"),
                Input("store-selected-subplot", "data"),
                Input("store-assignments", "data"),
                Input("store-links", "data"),
                Input("store-signal-props", "data"),
                Input("store-highlighted", "data"),
                Input("store-refresh-trigger", "data"),
                Input("store-derived", "data"),
                Input("store-search-filters", "data"),
            ],
            [
                State("store-time-offsets", "data"),
            ],
            prevent_initial_call=False,
        )
        def update_tree(
            files,
            search,
            tab,
            subplot,
            assignments,
            links,
            props,
            highlighted,
            refresh_trigger,
            derived,
            search_filters,
            time_offsets,
        ):
            import time as _time
            _start = _time.perf_counter()
            
            try:
                tab = tab or 0
                subplot = subplot or 0
                target = f"Tab {tab+1}, Sub {subplot+1}"
                highlighted = highlighted or []

                assigned = set()
                if assignments:
                    assignment_data = assignments.get(str(tab), {}).get(
                        str(subplot), []
                    )
                    # Assignments are always a list in both time and X-Y modes
                    # FIX: Ensure all items are dicts with proper structure to prevent unhashable dict errors
                    if isinstance(assignment_data, list):
                        for s in assignment_data:
                            if isinstance(s, dict) and 'csv_idx' in s and 'signal' in s:
                                assigned.add(f"{s['csv_idx']}:{s['signal']}")

                linked_csvs = {}
                for lg in links or []:
                    for idx in lg.get("csv_indices", []):
                        linked_csvs[idx] = lg.get("name", "L")

                tree = []
                search = (search or "").lower()
                props = props or {}

                # Use files from store if available, otherwise use data_manager
                csv_files = files if files else self.data_manager.csv_file_paths

                # PERFORMANCE: Early return if no files
                if not csv_files:
                    return [
                        html.Div(
                            [
                                html.I(className="fas fa-info-circle me-2", style={"color": "#4ea8de"}),
                                html.Span("Upload CSV files to see signals", className="text-muted small"),
                            ],
                            className="p-2",
                        )
                    ], target, str(len(highlighted))

                # Check for duplicate filenames to determine display names
                basenames = [os.path.basename(fp) for fp in csv_files]
                duplicate_names = set(n for n in basenames if basenames.count(n) > 1)

                def get_display_name(fp):
                    """Get display name - include folder if filename is duplicate"""
                    basename = os.path.basename(fp)
                    if basename in duplicate_names:
                        # Include parent folder
                        parent = os.path.basename(os.path.dirname(fp))
                        return f"{parent}/{basename}" if parent else basename
                    return basename

                for csv_idx, fp in enumerate(csv_files):
                    while len(self.data_manager.data_tables) <= csv_idx:
                        self.data_manager.data_tables.append(None)

                    # PERFORMANCE: Only read if not already loaded
                    if self.data_manager.data_tables[csv_idx] is None:
                        try:
                            if os.path.exists(fp):
                                self.data_manager.read_initial_data(csv_idx)
                        except Exception as e:
                            logger.warning(f"Error reading CSV {csv_idx}: {e}")
                            continue

                    # PERFORMANCE: Use cached signal names
                    signals = self.get_signal_names_cached(csv_idx)
                    if not signals:
                        continue

                    fname = get_display_name(fp)
                    # Note: signals already filtered by get_signal_names_cached

                    # Apply search filters (filter list OR current search)
                    active_filters = list(search_filters or [])
                    if search and search.strip():
                        active_filters.append(search.strip())

                    if active_filters:
                        # Signal must match at least one filter
                        filtered_signals = []
                        for sig in signals:
                            sig_lower = sig.lower()
                            if any(f.lower() in sig_lower for f in active_filters):
                                filtered_signals.append(sig)
                        signals = filtered_signals

                    if not signals:
                        continue
                    
                    total_signals_in_csv = len(signals)
                    
                    # Get time offset for this CSV
                    time_offsets = time_offsets or {}
                    csv_offset = time_offsets.get(str(csv_idx), 0)

                    link_badge = (
                        dbc.Badge(
                            "🔗",
                            color="info",
                            className="ms-1",
                            pill=True,
                            style={"fontSize": "8px"},
                        )
                        if csv_idx in linked_csvs
                        else ""
                    )

                    # PERFORMANCE: Limit signals shown per CSV to reduce DOM elements
                    # Show all signals in original order - no priority sorting for performance
                    # No limit - show all signals for full visibility
                    hidden_count = 0  # No signals hidden

                    sig_items = []
                    for sig in signals:
                        key = f"{csv_idx}:{sig}"
                        is_assigned = key in assigned
                        is_highlighted = key in highlighted
                        p = props.get(key, {})
                        display = (
                            p.get("display_name", sig)
                            if p.get("show_in_tree", True)
                            else sig
                        )

                        sig_items.append(
                            html.Div(
                                [
                                    html.Div(
                                        [
                                            dbc.Checkbox(
                                                id={
                                                    "type": "sig-check",
                                                    "csv_idx": csv_idx,  # FIXED
                                                    "signal": sig,       # FIXED
                                                },
                                                value=is_assigned,
                                                className="d-inline me-1",
                                                label="",
                                            ),
                                            html.Small(
                                                "📊",
                                                className="text-muted me-1",
                                                style={"fontSize": "8px"},
                                                title="Assign to subplot",
                                            ),
                                            dbc.Checkbox(
                                                id={
                                                    "type": "sig-highlight",
                                                    "csv_idx": csv_idx,  # FIXED
                                                    "signal": sig,       # FIXED
                                                },
                                                value=is_highlighted,
                                                className="d-inline me-1",
                                                style={"accentColor": "#f4a261"},
                                                label="",
                                            ),
                                            html.Small(
                                                "⚙",
                                                className="text-warning me-2",
                                                style={"fontSize": "8px"},
                                                title="Select for operations",
                                            ),
                                            html.Span(
                                                display,
                                                style={
                                                    "fontSize": "10px",
                                                    "flex": "1",
                                                    "minWidth": "0",
                                                },
                                                className="me-2",
                                            ),
                                        ],
                                        className="d-flex align-items-center",
                                        style={"flex": "1", "minWidth": "0"},
                                    ),
                                    html.Div(
                                        [
                                            dbc.Button(
                                                "⚙",
                                                id={
                                                    "type": "sig-ops",
                                                    "csv_idx": csv_idx,  # FIXED
                                                    "signal": sig,       # FIXED
                                                },
                                                size="sm",
                                                color="link",
                                                className="p-0 me-1",
                                                style={
                                                    "fontSize": "9px",
                                                    "lineHeight": "1",
                                                },
                                                title="Single operation",
                                            ),
                                            dbc.Button(
                                                "✎",
                                                id={
                                                    "type": "sig-props",
                                                    "csv_idx": csv_idx,  # FIXED
                                                    "signal": sig,       # FIXED
                                                },
                                                size="sm",
                                                color="link",
                                                className="p-0",
                                                style={
                                                    "fontSize": "9px",
                                                    "lineHeight": "1",
                                                },
                                                title="Properties",
                                            ),
                                        ],
                                        className="d-flex align-items-center",
                                        style={"flexShrink": "0"},
                                    ),
                                ],
                                className="ms-2 py-1 px-1 d-flex justify-content-between align-items-center",
                                style={
                                    "backgroundColor": (
                                        "rgba(244,162,97,0.2)"
                                        if is_highlighted
                                        else "transparent"
                                    ),
                                    "borderRadius": "3px",
                                    "marginBottom": "2px",
                                },
                            )
                        )
                    
                    # PERFORMANCE: Show hidden count message if signals were truncated
                    if hidden_count > 0:
                        sig_items.append(
                            html.Div(
                                f"... {hidden_count} more signals (use search to filter)",
                                className="text-muted small fst-italic ms-2 py-1",
                                style={"fontSize": "9px"},
                            )
                        )

                    # Build offset badge if offset is set
                    offset_badge = ""
                    if csv_offset != 0:
                        offset_badge = dbc.Badge(
                            f"⏱{csv_offset:+.3g}",
                            color="warning",
                            className="ms-1",
                            pill=True,
                            style={"fontSize": "8px"},
                            title=f"Time offset: {csv_offset}s",
                        )
                    
                    # Collapsible CSV node with clickable header
                    # NOTE: Collapse state is managed purely clientside via CSS
                    tree.append(
                        html.Div(
                            [
                                # Clickable header - clientside JS handles toggle
                                html.Div(
                                    [
                                        html.Span(
                                            "▼",
                                            className="collapse-icon",
                                        ),
                                        html.I(
                                            className="fas fa-folder-open me-1",
                                            style={"color": "#f4a261"},
                                        ),
                                        html.Strong(fname, style={"fontSize": "10px"}),
                                        html.Small(
                                            f" ({total_signals_in_csv})",
                                            className="text-muted ms-1",
                                            style={"fontSize": "9px"},
                                        ),
                                        link_badge,
                                        offset_badge,
                                    ],
                                    className="csv-folder-header",
                                    title="Click to collapse/expand",
                                ),
                                # Signals list - CSS handles collapse animation
                                html.Div(
                                    sig_items,
                                    className="csv-signals-list",
                                ),
                            ],
                            className="csv-node",
                        )
                    )

                # Derived signals - with both checkboxes and operations
                derived = derived or {}
                if derived:
                    derived_items = []
                    for name in derived:
                        key = f"-1:{name}"
                        is_assigned = key in assigned
                        is_highlighted = key in highlighted
                        derived_items.append(
                            html.Div(
                                [
                                    html.Div(
                                        [
                                            dbc.Checkbox(
                                                id={
                                                    "type": "sig-check",
                                                    "csv_idx": -1,    # FIXED
                                                    "signal": name,   # FIXED
                                                },
                                                value=is_assigned,
                                                className="d-inline me-1",
                                                label="",
                                            ),
                                            html.Small(
                                                "📊",
                                                className="text-muted me-1",
                                                style={"fontSize": "8px"},
                                                title="Assign to subplot",
                                            ),
                                            dbc.Checkbox(
                                                id={
                                                    "type": "sig-highlight",
                                                    "csv_idx": -1,    # FIXED
                                                    "signal": name,   # FIXED
                                                },
                                                value=is_highlighted,
                                                className="d-inline me-1",
                                                style={"accentColor": "#f4a261"},
                                                label="",
                                            ),
                                            html.Small(
                                                "⚙",
                                                className="text-warning me-2",
                                                style={"fontSize": "8px"},
                                                title="Select for operations",
                                            ),
                                            html.Span(
                                                f"📐 {name}",
                                                style={
                                                    "fontSize": "10px",
                                                    "flex": "1",
                                                    "minWidth": "0",
                                                },
                                                className="me-2",
                                            ),
                                        ],
                                        className="d-flex align-items-center",
                                        style={"flex": "1", "minWidth": "0"},
                                    ),
                                    html.Div(
                                        [
                                            dbc.Button(
                                                "⚙",
                                                id={
                                                    "type": "sig-ops",
                                                    "csv_idx": -1,    # FIXED
                                                    "signal": name,   # FIXED
                                                },
                                                size="sm",
                                                color="link",
                                                className="p-0 me-1",
                                                style={
                                                    "fontSize": "9px",
                                                    "lineHeight": "1",
                                                },
                                                title="Single operation",
                                            ),
                                            dbc.Button(
                                                "✎",
                                                id={
                                                    "type": "sig-props",
                                                    "csv_idx": -1,    # FIXED
                                                    "signal": name,   # FIXED
                                                },
                                                size="sm",
                                                color="link",
                                                className="p-0 me-1",
                                                style={
                                                    "fontSize": "9px",
                                                    "lineHeight": "1",
                                                },
                                                title="Properties",
                                            ),
                                            dbc.Button(
                                                "×",
                                                id={
                                                    "type": "del-derived",
                                                    "name": name,
                                                },
                                                size="sm",
                                                color="link",
                                                className="p-0 text-danger",
                                                style={
                                                    "fontSize": "10px",
                                                    "lineHeight": "1",
                                                },
                                                title="Delete",
                                            ),
                                        ],
                                        className="d-flex align-items-center",
                                        style={"flexShrink": "0"},
                                    ),
                                ],
                                className="ms-2 py-1 px-1 d-flex justify-content-between align-items-center",
                                style={
                                    "backgroundColor": (
                                        "rgba(244,162,97,0.2)"
                                        if is_highlighted
                                        else "transparent"
                                    )
                                },
                            )
                        )

                    # Add the Derived node to tree (OUTSIDE the for loop)
                    if derived_items:
                        tree.append(
                            html.Div(
                                [
                                    html.Div(
                                        [
                                            html.I(
                                                className="fas fa-calculator me-1",
                                                style={"color": "#52b788"},
                                            ),
                                            html.Strong(
                                                "Derived", style={"fontSize": "11px"}
                                            ),
                                        ]
                                    ),
                                    html.Div(derived_items),
                                ],
                                className="mb-2 p-1 rounded",
                            )
                        )

                if not tree:
                    # Show helpful message when no signals
                    tree = [
                        html.Div(
                            [
                                html.I(
                                    className="fas fa-info-circle me-2",
                                    style={"color": "#4ea8de"},
                                ),
                                html.Span(
                                    "Upload CSV files to see signals",
                                    className="text-muted small",
                                ),
                            ],
                            className="p-2",
                        )
                    ]
                    if csv_files:
                        tree.append(
                            html.Small(
                                f"({len(csv_files)} files loaded but no signals found)",
                                className="text-warning d-block mt-1",
                            )
                        )
                        # Show more debug info
                        for i, fp in enumerate(csv_files):
                            has_data = (
                                i < len(self.data_manager.data_tables)
                                and self.data_manager.data_tables[i] is not None
                            )
                            status = "✓" if has_data else "✗"
                            tree.append(
                                html.Small(
                                    f"  {status} {os.path.basename(fp)}",
                                    className="text-muted d-block",
                                )
                            )

                # Performance timing
                _elapsed = (_time.perf_counter() - _start) * 1000
                if _elapsed > 100:
                    print(f"[PERF] 🐢 update_tree() completed: {_elapsed:.1f}ms ({len(tree)} items)")
                
                return tree, target, str(len(highlighted))
            except Exception as e:
                logger.exception(f"Error in update_tree: {e}")
                return (
                    [
                        html.Span(
                            f"Error: {str(e)}", className="text-danger small d-block"
                        ),
                        html.Small(
                            "Check console for details", className="text-muted d-block"
                        ),
                    ],
                    "Error",
                    "0",
                )

        # Handle highlighting for operations
        # FIX #4: Updated pattern-matching IDs
        @self.app.callback(
            Output("store-highlighted", "data"),
            [
                Input({"type": "sig-highlight", "csv_idx": ALL, "signal": ALL}, "value"),  # FIXED
                Input("btn-clear-highlight", "n_clicks"),
            ],
            [
                State({"type": "sig-highlight", "csv_idx": ALL, "signal": ALL}, "id"),  # FIXED
                State("store-highlighted", "data"),
            ],
            prevent_initial_call=True,
        )
        def handle_highlight(vals, clear_click, ids, current):
            ctx = callback_context
            trigger = ctx.triggered[0]["prop_id"] if ctx.triggered else ""

            if "btn-clear-highlight" in trigger:
                return []

            highlighted = []
            if vals and ids:
                for val, id_dict in zip(vals, ids):
                    if val:
                        highlighted.append(f"{id_dict.get('csv_idx', -1)}:{id_dict.get('signal', '')}")  # FIXED
            return highlighted

        # NOTE: CSV collapse/expand is now handled purely clientside via JavaScript
        # See assets/collapse.js - no server round-trip needed for instant response

        # Derived signals list (in Derived card) - simple list with delete buttons
        # Full controls (checkboxes, ops, edit) are in the signal tree
        @self.app.callback(
            Output("derived-list", "children"),
            [Input("store-derived", "data")],
            prevent_initial_call=False,
        )
        def update_derived_list(derived):
            if not derived:
                return [
                    html.Span(
                        "None - use signal operations to create",
                        className="text-muted small",
                    )
                ]

            # Just show names - delete functionality is in the signal tree
            items = []
            for name in derived.keys():
                items.append(
                    html.Div(
                        [
                            html.Span(f"📐 {name}", style={"fontSize": "10px"}),
                        ],
                        className="py-1",
                    )
                )
            return items

        # Delete derived signal (from tree or clear all button)
        @self.app.callback(
            [
                Output("store-derived", "data", allow_duplicate=True),
                Output("store-assignments", "data", allow_duplicate=True),
                Output("status-text", "children", allow_duplicate=True),
            ],
            [
                Input({"type": "del-derived", "name": ALL}, "n_clicks"),
                Input("btn-clear-derived", "n_clicks"),
            ],
            [
                State("store-derived", "data"),
                State("store-assignments", "data"),
            ],
            prevent_initial_call=True,
        )
        def delete_derived(del_clicks_tree, clear_click, derived, assignments):
            derived = derived or {}
            assignments = assignments or {}

            # Use triggered_id for pattern-matching callbacks (Dash 2.x)
            from dash import ctx

            if not ctx.triggered_id:
                return dash.no_update, dash.no_update, dash.no_update

            def remove_from_assignments(signal_name):
                """Remove a derived signal from all assignments"""
                for tab_key in assignments:
                    for sp_key in assignments[tab_key]:
                        assignments[tab_key][sp_key] = [
                            s
                            for s in assignments[tab_key][sp_key]
                            if not (
                                s.get("csv_idx") == -1
                                and s.get("signal") == signal_name
                            )
                        ]

            # Check if clear-all button was clicked
            if ctx.triggered_id == "btn-clear-derived" and clear_click:
                # Remove all derived signals from assignments
                for name in list(derived.keys()):
                    remove_from_assignments(name)
                self.derived_signals = {}
                return {}, assignments, "Cleared all derived"

            # Check if a specific delete button was clicked
            if (
                isinstance(ctx.triggered_id, dict)
                and ctx.triggered_id.get("type") == "del-derived"
            ):
                name = ctx.triggered_id.get("name", "")
                # Find the index of this button and check if it was actually clicked
                if del_clicks_tree:
                    for clicks in del_clicks_tree:
                        if clicks and clicks > 0:
                            if name and name in derived:
                                del derived[name]
                                remove_from_assignments(name)
                                self.derived_signals = derived
                                return derived, assignments, f"Deleted {name}"
                            break

            return dash.no_update, dash.no_update, dash.no_update

        # Update selected subplot store when dropdown changes
        @self.app.callback(
            Output("store-selected-subplot", "data", allow_duplicate=True),
            Input("subplot-select", "value"),
            prevent_initial_call=True,
        )
        def update_selected_subplot(val):
            return int(val) if val is not None else 0

        # Update subplot from typed input
        @self.app.callback(
            [
                Output("store-selected-subplot", "data", allow_duplicate=True),
                Output("subplot-select", "value", allow_duplicate=True),
                Output("subplot-input", "value"),
            ],
            Input("subplot-input", "value"),
            [
                State("rows-input", "value"),
                State("cols-input", "value"),
            ],
            prevent_initial_call=True,
        )
        def update_subplot_from_input(typed_value, rows, cols):
            if typed_value is None:
                return dash.no_update, dash.no_update, dash.no_update

            rows = int(rows or 1)
            cols = int(cols or 1)
            max_subplots = rows * cols

            # Convert 1-based input to 0-based index
            subplot_idx = int(typed_value) - 1

            # Clamp to valid range
            if subplot_idx < 0:
                subplot_idx = 0
            elif subplot_idx >= max_subplots:
                subplot_idx = max_subplots - 1

            return subplot_idx, str(subplot_idx), None  # Clear input after setting

        # Click on plot to select subplot
        # Uses clickData (click on trace), selectedData (selection), and relayoutData (double-click on axes)
        # CRITICAL: Handle plot interactions including double-click on subplot background
        @self.app.callback(
            [
                Output("store-selected-subplot", "data", allow_duplicate=True),
                Output("subplot-select", "value", allow_duplicate=True),
            ],
            [
                Input("plot", "clickData"),
                Input("plot", "selectedData"),
                Input("plot", "relayoutData"),  # Re-added for double-click detection
            ],
            [
                State("store-layouts", "data"),
                State("store-selected-tab", "data"),
                State("store-selected-subplot", "data"),
            ],
            prevent_initial_call=True,
        )
        def select_subplot_by_click(click_data, selected_data, relayout_data, layouts, tab_idx, current_subplot):
            """
            Handle subplot selection via:
            1. Single-click on trace
            2. Double-click on subplot (anywhere in subplot area)
            
            CRITICAL: Only process relayoutData for double-click events, not zoom/pan
            """
            ctx = callback_context
            if not ctx.triggered:
                return dash.no_update, dash.no_update

            trigger = ctx.triggered[0]["prop_id"]
            subplot_idx = None

            # Handle clickData (clicking on a trace)
            if "clickData" in trigger and click_data:
                point = click_data.get("points", [{}])[0]
                x_axis = point.get("xaxis", "x")

                if x_axis == "x":
                    subplot_idx = 0
                else:
                    try:
                        subplot_idx = int(x_axis.replace("x", "")) - 1
                    except:
                        subplot_idx = 0
            
            # Handle selectedData (box selection)
            elif "selectedData" in trigger and selected_data:
                point = selected_data.get("points", [{}])[0]
                x_axis = point.get("xaxis", "x")

                if x_axis == "x":
                    subplot_idx = 0
                else:
                    try:
                        subplot_idx = int(x_axis.replace("x", "")) - 1
                    except:
                        subplot_idx = 0
            
            # Handle relayoutData (double-click on subplot background)
            elif "relayoutData" in trigger and relayout_data:
                # CRITICAL: Only handle double-click events, ignore zoom/pan
                # Double-click events have specific keys like 'xaxis.autorange', 'yaxis.autorange'
                # Zoom/pan events have keys like 'xaxis.range[0]', 'xaxis.range[1]'
                
                # Check if this is a double-click (autorange reset) not a zoom/pan
                is_double_click = False
                clicked_axis = None
                
                for key in relayout_data.keys():
                    if 'autorange' in key.lower():
                        # This is a double-click event
                        is_double_click = True
                        # Extract axis name (e.g., 'xaxis2.autorange' -> 'xaxis2')
                        clicked_axis = key.split('.')[0]
                        break
                
                if is_double_click and clicked_axis:
                    # Determine subplot from axis name
                    if clicked_axis in ['xaxis', 'yaxis']:
                        subplot_idx = 0
                    else:
                        try:
                            # Extract number from 'xaxis2', 'yaxis3', etc.
                            axis_num = int(''.join(filter(str.isdigit, clicked_axis)))
                            subplot_idx = axis_num - 1
                        except:
                            subplot_idx = 0
                else:
                    # This is zoom/pan, not double-click - ignore it
                    return dash.no_update, dash.no_update
            
            else:
                # Unknown trigger - don't change subplot
                return dash.no_update, dash.no_update
            
            # Validate subplot index
            if subplot_idx is None:
                return dash.no_update, dash.no_update
            
            layouts = layouts or {}
            tab_key = str(tab_idx or 0)
            layout = layouts.get(tab_key, {"rows": 1, "cols": 1})
            rows = layout.get("rows", 1)
            cols = layout.get("cols", 1)
            max_subplots = rows * cols
            
            if subplot_idx < 0 or subplot_idx >= max_subplots:
                return dash.no_update, dash.no_update
            
            return subplot_idx, subplot_idx

        # FIX: Update signal tree and assigned list highlights when subplot changes via double-click
        # This ensures the UI reflects the selected subplot by triggering refresh
        @self.app.callback(
            Output("store-refresh-trigger", "data", allow_duplicate=True),
            Input("store-selected-subplot", "data"),
            prevent_initial_call=True,
        )
        def trigger_refresh_on_subplot_change(sel_subplot):
            """Trigger refresh when subplot changes to update highlights"""
            # Return a new value to trigger dependent callbacks
            import time
            return int(time.time() * 1000)

        # Handle signal assignments with proper linking (fixed bidirectional issue)
        # FIX #2: Changed pattern-matching IDs from csv/sig to csv_idx/signal
        @self.app.callback(
            [
                Output("store-assignments", "data", allow_duplicate=True),
                Output("plot", "figure"),
                Output("assigned-list", "children"),
                Output("store-selected-tab", "data", allow_duplicate=True),
                Output("store-selected-subplot", "data", allow_duplicate=True),
                Output("store-layouts", "data", allow_duplicate=True),
                Output("store-x-axis-signal", "data", allow_duplicate=True),
            ],
            [
                Input({"type": "sig-check", "csv_idx": ALL, "signal": ALL}, "value"),  # FIXED
                Input("tabs", "value"),
                Input("subplot-select", "value"),
                Input("store-selected-subplot", "data"),  # ADDED: Handle double-click subplot selection
                Input("rows-input", "value"),
                Input("cols-input", "value"),
                Input("btn-remove", "n_clicks"),
                Input("theme-switch", "value"),
                Input("store-derived", "data"),
                Input("store-signal-props", "data"),
                Input("link-axes-check", "value"),
                Input("store-refresh-trigger", "data"),
                Input("store-x-axis-signal", "data"),  # ADDED: Handle X-Y mode X-axis changes
                Input("time-cursor-check", "value"),
                Input("store-cursor-x", "data"),
                Input("store-subplot-modes", "data"),  # Listen to mode changes
                Input("store-time-columns", "data"),
                Input("store-subplot-metadata", "data"),
                Input("store-time-offsets", "data"),
                Input("store-display-options", "data"),
                Input("store-annotations", "data"),
            ],
            [
                State({"type": "sig-check", "csv_idx": ALL, "signal": ALL}, "id"),  # FIXED
                State("store-assignments", "data"),
                State("store-layouts", "data"),
                State("store-selected-tab", "data"),
                State("store-links", "data"),
                State({"type": "remove-check", "idx": ALL}, "value"),
                State({"type": "remove-check", "idx": ALL}, "id"),
            ],
            prevent_initial_call=True,
        )
        def handle_assignments(
            check_vals,
            active_tab,
            subplot_val,
            sel_subplot_store,  # NEW: Input from store-selected-subplot (double-click)
            rows,
            cols,
            remove_click,
            is_dark,
            derived,
            props,
            link_axes,
            refresh_trigger,
            x_axis_signals,  # NEW: Input from store-x-axis-signal (X-Y mode changes)
            time_cursor,
            cursor_data,
            subplot_modes,
            time_columns,
            subplot_metadata,
            time_offsets,
            display_options,
            annotations_data,
            check_ids,
            assignments,
            layouts,
            sel_tab,
            links,
            remove_vals,
            remove_ids,
        ):
            # Performance timing
            import time as _time
            _start = _time.perf_counter()
            
            # FIX #9: Add detailed logging to catch unhashable dict errors
            try:
                logger.debug(f"handle_assignments called with check_vals type: {type(check_vals)}")
                logger.debug(f"check_ids type: {type(check_ids)}, length: {len(check_ids) if check_ids else 0}")
                if check_ids and len(check_ids) > 0:
                    logger.debug(f"First check_id: {check_ids[0]}, type: {type(check_ids[0])}")
            except Exception as log_err:
                logger.error(f"Error in initial logging: {log_err}")
            
            ctx = callback_context

            if not ctx.triggered:
                trigger = ""
            else:
                trigger = ctx.triggered[0]["prop_id"] if ctx.triggered else ""
            
            print(f"[PERF] 🔧 handle_assignments() triggered by: {trigger[:50]}...")

            # CRITICAL: Defensive checks to prevent callback errors
            # FIX #6: Added extra safety checks
            assignments = assignments or {}
            layouts = layouts or {}
            derived = derived or {}
            props = props or {}
            subplot_modes = subplot_modes or {}
            time_columns = time_columns or {}
            subplot_metadata = subplot_metadata or {}
            time_offsets = time_offsets or {}
            display_options = display_options or {}
            annotations_data = annotations_data or {}
            x_axis_signals = x_axis_signals or {}
            links = links or []
            check_vals = check_vals or []  # ADDED: Prevent None errors
            check_ids = check_ids or []    # ADDED: Prevent None errors
            remove_vals = remove_vals or []  # ADDED: Prevent None errors
            remove_ids = remove_ids or []    # ADDED: Prevent None errors
            
            theme = "dark" if is_dark else "light"

            # ============================================================
            # PERFORMANCE OPTIMIZATION: Fast paths for specific changes
            # ============================================================
            
            # Fast path 1: Cursor updates - rebuild figure but with cached signal data
            # The actual optimization is in create_figure which caches signal data
            # We still need to rebuild to show cursor position + values
            
            # Fast path 2: Subplot selection changes - use Patch
            # When only subplot selection changes, update assigned-list and
            # subplot highlight using Patch (no full figure rebuild)
            
            # DEBUG: Log the trigger to understand what's happening
            if trigger:
                logger.info(f"handle_assignments triggered by: {trigger}")
            
            # CRITICAL: Check if this is a refresh/clear operation - always force full rebuild
            is_refresh_or_clear = "store-refresh-trigger" in trigger
            
            # Check if assignments are effectively empty (data cleared)
            has_any_assignments = False
            for tab_key, tab_data in assignments.items():
                if isinstance(tab_data, dict):
                    for sp_key, sp_data in tab_data.items():
                        if isinstance(sp_data, list) and len(sp_data) > 0:
                            has_any_assignments = True
                            break
                if has_any_assignments:
                    break
            
            # Check if this is ONLY a tab selection change (no other changes)
            is_tab_only = (
                trigger == "tabs.value"
                and active_tab is not None
                and not is_refresh_or_clear  # Never use fast path on refresh
            )
            
            # Check if this is ONLY a subplot selection change (no other changes)
            is_subplot_only = (
                trigger == "subplot-select.value" 
                and subplot_val is not None
                and not is_refresh_or_clear  # Never use fast path on refresh
            )
            
            # DEBUG: Log why fast path is or isn't being used
            logger.info(f"  → trigger == 'subplot-select.value': {trigger == 'subplot-select.value'}")
            logger.info(f"  → trigger == 'tabs.value': {trigger == 'tabs.value'}")
            logger.info(f"  → subplot_val is not None: {subplot_val is not None} (value={subplot_val})")
            logger.info(f"  → is_subplot_only: {is_subplot_only}")
            logger.info(f"  → is_tab_only: {is_tab_only}")
            
            # FAST PATH 1: Tab-only change - DISABLED because it causes issues when switching tabs
            # Tab switches need to rebuild the figure to show the correct tab's data
            if is_tab_only and False:  # DISABLED
                logger.info(f"  → Using TAB-ONLY FAST PATH (Patch highlight only)")
                # Extract tab number from active_tab (e.g., "tab-0" -> 0)
                tab_idx = int(active_tab.split("-")[1]) if "-" in active_tab else 0
                tab_key = str(tab_idx)
                
                # Get current subplot for this tab (default to 0)
                # FIX: Use sel_subplot_store which is the input from store-selected-subplot
                current_subplot = sel_subplot_store if sel_subplot_store is not None else 0
                subplot_key = str(current_subplot)
                
                # Ensure tab exists in data structures
                if tab_key not in assignments:
                    assignments[tab_key] = {}
                if tab_key not in layouts:
                    layouts[tab_key] = {"rows": 1, "cols": 1}
                
                # Get layout for this tab
                layout = layouts.get(tab_key, {"rows": 1, "cols": 1})
                rows = layout.get("rows", 1)
                cols = layout.get("cols", 1)
                total_subplots = rows * cols
                
                # CRITICAL: Only use Patch if we have a cached figure AND layout matches
                # If layout changed, we need to rebuild (but this shouldn't happen on tab switch)
                if hasattr(self, '_last_figure') and self._last_figure is not None:
                    # Check if layout matches cached figure
                    cached_layout = getattr(self, '_last_layout', None)
                    if cached_layout and cached_layout != layout:
                        # Layout changed - need to rebuild (fall through)
                        logger.info(f"  → Layout changed, cannot use Patch - need rebuild")
                        pass
                    else:
                        # Layout matches - can use Patch
                        # Build assigned-list for this subplot
                        assigned = assignments.get(tab_key, {}).get(subplot_key, [])
                        items = []
                        
                        if isinstance(assigned, list):
                            csv_paths = self.data_manager.csv_file_paths
                            for i, s in enumerate(assigned):
                                csv_idx = s.get("csv_idx", -1)
                                sig_name = s.get("signal", "")
                                if csv_idx == -1:
                                    lbl = f"{sig_name} (D)"
                                else:
                                    if csv_idx < len(csv_paths):
                                        csv_path = csv_paths[csv_idx]
                                        csv_filename = os.path.splitext(os.path.basename(csv_path))[0]
                                    else:
                                        csv_filename = f"C{csv_idx+1}"
                                    lbl = f"{sig_name} ({csv_filename})"
                                items.append(
                                    dbc.Checkbox(
                                        id={"type": "remove-check", "idx": i},
                                        label=lbl,
                                        value=False,
                                        style={"fontSize": "10px"},
                                    )
                                )
                        
                        if not items:
                            items = [html.Span("None", className="text-muted small")]
                        
                        # Get theme from is_dark parameter
                        theme_name = "dark" if is_dark else "light"
                        colors = THEMES[theme_name]
                        fig_patch = Patch()
                        
                        # Update highlight for all subplots
                        for i in range(total_subplots):
                            is_selected = (i == current_subplot)
                            border_color = "#4ea8de" if is_selected else colors["grid"]
                            border_width = 3 if is_selected else 1
                            subplot_bg = "rgba(78, 168, 222, 0.08)" if is_selected else colors["plot_bg"]
                            
                            # Axis naming: xaxis, xaxis2, xaxis3, etc.
                            x_key = "xaxis" if i == 0 else f"xaxis{i + 1}"
                            y_key = "yaxis" if i == 0 else f"yaxis{i + 1}"
                            
                            fig_patch["layout"][x_key]["linecolor"] = border_color
                            fig_patch["layout"][x_key]["linewidth"] = border_width
                            fig_patch["layout"][y_key]["linecolor"] = border_color
                            fig_patch["layout"][y_key]["linewidth"] = border_width
                            fig_patch["layout"][y_key]["plotbackgroundcolor"] = subplot_bg
                        
                        return (
                            no_update,    # assignments - no change
                            fig_patch,    # figure - Patch for highlight only
                            items,        # assigned-list - update this
                            tab_idx,      # sel_tab - update (use tab_idx, not sel_tab)
                            no_update,    # sel_subplot - keep current
                            no_update,    # layouts
                            no_update,    # x_axis_signals
                        )
                
                # No cached figure - need to rebuild figure for the new tab
                logger.info(f"  → No cached figure - rebuilding for tab {tab_key}")
                # Fall through to normal rebuild but don't return early
            
            # FAST PATH 2: Subplot-only change - update highlight with Patch, no rebuild
            if is_subplot_only:
                logger.info(f"  → Using FAST PATH (no figure update at all)")
                logger.info(f"  → subplot_val={subplot_val}, sel_tab={sel_tab}, active_tab={active_tab}")
                # Fast path: Update assigned-list ONLY, don't touch the figure
                sel_tab = sel_tab or 0
                if active_tab and "-" in active_tab:
                    sel_tab = int(active_tab.split("-")[1])
                    
                sel_subplot = int(subplot_val) if subplot_val is not None else 0
                tab_key = str(sel_tab)
                subplot_key = str(sel_subplot)
                
                logger.info(f"  → FAST PATH: tab_key={tab_key}, subplot_key={subplot_key}")
                
                # Build assigned-list for this subplot
                assigned = assignments.get(tab_key, {}).get(subplot_key, [])
                items = []
                
                if isinstance(assigned, list):
                    csv_paths = self.data_manager.csv_file_paths
                    items = []
                    for i, s in enumerate(assigned):
                        csv_idx = s.get("csv_idx", -1)
                        sig_name = s.get("signal", "")
                        if csv_idx == -1:
                            lbl = f"{sig_name} (D)"
                        else:
                            if csv_idx < len(csv_paths):
                                csv_path = csv_paths[csv_idx]
                                csv_filename = os.path.splitext(os.path.basename(csv_path))[0]
                            else:
                                csv_filename = f"C{csv_idx+1}"
                            lbl = f"{sig_name} ({csv_filename})"
                        items.append(
                            dbc.Checkbox(
                                id={"type": "remove-check", "idx": i},
                                label=lbl,
                                value=False,
                                style={"fontSize": "10px"},
                            )
                        )
                
                if not items:
                    items = [html.Span("None", className="text-muted small")]
                
                # Use Patch to update subplot highlight only (no rebuild!)
                if hasattr(self, '_last_figure') and self._last_figure is not None:
                    # Get layout for current tab
                    layout = layouts.get(tab_key, {"rows": 1, "cols": 1})
                    rows = layout.get("rows", 1)
                    cols = layout.get("cols", 1)
                    total_subplots = rows * cols
                    
                    # Get theme from is_dark parameter
                    theme_name = "dark" if is_dark else "light"
                    colors = THEMES[theme_name]
                    fig_patch = Patch()
                    
                    # Update highlight for all subplots
                    for i in range(total_subplots):
                        is_selected = (i == sel_subplot)
                        border_color = "#4ea8de" if is_selected else colors["grid"]
                        border_width = 3 if is_selected else 1
                        subplot_bg = "rgba(78, 168, 222, 0.08)" if is_selected else colors["plot_bg"]
                        
                        # Axis naming: xaxis, xaxis2, xaxis3, etc.
                        x_key = "xaxis" if i == 0 else f"xaxis{i + 1}"
                        y_key = "yaxis" if i == 0 else f"yaxis{i + 1}"
                        
                        fig_patch["layout"][x_key]["linecolor"] = border_color
                        fig_patch["layout"][x_key]["linewidth"] = border_width
                        fig_patch["layout"][y_key]["linecolor"] = border_color
                        fig_patch["layout"][y_key]["linewidth"] = border_width
                        fig_patch["layout"][y_key]["plotbackgroundcolor"] = subplot_bg
                    
                    return (
                        no_update,    # assignments - no change
                        fig_patch,    # figure - Patch for highlight only
                        items,        # assigned-list - update this
                        no_update,    # sel_tab
                        sel_subplot,  # sel_subplot - update
                        no_update,    # layouts
                        no_update,    # x_axis_signals
                    )
                else:
                    # No cached figure - return no_update to avoid white flash
                    return (
                        no_update,    # assignments - no change
                        no_update,    # figure - DON'T UPDATE (no white flash!)
                        items,        # assigned-list - update this
                        no_update,    # sel_tab
                        sel_subplot,  # sel_subplot - update
                        no_update,    # layouts
                        no_update,    # x_axis_signals
                    )

            if "tabs" in trigger and active_tab:
                sel_tab = int(active_tab.split("-")[1]) if "-" in active_tab else 0
            elif active_tab:
                sel_tab = int(active_tab.split("-")[1]) if "-" in active_tab else 0
            else:
                sel_tab = sel_tab or 0

            # Track if we should output sel_subplot or use no_update
            output_subplot = False

            # Determine current subplot from dropdown or store (store updated by double-click)
            if "subplot-select" in trigger and subplot_val is not None:
                # Dropdown changed
                try:
                    sel_subplot = int(subplot_val)
                except (ValueError, TypeError):
                    sel_subplot = 0
                output_subplot = True
            elif "store-selected-subplot" in trigger and sel_subplot_store is not None:
                # Double-click changed store
                try:
                    sel_subplot = int(sel_subplot_store)
                except (ValueError, TypeError):
                    sel_subplot = 0
                output_subplot = True
            elif sel_subplot_store is not None:
                # Use store value as current
                try:
                    sel_subplot = int(sel_subplot_store)
                except (ValueError, TypeError):
                    sel_subplot = 0
            elif subplot_val is not None:
                # Use dropdown value
                try:
                    sel_subplot = int(subplot_val)
                except (ValueError, TypeError):
                    sel_subplot = 0
            else:
                # Default
                sel_subplot = 0

            tab_key = str(sel_tab)
            subplot_key = str(sel_subplot)
            
            # CRITICAL: Ensure tab exists in all data structures
            # FIX: Add extra validation for tab "0" and ensure proper structure
            # This is especially important for tab "0" which can get corrupted
            if not isinstance(assignments, dict):
                assignments = {}
            if tab_key not in assignments:
                assignments[tab_key] = {}
            elif not isinstance(assignments[tab_key], dict):
                # Fix corrupted tab structure - especially common for tab "0"
                assignments[tab_key] = {}
            
            if tab_key not in layouts:
                layouts[tab_key] = {"rows": 1, "cols": 1}
            elif not isinstance(layouts[tab_key], dict):
                # Fix corrupted layout structure
                layouts[tab_key] = {"rows": 1, "cols": 1}

            # Get subplot modes for current tab FIRST (needed for assignment initialization)
            tab_subplot_modes = (subplot_modes or {}).get(tab_key, {})
            current_mode = tab_subplot_modes.get(subplot_key, "time")

            # Initialize assignment as list (same format for both time and xy modes)
            # FIX: Extra validation to prevent corruption, especially for tab "0"
            if subplot_key not in assignments[tab_key]:
                assignments[tab_key][subplot_key] = []
            else:
                # Sanitize existing assignment to ensure it's a proper list
                current_assignment = assignments[tab_key][subplot_key]
                if isinstance(current_assignment, dict):
                    # Convert old dict format to list
                    assignments[tab_key][subplot_key] = []
                elif isinstance(current_assignment, list):
                    # Validate list items are proper dicts
                    sanitized = [
                        s for s in current_assignment
                        if isinstance(s, dict) and 'csv_idx' in s and 'signal' in s
                    ]
                    assignments[tab_key][subplot_key] = sanitized
                else:
                    # Invalid format - reset to empty list
                    assignments[tab_key][subplot_key] = []

            rows = int(rows) if rows else 1
            cols = int(cols) if cols else 1

            # Get old layout to remap subplots if layout changed
            old_layout = layouts.get(tab_key, {"rows": 1, "cols": 1})
            old_rows = old_layout.get("rows", 1)
            old_cols = old_layout.get("cols", 1)

            # Remap assignments if layout changed (preserve row/col positions)
            if (
                (old_rows != rows or old_cols != cols)
                and "rows-input" in trigger
                or "cols-input" in trigger
            ):
                if tab_key in assignments:
                    old_assignments = assignments[tab_key].copy()
                    new_assignments = {}

                    for old_sp_key, signals in old_assignments.items():
                        old_sp = int(old_sp_key)
                        # Calculate old row/col
                        old_row = old_sp // old_cols
                        old_col = old_sp % old_cols
                        # Calculate new index (preserving row/col position)
                        if old_row < rows and old_col < cols:
                            new_sp = old_row * cols + old_col
                            new_sp_key = str(new_sp)
                            if new_sp_key not in new_assignments:
                                new_assignments[new_sp_key] = []
                            new_assignments[new_sp_key].extend(signals)

                    assignments[tab_key] = new_assignments

            layouts[tab_key] = {"rows": rows, "cols": cols}

            if sel_subplot >= rows * cols:
                sel_subplot = 0
                subplot_key = "0"

            # Handle signal checkbox changes - FIXED linking logic
            # FIX #3: Updated to use csv_idx/signal instead of csv/sig
            # FIX #7: Added try-catch to prevent corruption on errors
            if (
                check_vals is not None
                and check_ids is not None
                and len(check_vals) > 0
                and "sig-check" in trigger
            ):
                try:
                    # Find which checkbox was actually clicked
                    clicked_csv = None
                    clicked_sig = None
                    clicked_new_val = None

                    import json as js

                    try:
                        id_str = trigger.split(".")[0]
                        clicked_id = js.loads(id_str)
                        clicked_csv = clicked_id.get("csv_idx")  # FIXED: was "csv"
                        clicked_sig = clicked_id.get("signal")   # FIXED: was "sig"
                        # Find the value for this specific checkbox
                        for val, id_dict in zip(check_vals, check_ids):
                            if (
                                id_dict.get("csv_idx") == clicked_csv  # FIXED
                                and id_dict.get("signal") == clicked_sig  # FIXED
                            ):
                                clicked_new_val = val
                                break
                    except Exception as e:
                        logger.warning(f"Error parsing clicked ID: {e}")

                    # Handle X-Y mode assignment
                    # Both time and X-Y modes use the same list-based assignment
                    # In X-Y mode, X-axis is selected separately via xy-x-select dropdown
                    if True:
                        # List-based signal assignment logic
                        current = assignments[tab_key][subplot_key]
                        if not isinstance(current, list):
                            current = []
                            assignments[tab_key][subplot_key] = current
                        # FIX: Ensure all items are dicts with proper structure to prevent unhashable dict errors
                        current = [
                            s for s in current 
                            if isinstance(s, dict) and 'csv_idx' in s and 'signal' in s
                        ]
                        assignments[tab_key][subplot_key] = current
                        current_keys = {f"{s['csv_idx']}:{s['signal']}" for s in current}

                        # Only process the clicked checkbox and apply linking
                        if clicked_csv is not None and clicked_sig is not None:
                            csv_idx = clicked_csv
                            sig = clicked_sig
                            key = f"{csv_idx}:{sig}"

                            is_currently_assigned = key in current_keys
                            should_be_assigned = (
                                clicked_new_val if clicked_new_val else False
                            )

                            if should_be_assigned and not is_currently_assigned:
                                # ADD signal - check for linked CSVs
                                linked_indices = [csv_idx]
                                for lg in links or []:
                                    if csv_idx in lg.get("csv_indices", []):
                                        linked_indices = lg["csv_indices"]
                                        break

                                # Add clicked signal and all linked versions
                                for linked_csv_idx in linked_indices:
                                    linked_key = f"{linked_csv_idx}:{sig}"
                                    if linked_key not in current_keys:
                                        if linked_csv_idx >= 0 and linked_csv_idx < len(
                                            self.data_manager.data_tables
                                        ):
                                            df = self.data_manager.data_tables[
                                                linked_csv_idx
                                            ]
                                            if df is not None and sig in df.columns:
                                                assignments[tab_key][subplot_key].append(
                                                    {
                                                        "csv_idx": linked_csv_idx,
                                                        "signal": sig,
                                                    }
                                                )
                                                current_keys.add(linked_key)
                                        elif linked_csv_idx == -1:
                                            # FIX #8: Use correct reference to derived signals
                                            if sig in (derived or {}):
                                                assignments[tab_key][subplot_key].append(
                                                    {"csv_idx": -1, "signal": sig}
                                                )
                                                current_keys.add(linked_key)

                            elif not should_be_assigned and is_currently_assigned:
                                # REMOVE signal - check for linked CSVs
                                linked_indices = [csv_idx]
                                for lg in links or []:
                                    if csv_idx in lg.get("csv_indices", []):
                                        linked_indices = lg["csv_indices"]
                                        break

                                # Remove clicked signal and all linked versions
                                for linked_csv_idx in linked_indices:
                                    linked_key = f"{linked_csv_idx}:{sig}"
                                    assignments[tab_key][subplot_key] = [
                                        s
                                        for s in assignments[tab_key][subplot_key]
                                        if f"{s['csv_idx']}:{s['signal']}" != linked_key
                                    ]

                except Exception as e:
                    # Enable error to see details - don't hide it
                    import traceback
                    error_msg = f"Error in signal assignment: {e}\nTraceback: {traceback.format_exc()}\nClicked CSV: {clicked_csv}, Signal: {clicked_sig}"
                    logger.error(error_msg)
                    # Re-raise to see the actual error
                    raise

            if "btn-remove" in trigger and remove_vals:
                # FIX #5: Use .get() to safely extract idx values, not dict objects
                to_remove = {
                    id_dict.get("idx", -1)  # FIXED: Extract int, not dict
                    for val, id_dict in zip(remove_vals, remove_ids)
                    if val
                }

                assignments[tab_key][subplot_key] = [
                    s
                    for i, s in enumerate(assignments[tab_key][subplot_key])
                    if i not in to_remove
                ]

            # After any assignment changes, check if X-axis signal still exists
            # If not, reset to time
            # FIX: Ensure we always work with a dict, not None
            if x_axis_signals is None:
                x_axis_signals = {}
            elif not isinstance(x_axis_signals, dict):
                logger.warning(f"x_axis_signals is not a dict: {type(x_axis_signals)}, resetting")
                x_axis_signals = {}
            
            # FIX: Check and fix corrupted structure where tab_key value is a string instead of dict
            if tab_key in x_axis_signals:
                tab_val = x_axis_signals[tab_key]
                if not isinstance(tab_val, dict):
                    # CORRUPTED: tab_key value is a string or other non-dict type
                    logger.error(f"CRITICAL: x_axis_signals[{tab_key}] is corrupted: {type(tab_val)} = {tab_val}")
                    if isinstance(tab_val, str):
                        # Old/corrupted format: {'0': 'time'} -> should be {'0': {'0': 'time'}}
                        # Assume it was for subplot 0
                        logger.error(f"  Fixing: Converting string value '{tab_val}' to dict structure")
                        x_axis_signals[tab_key] = {"0": tab_val}  # FIX: Keep the value for subplot 0
                    else:
                        x_axis_signals[tab_key] = {"0": "time"}  # FIX: Default structure
            
            if tab_key not in x_axis_signals:
                x_axis_signals[tab_key] = {}
            
            # FIX: Ensure the subplot_key exists and is a string, not a dict
            if subplot_key not in x_axis_signals[tab_key]:
                x_axis_signals[tab_key][subplot_key] = "time"
            elif not isinstance(x_axis_signals[tab_key][subplot_key], str):
                # CORRUPTED: subplot_key value is not a string
                logger.error(f"CRITICAL: x_axis_signals[{tab_key}][{subplot_key}] is not a string: {type(x_axis_signals[tab_key][subplot_key])} = {x_axis_signals[tab_key][subplot_key]}")
                x_axis_signals[tab_key][subplot_key] = "time"

            # FIX: Double-check that tab_key value is a dict before accessing subplot_key
            tab_x_dict = x_axis_signals.get(tab_key, {})
            if not isinstance(tab_x_dict, dict):
                logger.error(f"CRITICAL: x_axis_signals[{tab_key}] is still not a dict after fix: {type(tab_x_dict)} = {tab_x_dict}")
                tab_x_dict = {}
                x_axis_signals[tab_key] = tab_x_dict
            
            # FIX: Get current_x_axis and ensure it's a string, not a dict
            current_x_axis_raw = tab_x_dict.get(subplot_key, "time")
            if isinstance(current_x_axis_raw, dict):
                # CORRUPTED: subplot_key value is a dict instead of string
                logger.error(f"CRITICAL: x_axis_signals[{tab_key}][{subplot_key}] is a dict: {current_x_axis_raw}, resetting to 'time'")
                current_x_axis = "time"
                x_axis_signals[tab_key][subplot_key] = "time"
            elif not isinstance(current_x_axis_raw, str):
                logger.error(f"CRITICAL: current_x_axis is not a string: {type(current_x_axis_raw)} = {current_x_axis_raw}, resetting to 'time'")
                current_x_axis = "time"
                x_axis_signals[tab_key][subplot_key] = "time"
            else:
                current_x_axis = current_x_axis_raw
            
            if current_x_axis != "time":
                # Check if this signal is still in assignments
                # FIX: Ensure all items are dicts with proper structure to prevent unhashable dict errors
                assignment_list = assignments.get(tab_key, {}).get(subplot_key, [])
                if not isinstance(assignment_list, list):
                    assignment_list = []
                try:
                    assigned_signal_keys = {
                        f"{s['csv_idx']}:{s['signal']}"
                        for s in assignment_list
                        if isinstance(s, dict) and 'csv_idx' in s and 'signal' in s
                    }
                    # FIX: Double-check current_x_axis is a string before using 'in' operator
                    if not isinstance(current_x_axis, str):
                        logger.error(f"CRITICAL: current_x_axis is not a string before 'in' check: {type(current_x_axis)} = {current_x_axis}")
                        current_x_axis = "time"
                        x_axis_signals[tab_key][subplot_key] = "time"
                    elif current_x_axis not in assigned_signal_keys:
                        # X-axis signal no longer assigned, reset to time
                        x_axis_signals[tab_key][subplot_key] = "time"
                except TypeError as e:
                    logger.error(f"CRITICAL: TypeError in assigned_signal_keys check: {e}")
                    logger.error(f"  current_x_axis type: {type(current_x_axis)}, value: {current_x_axis}")
                    logger.error(f"  assignment_list: {assignment_list}")
                    raise

            self.signal_properties = props or {}
            self.derived_signals = derived or {}

            # Get cursor X position
            cursor_x = None
            if time_cursor and cursor_data:
                cursor_x = cursor_data.get("x")

            # Get subplot modes for current tab
            # FIX: Ensure tab_subplot_modes is always a dict, not a string
            raw_tab_modes = (subplot_modes or {}).get(tab_key, {})
            if isinstance(raw_tab_modes, dict):
                tab_subplot_modes = raw_tab_modes
            else:
                logger.warning(f"tab_subplot_modes is not a dict for tab {tab_key}: {type(raw_tab_modes)} = {raw_tab_modes}")
                tab_subplot_modes = {}

            # Get X-axis signals for this tab
            # FIX: Ensure x_axis_signals is a dict before accessing
            if x_axis_signals is None:
                logger.warning(f"x_axis_signals is None when getting tab_x_axis_signals for tab {tab_key}")
                x_axis_signals = {}
            tab_x_axis_signals = x_axis_signals.get(tab_key, {})
            if not isinstance(tab_x_axis_signals, dict):
                logger.warning(f"tab_x_axis_signals is not a dict for tab {tab_key}: {type(tab_x_axis_signals)}")
                tab_x_axis_signals = {}

            # Get subplot metadata for this tab
            tab_subplot_metadata = (subplot_metadata or {}).get(tab_key, {})

            # Get time offsets from store
            time_offsets_data = time_offsets or {}
            
            # ================================================================
            # PERFORMANCE: Check if figure rebuild is actually needed
            # Skip expensive create_figure if only cursor or subplot selection changed
            # ================================================================
            
            # Cursor-only change: just update cursor position, not the whole figure
            cursor_only_triggers = ["store-cursor-x.data"]
            is_cursor_only = trigger in cursor_only_triggers
            
            # PERFORMANCE BOOST: Enhanced figure caching system
            # Compute state hash to check if we can reuse a cached figure
            logger.info(f"Rendering figure for tab {tab_key}, subplot {subplot_key}")
            logger.info(f"Assignments: {assignments.get(tab_key, {})}")

            import time
            render_start = time.time()

            current_state_hash = self._compute_figure_hash(
                assignments, tab_key, layouts, self.signal_properties,
                display_options, cursor_x, link_axes, time_columns,
                x_axis_signals=x_axis_signals, subplot_modes=subplot_modes
            )

            # Check if we can reuse cached figure
            # CRITICAL: Never use cache on refresh/clear - always rebuild
            cache_key = (int(tab_key), current_state_hash)
            use_cache = cache_key in self._figure_cache and not is_refresh_or_clear
            
            if use_cache:
                fig = self._figure_cache[cache_key]
                self._cache_hit_count += 1
                logger.info(f"[CACHE HIT] Reusing cached figure (0ms) - {self._cache_hit_count} hits so far")

                # CRITICAL FIX: Update subplot highlight even on cache hit
                # When only subplot selection changes, we need to update the border highlight
                # Handle both dropdown changes and double-click subplot selection
                is_subplot_change = (
                    "subplot-select" in trigger or
                    "store-selected-subplot" in trigger
                )

                # PERFORMANCE: Handle cursor-only changes efficiently
                if is_cursor_only and cursor_x is not None:
                    # Cursor only - update just the cursor line shapes without rebuilding
                    logger.info(f"[CACHE HIT + CURSOR ONLY] Fast cursor update via Patch")
                    fig_patch = Patch()
                    layout = layouts.get(tab_key, {"rows": 1, "cols": 1})
                    rows_layout = layout.get("rows", 1)
                    cols_layout = layout.get("cols", 1)
                    total_subplots = rows_layout * cols_layout
                    
                    # Update cursor line shapes (one per subplot)
                    for shape_idx in range(total_subplots):
                        fig_patch["layout"]["shapes"][shape_idx]["x0"] = cursor_x
                        fig_patch["layout"]["shapes"][shape_idx]["x1"] = cursor_x
                    
                    fig = fig_patch
                elif is_subplot_change:
                    logger.info(f"[CACHE HIT + SUBPLOT CHANGE] Updating highlight with Patch (trigger: {trigger})")
                    fig_patch = Patch()
                    layout = layouts.get(tab_key, {"rows": 1, "cols": 1})
                    rows_layout = layout.get("rows", 1)
                    cols_layout = layout.get("cols", 1)
                    total_subplots = rows_layout * cols_layout
                    theme_name = "dark" if is_dark else "light"
                    colors = THEMES[theme_name]

                    for i in range(total_subplots):
                        is_selected = (i == sel_subplot)
                        border_color = "#4ea8de" if is_selected else colors["grid"]
                        border_width = 3 if is_selected else 1
                        subplot_bg = "rgba(78, 168, 222, 0.08)" if is_selected else colors["plot_bg"]

                        x_key = "xaxis" if i == 0 else f"xaxis{i + 1}"
                        y_key = "yaxis" if i == 0 else f"yaxis{i + 1}"

                        fig_patch["layout"][x_key]["linecolor"] = border_color
                        fig_patch["layout"][x_key]["linewidth"] = border_width
                        fig_patch["layout"][y_key]["linecolor"] = border_color
                        fig_patch["layout"][y_key]["linewidth"] = border_width
                        fig_patch["layout"][y_key]["plotbackgroundcolor"] = subplot_bg

                    fig = fig_patch
            
            if not use_cache:
                # Cache miss or forced rebuild - need to build figure
                self._cache_miss_count += 1
                if is_refresh_or_clear:
                    logger.info(f"[FORCED REBUILD] Refresh/clear triggered - rebuilding figure")
                else:
                    logger.info(f"[CACHE MISS] Building new figure - {self._cache_miss_count} misses so far")

                try:
                    fig = self.create_figure(
                        rows,
                        cols,
                        theme,
                        sel_subplot,
                        assignments,
                        tab_key,
                        link_axes,
                        time_cursor,
                        cursor_x,
                        tab_subplot_modes,
                        tab_x_axis_signals,
                        time_columns,
                        tab_subplot_metadata,
                        time_offsets_data,
                        display_options,
                        annotations_data,
                    )

                    # CRITICAL: Ensure figure is valid before caching
                    is_valid = (
                        fig is not None
                        and hasattr(fig, 'data')
                        and hasattr(fig, 'layout')
                        and len(fig.data) > 0
                    )

                    if is_valid:
                        # PERFORMANCE BOOST: Store valid figure in cache
                        self._figure_cache[cache_key] = fig
                        self._state_hash_cache[int(tab_key)] = current_state_hash

                        # Limit cache size to prevent memory issues (max 20 figures)
                        if len(self._figure_cache) > 20:
                            oldest_key = min(self._figure_cache.keys(), key=lambda k: k[1])
                            del self._figure_cache[oldest_key]
                            logger.info(f"[CACHE] Evicted oldest entry, size: {len(self._figure_cache)}")

                        render_time_ms = (time.time() - render_start) * 1000
                        logger.info(f"[RENDER] Figure built and cached in {render_time_ms:.0f}ms")
                    else:
                        logger.warning(f"Figure invalid, data traces: {len(fig.data) if fig else 0}")
                        fig = dash.no_update

                except Exception as e:
                    logger.error(f"ERROR creating figure: {e}")
                    import traceback
                    logger.error(traceback.format_exc())
                    fig = dash.no_update

            # FIX: Sanitize assigned list to prevent unhashable dict errors
            tab_assignments = assignments.get(tab_key, {})
            if not isinstance(tab_assignments, dict):
                tab_assignments = {}
            assigned_raw = tab_assignments.get(subplot_key, [])
            
            # Ensure assigned is always a list with proper dict items
            if isinstance(assigned_raw, list):
                assigned = [
                    s for s in assigned_raw
                    if isinstance(s, dict) and 'csv_idx' in s and 'signal' in s
                ]
            elif isinstance(assigned_raw, dict):
                # Old dict format - convert to empty list
                assigned = []
            else:
                # Invalid format - use empty list
                assigned = []
            
            current_mode = tab_subplot_modes.get(subplot_key, "time")
            items = []

            # Handle X-Y mode display
            if current_mode == "xy" and isinstance(assigned, dict):
                # Show X and Y signal info
                x_info = assigned.get("x", {})
                y_info = assigned.get("y", {})

                for axis, info in [("X", x_info), ("Y", y_info)]:
                    if info:
                        csv_idx = info.get("csv_idx", -1)
                        sig_name = info.get("signal", "")
                        if csv_idx == -1:
                            lbl = f"{sig_name} (D)"
                        else:
                            if csv_idx < len(self.data_manager.csv_file_paths):
                                csv_path = self.data_manager.csv_file_paths[csv_idx]
                                csv_filename = os.path.splitext(
                                    os.path.basename(csv_path)
                                )[0]
                            else:
                                csv_filename = f"C{csv_idx+1}"
                            lbl = f"{sig_name} ({csv_filename})"

                        color = "info" if axis == "X" else "warning"
                        items.append(
                            html.Div(
                                [
                                    html.Span(
                                        f"{axis}: ",
                                        className=f"text-{color} fw-bold small",
                                    ),
                                    html.Span(lbl, className="small"),
                                    dbc.Button(
                                        "×",
                                        id={"type": "xy-remove", "axis": axis.lower()},
                                        size="sm",
                                        color="danger",
                                        outline=True,
                                        className="ms-2 py-0 px-1",
                                        style={"fontSize": "10px"},
                                    ),
                                ],
                                className="d-flex align-items-center mb-1",
                            )
                        )
                if not items:
                    items = [
                        html.Span(
                            "Assign X and Y signals", className="text-muted small"
                        )
                    ]
            else:
                # Time mode: show list of signals with checkboxes - DRAGGABLE
                if isinstance(assigned, list):
                    for i, s in enumerate(assigned):
                        csv_idx = s.get("csv_idx", -1)
                        sig_name = s.get("signal", "")
                        if csv_idx == -1:
                            lbl = f"{sig_name} (D)"
                        else:
                            # Get CSV filename for display
                            if csv_idx < len(self.data_manager.csv_file_paths):
                                csv_path = self.data_manager.csv_file_paths[csv_idx]
                                csv_filename = os.path.splitext(
                                    os.path.basename(csv_path)
                                )[0]
                            else:
                                csv_filename = f"C{csv_idx+1}"
                            lbl = f"{sig_name} ({csv_filename})"
                        # Wrap in draggable div for drag & drop + right-click support
                        items.append(
                            html.Div(
                                [
                                    html.Span("⋮⋮", className="drag-handle text-muted me-2", 
                                              style={"cursor": "grab", "fontSize": "10px", "opacity": "0.5"}),
                                    dbc.Checkbox(
                                        id={"type": "remove-check", "idx": i},
                                        label=lbl,
                                        value=False,
                                        style={"fontSize": "10px", "flex": "1"},
                                        className="d-inline-block",
                                    ),
                                ],
                                className="draggable-signal d-flex align-items-center py-1",
                                draggable="true",
                                **{
                                    "data-signal-idx": str(i),
                                    "data-signal-name": sig_name,
                                    "data-csv-idx": str(csv_idx),
                                }
                            )
                        )
                if not items:
                    items = [html.Span("None", className="text-muted small")]

            # Only output sel_subplot if explicitly changed, otherwise use no_update
            subplot_output = sel_subplot if output_subplot else dash.no_update
            
            # FIX: Ensure x_axis_signals is always a dict before returning (not None)
            # This is critical - if it's None, we lose all state
            if x_axis_signals is None:
                logger.error(f"CRITICAL: x_axis_signals is None when returning from handle_assignments for tab {tab_key}!")
                x_axis_signals = {}
            elif not isinstance(x_axis_signals, dict):
                logger.error(f"CRITICAL: x_axis_signals is not a dict when returning: {type(x_axis_signals)}")
                x_axis_signals = {}
            
            logger.info(f"Returning from handle_assignments: tab={tab_key}, subplot={subplot_key}")
            logger.info(f"  x_axis_signals keys: {list(x_axis_signals.keys())}")
            if tab_key in x_axis_signals:
                logger.info(f"  Tab {tab_key} x_axis_signals: {x_axis_signals[tab_key]}")
            
            # CRITICAL: If fig is None, use cached figure or no_update to prevent white plot
            # This should never happen, but safety check to prevent white plots
            if fig is None:
                logger.warning(f"CRITICAL: fig is None! Using cached figure or no_update to prevent white plot")
                if hasattr(self, '_last_figure') and self._last_figure is not None:
                    # Use cached figure with Patch for highlight update
                    fig = Patch()
                    layout = layouts.get(tab_key, {"rows": 1, "cols": 1})
                    rows = layout.get("rows", 1)
                    cols = layout.get("cols", 1)
                    total_subplots = rows * cols
                    theme_name = "dark" if is_dark else "light"
                    colors = THEMES[theme_name]
                    for i in range(total_subplots):
                        is_selected = (i == sel_subplot)
                        border_color = "#4ea8de" if is_selected else colors["grid"]
                        border_width = 3 if is_selected else 1
                        subplot_bg = "rgba(78, 168, 222, 0.08)" if is_selected else colors["plot_bg"]
                        x_key = "xaxis" if i == 0 else f"xaxis{i + 1}"
                        y_key = "yaxis" if i == 0 else f"yaxis{i + 1}"
                        fig["layout"][x_key]["linecolor"] = border_color
                        fig["layout"][x_key]["linewidth"] = border_width
                        fig["layout"][y_key]["linecolor"] = border_color
                        fig["layout"][y_key]["linewidth"] = border_width
                        fig["layout"][y_key]["plotbackgroundcolor"] = subplot_bg
                else:
                    # No cached figure - use no_update to prevent white plot
                    fig = dash.no_update
            
            return (
                assignments,
                fig,
                items,
                sel_tab,
                subplot_output,
                layouts,
                x_axis_signals,  # Return the actual dict, not x_axis_signals or {}
            )

        # Subplot selector
        @self.app.callback(
            Output("subplot-select", "options"),
            [Input("rows-input", "value"), Input("cols-input", "value")],
        )
        def update_subplot_options(rows, cols):
            rows, cols = int(rows or 1), int(cols or 1)
            return [{"label": str(i + 1), "value": i} for i in range(rows * cols)]

        # Sync rows/cols inputs when tab changes (each tab has its own layout)
        @self.app.callback(
            [
                Output("rows-input", "value"),
                Output("cols-input", "value"),
                Output("subplot-select", "value"),
            ],
            [Input("tabs", "value")],
            [State("store-layouts", "data")],
            prevent_initial_call=True,
        )
        def sync_layout_on_tab_change(active_tab, layouts):
            if not active_tab:
                return 1, 1, 0
            tab_idx = int(active_tab.split("-")[1]) if "-" in active_tab else 0
            tab_key = str(tab_idx)
            layouts = layouts or {}
            layout = layouts.get(tab_key, {"rows": 1, "cols": 1})
            return layout.get("rows", 1), layout.get("cols", 1), 0

        # Handle subplot mode toggle (Time vs X-Y)
        @self.app.callback(
            [
                Output("store-subplot-modes", "data"),
                Output("xy-controls", "style"),
                Output("xy-x-signal", "children"),
                Output("xy-y-signal", "children"),
                Output("xy-x-select", "options"),
                Output("xy-y-select", "options"),
                Output("xy-x-select", "value"),
                Output("xy-y-select", "value"),
            ],
            [
                Input("subplot-mode-toggle", "value"),
                Input("store-selected-subplot", "data"),
                Input("tabs", "value"),
                Input("store-assignments", "data"),  # Listen to assignment changes
                Input("store-csv-files", "data"),  # Listen to CSV changes
            ],
            [
                State("store-subplot-modes", "data"),
                State("store-derived", "data"),
                State("store-x-axis-signal", "data"),  # Read X-axis (not Input to avoid cycle)
            ],
            prevent_initial_call=True,
        )
        def handle_subplot_mode(
            mode,
            sel_subplot,
            active_tab,
            assignments,
            csv_files,
            modes,
            derived,
            x_axis_signals,
        ):
            ctx = callback_context
            modes = modes or {}
            assignments = assignments or {}

            tab_idx = (
                int(active_tab.split("-")[1]) if active_tab and "-" in active_tab else 0
            )
            tab_key = str(tab_idx)
            subplot_key = str(sel_subplot or 0)

            if tab_key not in modes:
                modes[tab_key] = {}

            trigger = ctx.triggered[0]["prop_id"] if ctx.triggered else ""

            # If mode toggle triggered, update the mode
            if "subplot-mode-toggle" in trigger:
                modes[tab_key][subplot_key] = mode
                # FIX: Invalidate cache when mode changes to ensure plot updates
                tab_key_int = int(tab_key)
                if tab_key_int in self._state_hash_cache:
                    del self._state_hash_cache[tab_key_int]
                # Clear figure cache for this tab
                cache_keys_to_remove = [k for k in self._figure_cache.keys() if k[0] == tab_key_int]
                for key in cache_keys_to_remove:
                    del self._figure_cache[key]
                logger.info(f"[CACHE] Invalidated cache for tab {tab_key} due to mode change to {mode}")
            else:
                # Otherwise, get current mode for this subplot
                mode = modes.get(tab_key, {}).get(subplot_key, "time")

            # Show/hide X-Y controls based on mode
            xy_style = {"display": "flex"} if mode == "xy" else {"display": "none"}

            # Build X-axis options: Time (default) + ONLY assigned signals for this subplot
            x_axis_options = [{"label": "⏱ Time (default)", "value": "time"}]

            # Get assigned signals for this subplot
            assigned_signals = assignments.get(tab_key, {}).get(subplot_key, [])
            if isinstance(assigned_signals, list):
                for sig in assigned_signals:
                    csv_idx = sig.get("csv_idx", -1)
                    sig_name = sig.get("signal", "")

                    if csv_idx == -1:
                        # Derived signal
                        x_axis_options.append(
                            {"label": f"{sig_name} (D)", "value": f"-1:{sig_name}"}
                        )
                    elif csv_idx >= 0 and csv_idx < len(csv_files or []):
                        # CSV signal - use folder prefix for duplicate names
                        csv_display = get_csv_display_name(csv_files[csv_idx], csv_files)
                        csv_name = os.path.splitext(csv_display)[0]
                        x_axis_options.append(
                            {
                                "label": f"{sig_name} ({csv_name})",
                                "value": f"{csv_idx}:{sig_name}",
                            }
                        )

            # Get current X-axis value from store
            x_axis_signals = x_axis_signals or {}
            x_value = x_axis_signals.get(tab_key, {}).get(subplot_key, "time")

            # Get display text for current X-axis
            x_signal_display = "Time (default)"
            if x_value and x_value != "time":
                # Find the label for this value
                for opt in x_axis_options:
                    if opt["value"] == x_value:
                        x_signal_display = opt["label"]
                        break

            y_signal = ""  # Not used in simplified mode
            y_value = None  # Not used

            return (
                modes,
                xy_style,
                x_signal_display,
                y_signal,
                x_axis_options,
                [],
                x_value,
                y_value,
            )

        # Sync mode toggle and X-axis when subplot changes
        @self.app.callback(
            [
                Output("subplot-mode-toggle", "value"),
                Output("xy-x-select", "value", allow_duplicate=True),
            ],
            [
                Input("store-selected-subplot", "data"),
                Input("tabs", "value"),
            ],
            [
                State("store-subplot-modes", "data"),
                State("store-x-axis-signal", "data"),
            ],
            prevent_initial_call=True,
        )
        def sync_mode_on_subplot_change(sel_subplot, active_tab, modes, x_axis_signals):
            modes = modes or {}
            x_axis_signals = x_axis_signals or {}
            # CRITICAL: Ensure x_axis_signals structure is valid
            if not isinstance(x_axis_signals, dict):
                x_axis_signals = {}
            
            tab_idx = (
                int(active_tab.split("-")[1]) if active_tab and "-" in active_tab else 0
            )
            tab_key = str(tab_idx)
            subplot_key = str(sel_subplot or 0)

            mode = modes.get(tab_key, {}).get(subplot_key, "time")
            # CRITICAL: Defensive check for x_axis_signals structure
            tab_x_axis = x_axis_signals.get(tab_key, {})
            if not isinstance(tab_x_axis, dict):
                tab_x_axis = {}
                x_axis_signals[tab_key] = tab_x_axis
            x_value = tab_x_axis.get(subplot_key, "time")

            return mode, x_value

        # Handle X-axis selection for X-Y mode (store in separate store, not assignments)
        @self.app.callback(
            Output("store-x-axis-signal", "data"),
            Input("xy-x-select", "value"),
            [
                State("store-x-axis-signal", "data"),
                State("store-selected-tab", "data"),
                State("store-selected-subplot", "data"),
            ],
            prevent_initial_call=True,
        )
        def handle_x_axis_selection(x_value, x_axis_signals, sel_tab, sel_subplot):
            if not x_value:
                return dash.no_update

            x_axis_signals = x_axis_signals or {}
            tab_key = str(sel_tab or 0)
            subplot_key = str(sel_subplot or 0)

            if tab_key not in x_axis_signals:
                x_axis_signals[tab_key] = {}

            old_value = x_axis_signals[tab_key].get(subplot_key)
            
            # Only update if value actually changed
            if old_value == x_value:
                return dash.no_update
            
            x_axis_signals[tab_key][subplot_key] = x_value

            # FIX: Invalidate ALL caches when X axis signal changes in X-Y mode
            # This ensures the plot updates when X axis is selected
            print(f"[X-Y MODE] X-axis changed from {old_value} to {x_value} for tab {tab_key}, subplot {subplot_key}")
            
            tab_key_int = int(tab_key)
            
            # Clear state hash cache for this tab
            if tab_key_int in self._state_hash_cache:
                del self._state_hash_cache[tab_key_int]
                print(f"   [CACHE] Cleared state hash cache for tab {tab_key}")
            
            # Clear figure cache for this tab
            cache_keys_to_remove = [k for k in list(self._figure_cache.keys()) if k[0] == tab_key_int]
            for key in cache_keys_to_remove:
                del self._figure_cache[key]
                print(f"   [CACHE] Cleared figure cache entry: {key}")
            
            # Force invalidate all figure-related caches
            self._last_figure_hash = None
            self._last_figure = None
            
            logger.info(f"[CACHE] Invalidated ALL caches for tab {tab_key} due to X-axis signal change to {x_value}")

            return x_axis_signals

        # Note: X-Y mode now uses list-based assignments like time mode
        # The xy-remove callback is no longer needed

        # Tab management
        @self.app.callback(
            [
                Output("tabs", "children"),
                Output("store-num-tabs", "data"),
                Output("store-layouts", "data", allow_duplicate=True),
                Output("tabs", "value", allow_duplicate=True),
                Output("store-assignments", "data", allow_duplicate=True),
            ],
            [Input("btn-add-tab", "n_clicks"), Input("btn-del-tab", "n_clicks")],
            [
                State("tabs", "children"),
                State("tabs", "value"),
                State("store-layouts", "data"),
                State("store-assignments", "data"),
            ],
            prevent_initial_call=True,
        )
        def manage_tabs(add, delete, tabs, current, layouts, assignments):
            ctx = callback_context
            trigger = ctx.triggered[0]["prop_id"].split(".")[0] if ctx.triggered else ""
            tabs, layouts = tabs or [], layouts or {}
            assignments = assignments or {}
            new_active = current

            if trigger == "btn-add-tab":
                idx = len(tabs)
                tabs.append(dcc.Tab(label=f"Tab {idx+1}", value=f"tab-{idx}"))
                layouts[str(idx)] = {"rows": 1, "cols": 1}

            elif trigger == "btn-del-tab" and len(tabs) > 1:
                # Get current tab index
                del_idx = (
                    int(current.split("-")[1]) if current and "-" in current else 0
                )

                # Remove the tab
                tabs = [t for i, t in enumerate(tabs) if i != del_idx]

                # CRITICAL: Rebuild layouts and assignments with reindexed keys
                # This ensures remaining tabs keep their properties but get new indices
                new_layouts = {}
                new_assignments = {}

                new_i = 0
                for old_i in range(len(tabs) + 1):  # +1 because we removed one
                    if old_i == del_idx:
                        continue  # Skip deleted tab
                    
                    old_key = str(old_i)
                    new_key = str(new_i)
                    
                    # Copy layout (deep copy)
                    if old_key in layouts:
                        new_layouts[new_key] = layouts[old_key].copy()
                    
                    # Copy assignments (deep copy to avoid reference issues)
                    if old_key in assignments:
                        new_assignments[new_key] = {}
                        for sp_key, sp_data in assignments[old_key].items():
                            if isinstance(sp_data, list):
                                new_assignments[new_key][sp_key] = [s.copy() for s in sp_data]
                            else:
                                new_assignments[new_key][sp_key] = sp_data
                    
                    new_i += 1

                layouts = new_layouts
                assignments = new_assignments

                # Rename tabs to maintain sequential numbering
                for i, t in enumerate(tabs):
                    if isinstance(t, dict):
                        t["props"]["label"] = f"Tab {i+1}"
                        t["props"]["value"] = f"tab-{i}"
                    elif hasattr(t, "label"):
                        t.label = f"Tab {i+1}"
                        t.value = f"tab-{i}"
                
                # Set new active tab (first tab if deleted was first, otherwise previous)
                if del_idx == 0:
                    new_active = "tab-0"
                elif del_idx < len(tabs):
                    new_active = f"tab-{del_idx-1}"
                else:
                    new_active = f"tab-{len(tabs)-1}"

            return tabs, len(tabs), layouts, new_active, assignments

        # Link modal
        @self.app.callback(
            [Output("modal-link", "is_open"), Output("link-checks", "options")],
            [
                Input("btn-link", "n_clicks"),
                Input("btn-close-link", "n_clicks"),
                Input("btn-create-link", "n_clicks"),
            ],
            [State("modal-link", "is_open")],
            prevent_initial_call=True,
        )
        def toggle_link(o, c, cr, is_open):
            if "btn-link" in callback_context.triggered[0]["prop_id"]:
                return True, [
                    {"label": os.path.basename(f), "value": i}
                    for i, f in enumerate(self.data_manager.csv_file_paths)
                ]
            return False, []

        @self.app.callback(
            [
                Output("store-links", "data", allow_duplicate=True),
                Output("status-text", "children", allow_duplicate=True),
            ],
            Input("btn-create-link", "n_clicks"),
            [
                State("link-checks", "value"),
                State("link-name", "value"),
                State("store-links", "data"),
            ],
            prevent_initial_call=True,
        )
        def create_link(n, selected, name, links):
            if not n or not selected or len(selected) < 2:
                return dash.no_update, "Select 2+ CSVs"
            links = links or []
            links.append(
                {"csv_indices": selected, "name": name or f"Link{len(links)+1}"}
            )
            return links, f"[OK] Linked"

        # Properties modal
        @self.app.callback(
            [
                Output("modal-props", "is_open"),
                Output("props-title", "children"),
                Output("store-context-signal", "data"),
                Output("prop-name", "value"),
                Output("prop-scale", "value"),
                Output("prop-color", "value"),
                Output("prop-width", "value"),
                Output("prop-time-offset", "value"),
                Output("prop-state-signal", "value"),
            ],
            [
                Input({"type": "sig-props", "csv_idx": ALL, "signal": ALL}, "n_clicks"),  # FIXED
                Input("btn-close-props", "n_clicks"),
            ],
            [
                State("store-signal-props", "data"),
                State("store-time-offsets", "data"),
            ],
            prevent_initial_call=True,
        )
        def toggle_props(clicks, close, props, time_offsets):
            trigger = (
                callback_context.triggered[0]["prop_id"]
                if callback_context.triggered
                else ""
            )
            if "sig-props" in trigger:
                for c in clicks or []:
                    if c:
                        import json as js

                        id_dict = js.loads(trigger.split(".")[0])
                        key = f"{id_dict.get('csv_idx', -1)}:{id_dict.get('signal', '')}"  # FIXED
                        p = (props or {}).get(key, {})
                        # Get time offset (check signal-specific, then CSV-wide)
                        time_offsets = time_offsets or {}
                        offset = time_offsets.get(key, time_offsets.get(str(id_dict.get('csv_idx', -1)), 0))  # FIXED
                        return (
                            True,
                            f"Props: {id_dict.get('signal', '')}",  # FIXED
                            key,
                            p.get("display_name", id_dict.get("signal", "")),  # FIXED
                            p.get("scale", 1.0),
                            p.get("color", "#2E86AB"),
                            p.get("width", 1.5),
                            offset,
                            p.get("is_state", False),
                        )
            return False, "", None, "", 1.0, "#2E86AB", 1.5, 0, False

        @self.app.callback(
            [
                Output("store-signal-props", "data", allow_duplicate=True),
                Output("store-time-offsets", "data", allow_duplicate=True),
                Output("status-text", "children", allow_duplicate=True),
            ],
            Input("btn-apply-props", "n_clicks"),
            [
                State("store-context-signal", "data"),
                State("prop-name", "value"),
                State("prop-scale", "value"),
                State("prop-color", "value"),
                State("prop-width", "value"),
                State("prop-time-offset", "value"),
                State("prop-apply-tree", "value"),
                State("prop-state-signal", "value"),
                State("store-signal-props", "data"),
                State("store-time-offsets", "data"),
            ],
            prevent_initial_call=True,
        )
        def apply_props(n, key, name, scale, color, width, time_offset, show, is_state, props, time_offsets):
            if not n or not key:
                return dash.no_update, dash.no_update, dash.no_update
            props = props or {}
            time_offsets = time_offsets or {}
            
            # Save time offset for this signal
            offset_val = float(time_offset) if time_offset else 0.0
            if offset_val != 0:
                time_offsets[key] = offset_val
            elif key in time_offsets:
                del time_offsets[key]
            
            old_display_name = props.get(key, {}).get("display_name", "")
            props[key] = {
                "display_name": name,
                "scale": float(scale or 1),
                "color": color,
                "width": float(width or 1.5),
                "show_in_tree": show,
                "is_state": is_state,
            }
            
            # Update derived signals that reference this signal
            # When a signal is renamed, derived signals should reflect the new name
            csv_idx, signal_name = key.split(":", 1)
            csv_idx = int(csv_idx)
            
            # Find derived signals that use this source signal
            for derived_name, derived_info in self.derived_signals.items():
                source_signal_key = derived_info.get("source_signal_key")
                source_signals = derived_info.get("source_signals", [])
                
                # Check if this derived signal references the renamed signal
                should_update = False
                if source_signal_key == key:
                    should_update = True
                elif source_signals and key in source_signals:
                    should_update = True
                
                if should_update:
                    # Update the derived signal's display to reflect renamed source
                    # The derived signal keeps its own name, but we track the source rename
                    if key not in props:
                        props[key] = {}
                    # Store the original signal name for reference
                    if "original_signal_name" not in props[key]:
                        props[key]["original_signal_name"] = signal_name
            
            self.signal_properties = props
            return props, time_offsets, "[OK] Saved"

        # Single signal operations
        @self.app.callback(
            [
                Output("modal-ops", "is_open"),
                Output("ops-title", "children"),
                Output("op-result-name", "value"),
            ],
            [
                Input({"type": "sig-ops", "csv_idx": ALL, "signal": ALL}, "n_clicks"),  # FIXED
                Input("btn-close-ops", "n_clicks"),
            ],
            prevent_initial_call=True,
        )
        def toggle_ops(clicks, close):
            trigger = (
                callback_context.triggered[0]["prop_id"]
                if callback_context.triggered
                else ""
            )
            if "sig-ops" in trigger:
                for c in clicks or []:
                    if c:
                        import json as js

                        id_dict = js.loads(trigger.split(".")[0])
                        # Return empty name - let compute_op generate meaningful default based on operation
                        return True, f"Op: {id_dict.get('signal', '')}", ""  # FIXED
            return False, "", ""

        @self.app.callback(
            [
                Output("store-derived", "data", allow_duplicate=True),
                Output("status-text", "children", allow_duplicate=True),
            ],
            Input("btn-compute-op", "n_clicks"),
            [
                State("ops-title", "children"),
                State("op-type", "value"),
                State("op-result-name", "value"),
                State("store-derived", "data"),
            ],
            prevent_initial_call=True,
        )
        def compute_op(n, title, op, name, derived):
            if not n or not title:
                return dash.no_update, dash.no_update
            sig = title.replace("Op: ", "")
            derived = derived or {}
            time_data, sig_data = None, None

            # Find the signal and apply properties
            source_csv_idx = None
            source_signal_key = None
            
            # Check if it's a derived signal - check both self.derived_signals and derived parameter
            ds = None
            if derived and sig in derived:
                ds = derived.get(sig)
            elif self.derived_signals and sig in self.derived_signals:
                ds = self.derived_signals.get(sig)
            
            if ds:
                # It's a derived signal
                time_data = ds.get("time", [])
                sig_data = ds.get("data", [])
                # FIX: Convert lists to numpy arrays for operations to work on derived signals
                if isinstance(time_data, list):
                    time_data = np.array(time_data)
                if isinstance(sig_data, list):
                    sig_data = np.array(sig_data)
                source_signal_key = f"-1:{sig}"  # Derived signal
                source_csv_idx = -1  # Derived signal
            else:
                # Regular signal from CSV - find it and apply properties
                for csv_idx, df in enumerate(self.data_manager.data_tables):
                    if df is not None and sig in df.columns:
                        source_csv_idx = csv_idx
                        source_signal_key = f"{csv_idx}:{sig}"
                        
                        time_col = "Time" if "Time" in df.columns else df.columns[0]
                        time_data = df[time_col].values
                        sig_data = df[sig].values
                        
                        # Apply signal properties (scale, time_offset)
                        props = self.signal_properties.get(source_signal_key, {})
                        scale = props.get("scale", 1.0)
                        time_offset = props.get("time_offset", 0.0)
                        
                        if scale != 1.0:
                            sig_data = sig_data * scale
                        if time_offset != 0.0:
                            time_data = time_data + time_offset
                        
                        break

            if time_data is None or len(time_data) == 0:
                return dash.no_update, "Signal not found"

            try:
                if op == "derivative":
                    result = np.gradient(sig_data, time_data)
                elif op == "integral":
                    result = np.cumsum(sig_data) * np.mean(np.diff(time_data))
                elif op == "abs":
                    result = np.abs(sig_data)
                elif op == "sqrt":
                    result = np.sqrt(np.abs(sig_data))
                elif op == "negate":
                    result = -sig_data
                else:
                    result = sig_data
                
                # Generate meaningful default name if not provided
                if not name or name.strip() == "":
                    # Use LaTeX-style names for better readability - just the symbol, not with signal name
                    op_names = {
                        "derivative": "\\delta",
                        "integral": "\\int",
                        "abs": "|x|",
                        "sqrt": "\\sqrt{x}",
                        "negate": "-x",
                    }
                    name = op_names.get(op, f"{op}_{sig}")
                
                derived[name] = {
                    "time": time_data.tolist(),
                    "data": result.tolist(),
                    "source": sig,
                    "source_signal_key": source_signal_key,  # Track source for cleanup
                    "source_csv_idx": source_csv_idx,  # Track CSV for cleanup
                    "op": op,
                }
                self.derived_signals = derived
                return derived, f"[OK] {name}"
            except Exception as e:
                return dash.no_update, f"[ERROR] {e}"

        # Multi-signal operations
        @self.app.callback(
            [
                Output("modal-multi-ops", "is_open"),
                Output("selected-signals-info", "children"),
            ],
            [
                Input("btn-operate-selected", "n_clicks"),
                Input("btn-close-multi", "n_clicks"),
                Input("btn-compute-multi", "n_clicks"),
            ],
            [State("store-highlighted", "data")],
            prevent_initial_call=True,
        )
        def toggle_multi(o, c, comp, highlighted):
            if "btn-operate-selected" in callback_context.triggered[0]["prop_id"]:
                info = (
                    f"Selected: {', '.join(h.split(':')[1] for h in (highlighted or []))}"
                    if highlighted
                    else "None selected"
                )
                return True, info
            return False, ""

        @self.app.callback(
            [
                Output("store-derived", "data", allow_duplicate=True),
                Output("status-text", "children", allow_duplicate=True),
            ],
            Input("btn-compute-multi", "n_clicks"),
            [
                State("store-highlighted", "data"),
                State("multi-op-type", "value"),
                State("multi-op-name", "value"),
                State("store-derived", "data"),
            ],
            prevent_initial_call=True,
        )
        def compute_multi(n, highlighted, op, name, derived):
            if not n or not highlighted or len(highlighted) < 2:
                return dash.no_update, "Select 2+ signals"
            derived = derived or {}
            signals_data, time_data = [], None

            for h in highlighted:
                csv_idx, sig = h.split(":", 1)
                csv_idx = int(csv_idx)
                signal_key = f"{csv_idx}:{sig}"

                # Get signal data and apply properties
                if csv_idx == -1 and sig in self.derived_signals:
                    ds = self.derived_signals[sig]
                    sig_data = np.array(ds.get("data", []))
                    if time_data is None:
                        time_data = np.array(ds.get("time", []))
                elif csv_idx >= 0 and csv_idx < len(self.data_manager.data_tables):
                    df = self.data_manager.data_tables[csv_idx]
                    if df is not None and sig in df.columns:
                        if time_data is None:
                            time_col = "Time" if "Time" in df.columns else df.columns[0]
                            time_data = df[time_col].values
                        sig_data = df[sig].values
                        
                        # Apply signal properties (scale, time_offset)
                        props = self.signal_properties.get(signal_key, {})
                        scale = props.get("scale", 1.0)
                        time_offset = props.get("time_offset", 0.0)
                        
                        if scale != 1.0:
                            sig_data = sig_data * scale
                        if time_offset != 0.0:
                            # Apply offset to time for this signal only if it's the first one
                            # For subsequent signals, we'll align by time
                            if time_data is None or len(time_data) == 0:
                                time_data = df[time_col].values + time_offset
                    else:
                        continue
                else:
                    continue

                signals_data.append(sig_data)

            if len(signals_data) < 2:
                return dash.no_update, "Not enough data"

            try:
                min_len = min(len(s) for s in signals_data)
                signals_data = [s[:min_len] for s in signals_data]
                time_data = time_data[:min_len]
                a, b = np.array(signals_data[0]), np.array(signals_data[1])
                if op == "add":
                    result = a + b
                elif op == "sub":
                    result = a - b
                elif op == "mul":
                    result = a * b
                elif op == "div":
                    result = np.divide(a, b, where=b != 0)
                elif op == "norm":
                    # PERFORMANCE: Vectorized norm calculation
                    signals_array = np.array(signals_data)  # Shape: (n_signals, n_points)
                    result = np.sqrt(np.sum(signals_array ** 2, axis=0))  # Vectorized across all signals
                elif op == "mean":
                    result = np.mean(signals_data, axis=0)
                else:
                    result = a
                
                # Generate meaningful default name if not provided
                if not name or name.strip() == "":
                    # Build name from signal names: "signal1+signal2" or "signal1-signal2" etc.
                    op_symbols = {
                        "add": "+",
                        "sub": "-",
                        "mul": "×",
                        "div": "÷",
                        "norm": "||",
                        "mean": "mean",
                    }
                    op_symbol = op_symbols.get(op, op)
                    
                    # Get signal names from highlighted list
                    signal_names = []
                    for h in highlighted:
                        csv_idx, sig = h.split(":", 1)
                        csv_idx = int(csv_idx)
                        if csv_idx == -1:
                            signal_names.append(sig)
                        else:
                            # Get short signal name
                            signal_names.append(sig)
                    
                    if op == "norm" or op == "mean":
                        # For norm/mean, use all signals: "||signal1,signal2||"
                        name = f"{op_symbol}{','.join(signal_names[:3])}{op_symbol}"  # Limit to 3 for readability
                    else:
                        # For binary ops: "signal1+signal2"
                        if len(signal_names) >= 2:
                            name = f"{signal_names[0]}{op_symbol}{signal_names[1]}"
                        else:
                            name = f"{op}_{signal_names[0] if signal_names else 'result'}"
                
                derived[name] = {
                    "time": time_data.tolist(),
                    "data": result.tolist(),
                    "source": "multi",
                    "op": op,
                    "source_signals": highlighted.copy(),  # Track source signals for issue #4
                }
                self.derived_signals = derived
                return derived, f"[OK] {name}"
            except Exception as e:
                return dash.no_update, f"[ERROR] {e}"

        # Save session - downloads to user's chosen location
        @self.app.callback(
            [
                Output("download-session", "data"),
                Output("status-text", "children", allow_duplicate=True),
            ],
            Input("btn-save", "n_clicks"),
            [
                State("store-csv-files", "data"),
                State("store-assignments", "data"),
                State("store-layouts", "data"),
                State("store-links", "data"),
                State("store-signal-props", "data"),
                State("store-derived", "data"),
                State("store-num-tabs", "data"),
                State("store-subplot-modes", "data"),
                State("store-cursor-x", "data"),
                State("store-theme", "data"),
                State("store-search-filters", "data"),
                State("store-time-columns", "data"),
                State("store-x-axis-signal", "data"),
                State("store-document-text", "data"),
                State("store-subplot-metadata", "data"),
                State("store-time-offsets", "data"),
                State("store-selected-tab", "data"),
                State("store-selected-subplot", "data"),
                State("tabs", "value"),
                State("subplot-select", "value"),
                State("rows-input", "value"),
                State("cols-input", "value"),
            ],
            prevent_initial_call=True,
        )
        def save_session(
            n,
            files,
            assign,
            layouts,
            links,
            props,
            derived,
            num_tabs,
            subplot_modes,
            cursor_x,
            theme,
            search_filters,
            time_columns,
            x_axis_signals,
            doc_text,
            subplot_metadata,
            time_offsets,
            selected_tab,
            selected_subplot,
            active_tab,
            active_subplot,
            rows,
            cols,
        ):
            if not n:
                return dash.no_update, dash.no_update
            try:
                # Get current tab/subplot from active values or stores
                current_tab = selected_tab if selected_tab is not None else 0
                if active_tab and "-" in active_tab:
                    current_tab = int(active_tab.split("-")[1])
                
                current_subplot = selected_subplot if selected_subplot is not None else 0
                if active_subplot is not None:
                    current_subplot = int(active_subplot)
                
                session_data = {
                    "version": "3.1",  # Version 3.1: includes full UI state
                    "files": files,  # Original file paths directly
                    "original_file_paths": self.original_file_paths,
                    "assignments": assign,
                    "layouts": layouts,
                    "links": links,
                    "props": props,
                    "derived": derived,
                    "num_tabs": num_tabs or 1,
                    "subplot_modes": subplot_modes or {},
                    "cursor_x": cursor_x,
                    "theme": theme or "dark",
                    "search_filters": search_filters or [],
                    "time_columns": time_columns or {},
                    "x_axis_signals": x_axis_signals or {},
                    "document_text": doc_text or {"introduction": "", "conclusion": ""},
                    "subplot_metadata": subplot_metadata or {},
                    "time_offsets": time_offsets or {},
                    "selected_tab": current_tab,  # Save current tab
                    "selected_subplot": current_subplot,  # Save current subplot
                    "display_options": {},  # Placeholder for future
                }
                # Generate filename with timestamp
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"signal_viewer_session_{timestamp}.json"
                return (
                    dict(content=json.dumps(session_data, indent=2), filename=filename),
                    "[OK] Session saved",
                )
            except Exception as e:
                return dash.no_update, f"[ERROR] {e}"

        # Load session - from uploaded file
        @self.app.callback(
            [
                Output("store-csv-files", "data", allow_duplicate=True),
                Output("csv-list", "children", allow_duplicate=True),  # Update CSV list UI
                Output("store-assignments", "data", allow_duplicate=True),
                Output("store-layouts", "data", allow_duplicate=True),
                Output("store-links", "data", allow_duplicate=True),
                Output("store-signal-props", "data", allow_duplicate=True),
                Output("store-derived", "data", allow_duplicate=True),
                Output("store-num-tabs", "data", allow_duplicate=True),
                Output("tabs", "children", allow_duplicate=True),
                Output("tabs", "value", allow_duplicate=True),  # Restore active tab
                Output("store-selected-tab", "data", allow_duplicate=True),  # Restore selected tab
                Output("store-selected-subplot", "data", allow_duplicate=True),  # Restore selected subplot
                Output("subplot-select", "value", allow_duplicate=True),  # Restore subplot selector
                Output("store-subplot-modes", "data", allow_duplicate=True),
                Output("store-cursor-x", "data", allow_duplicate=True),
                Output("theme-switch", "value"),
                Output("store-search-filters", "data", allow_duplicate=True),
                Output("store-time-columns", "data", allow_duplicate=True),
                Output("store-x-axis-signal", "data", allow_duplicate=True),
                Output("store-document-text", "data", allow_duplicate=True),
                Output("store-subplot-metadata", "data", allow_duplicate=True),
                Output("store-time-offsets", "data", allow_duplicate=True),
                Output("status-text", "children", allow_duplicate=True),
            ],
            Input("upload-session", "contents"),
            State("upload-session", "filename"),
            prevent_initial_call=True,
        )
        def load_session(contents, filename):
            if not contents:
                return [dash.no_update] * 23
            try:
                # Decode uploaded file
                content_type, content_string = contents.split(",")
                decoded = base64.b64decode(content_string).decode("utf-8")
                d = json.loads(decoded)

                files = d.get("files", [])
                
                # CRITICAL: Clear all state before loading to prevent conflicts
                self.data_manager.csv_file_paths = []
                self.data_manager.data_tables = []
                self.data_manager.last_read_rows = []
                self.data_manager.last_file_mod_times = []
                self.derived_signals = {}
                self.signal_properties = {}
                self.original_file_paths = {}
                self.invalidate_caches()
                
                # Now load the session data
                self.data_manager.csv_file_paths = files
                self.data_manager.data_tables = [None] * len(files)
                self.data_manager.last_read_rows = [0] * len(files)
                self.data_manager.last_file_mod_times = [0] * len(files)
                
                # Restore original file paths for streaming/refresh
                self.original_file_paths = d.get("original_file_paths", {})
                
                # Restore to data_manager as well
                for uploads_path, original_path in self.original_file_paths.items():
                    if uploads_path in files:
                        idx = files.index(uploads_path)
                        self.data_manager.original_source_paths[idx] = original_path
                
                # Load CSV data
                for i in range(len(files)):
                    if os.path.exists(files[i]):
                        try:
                            self.data_manager.read_initial_data(i)
                        except Exception as e:
                            print(f"[WARN] Could not load CSV {i}: {e}")
                
                # Restore derived signals and properties
                self.derived_signals = d.get("derived", {}) or {}
                self.signal_properties = d.get("props", {}) or {}

                # Restore tabs
                num_tabs = d.get("num_tabs", 1)
                tabs_children = [
                    dcc.Tab(label=f"Tab {i+1}", value=f"tab-{i}")
                    for i in range(num_tabs)
                ]
                
                # Restore selected tab and subplot
                selected_tab = d.get("selected_tab", 0)
                selected_subplot = d.get("selected_subplot", 0)
                
                # Ensure selected tab is valid
                if selected_tab >= num_tabs:
                    selected_tab = 0
                
                # Set active tab value
                active_tab_value = f"tab-{selected_tab}"
                
                # CRITICAL: Ensure all data structures have valid tab keys
                layouts = d.get("layouts", {}) or {}
                assignments = d.get("assignments", {}) or {}
                subplot_modes_data = d.get("subplot_modes", {}) or {}
                x_axis_signals_data = d.get("x_axis_signals", {}) or {}
                
                # CRITICAL: Validate x_axis_signals structure - ensure it's a dict of dicts
                if not isinstance(x_axis_signals_data, dict):
                    x_axis_signals_data = {}
                
                # Validate and fix tab keys in all structures
                for i in range(num_tabs):
                    tab_key = str(i)
                    if tab_key not in layouts:
                        layouts[tab_key] = {"rows": 1, "cols": 1}
                    if tab_key not in assignments:
                        assignments[tab_key] = {}
                    if tab_key not in subplot_modes_data:
                        subplot_modes_data[tab_key] = {}
                    if tab_key not in x_axis_signals_data:
                        x_axis_signals_data[tab_key] = {}
                    else:
                        # CRITICAL: Fix corrupted data - ensure tab_key value is a dict
                        if not isinstance(x_axis_signals_data[tab_key], dict):
                            x_axis_signals_data[tab_key] = {}
                
                # Theme (True = dark, False = light)
                theme = d.get("theme", "dark")
                is_dark = theme == "dark"
                
                # Update stores with selected tab/subplot after a short delay
                # This ensures tabs are created first
                import threading
                # Build CSV list UI with delete buttons
                csv_items = [
                    html.Div(
                        [
                            html.Span(
                                get_csv_display_name(f, files),
                                style={"fontSize": "10px"},
                                title=f,  # Show full path on hover
                            ),
                            html.A(
                                "×",
                                id={"type": "del-csv", "idx": i},
                                className="float-end text-danger",
                                style={"cursor": "pointer"},
                            ),
                        ],
                        className="py-1",
                    )
                    for i, f in enumerate(files)
                ]

                return (
                    files,
                    csv_items,  # CSV list UI with delete buttons
                    assignments,  # Use validated assignments
                    layouts,  # Use validated layouts
                    d.get("links", []) or [],
                    d.get("props", {}) or {},
                    d.get("derived", {}) or {},
                    num_tabs,
                    tabs_children,
                    active_tab_value,  # Set active tab
                    selected_tab,  # Store selected tab
                    selected_subplot,  # Store selected subplot
                    selected_subplot,  # Set subplot selector
                    subplot_modes_data,  # Use validated subplot_modes
                    d.get("cursor_x", {"x": None, "initialized": False}) or {"x": None, "initialized": False},
                    is_dark,
                    d.get("search_filters", []) or [],
                    d.get("time_columns", {}) or {},
                    x_axis_signals_data,  # Use validated x_axis_signals
                    d.get("document_text", {"introduction": "", "conclusion": ""}) or {"introduction": "", "conclusion": ""},
                    d.get("subplot_metadata", {}) or {},
                    d.get("time_offsets", {}) or {},
                    f"[OK] Loaded: {filename}",
                )
            except Exception as e:
                return [dash.no_update] * 22 + [f"[ERROR] {e}"]

        # Clientside callback to initialize Split.js after page loads
        self.app.clientside_callback(
            """
            function(n_intervals) {
                console.log('Split init attempt:', n_intervals);
                
                if (typeof Split === 'undefined') {
                    console.log('Split.js not loaded yet');
                    return window.dash_clientside.no_update;
                }
                
                var sidebar = document.getElementById('split-sidebar');
                var plot = document.getElementById('split-plot');
                var panel1 = document.getElementById('split-panel-1');
                var panel2 = document.getElementById('split-panel-2');
                var panel3 = document.getElementById('split-panel-3');
                
                if (!sidebar || !plot) {
                    console.log('Split elements not found yet');
                    return window.dash_clientside.no_update;
                }
                
                // Check if already initialized
                if (sidebar.dataset.splitInit === 'true') {
                    console.log('Already initialized');
                    return true;
                }
                
                try {
                    // Horizontal split: sidebar vs plot
                    Split(['#split-sidebar', '#split-plot'], {
                        sizes: [35, 65],
                        minSize: [300, 400],
                        gutterSize: 10,
                        cursor: 'col-resize',
                        direction: 'horizontal',
                        onDrag: function() {
                            window.dispatchEvent(new Event('resize'));
                        }
                    });
                    console.log('Horizontal split created');
                    
                    // Vertical split: panels in sidebar
                    if (panel1 && panel2 && panel3) {
                        Split(['#split-panel-1', '#split-panel-2', '#split-panel-3'], {
                            sizes: [18, 57, 25],
                            minSize: [80, 100, 80],
                            gutterSize: 8,
                            cursor: 'row-resize',
                            direction: 'vertical'
                        });
                        console.log('Vertical split created');
                    }
                    
                    sidebar.dataset.splitInit = 'true';
                    console.log('Split.js initialized successfully!');
                    return true;
                } catch (e) {
                    console.error('Split error:', e);
                    return window.dash_clientside.no_update;
                }
            }
            """,
            Output("store-split-init", "data"),
            Input("interval-split-init", "n_intervals"),
        )

        # Update cursor position from plot click (NOT from slider to avoid cycle)
        @self.app.callback(
            Output("store-cursor-x", "data"),
            [Input("plot", "clickData"), Input("time-cursor-check", "value")],
            [State("store-cursor-x", "data")],
            prevent_initial_call=True,
        )
        def update_cursor_position(click_data, cursor_enabled, current_cursor):
            if not cursor_enabled:
                return {"x": None, "initialized": False}

            current_cursor = current_cursor or {"x": None, "initialized": False}

            # Update from plot click only
            if click_data:
                try:
                    point = click_data.get("points", [{}])[0]
                    x_val = point.get("x", None)
                    if x_val is not None:
                        return {"x": float(x_val), "initialized": True}
                except:
                    pass

            return current_cursor
        
        # Update cursor from slider (separate callback to break cycle)
        # Use assignments to compute range, not figure
        # Also updates cursor text display for immediate feedback
        @self.app.callback(
            [
                Output("store-cursor-x", "data", allow_duplicate=True),
                Output("cursor-slider-value", "children", allow_duplicate=True),
            ],
            Input("cursor-slider", "value"),
            [
                State("time-cursor-check", "value"),
                State("store-assignments", "data"),
                State("store-selected-tab", "data"),
                State("store-selected-subplot", "data"),
                State("store-time-columns", "data"),
            ],
            prevent_initial_call=True,
        )
        def update_cursor_from_slider(slider_value, cursor_enabled, assignments, sel_tab, sel_subplot, time_columns):
            if not cursor_enabled or slider_value is None:
                return dash.no_update, dash.no_update
            
            # Compute time range from assignments (same logic as slider update)
            x_min, x_max = 0, 100
            
            if assignments:
                tab_key = str(sel_tab or 0)
                subplot_key = str(sel_subplot or 0)
                sp_assignment = assignments.get(tab_key, {}).get(subplot_key, [])
                
                if isinstance(sp_assignment, list) and sp_assignment:
                    all_times = []
                    for sig in sp_assignment:
                        csv_idx = sig.get("csv_idx", -1)
                        signal_name = sig.get("signal", "")
                        
                        try:
                            if csv_idx == -1 and signal_name in self.derived_signals:
                                ds = self.derived_signals[signal_name]
                                times = ds.get("time", [])
                                if times:
                                    all_times.extend(times)
                            elif csv_idx >= 0 and csv_idx < len(self.data_manager.data_tables):
                                df = self.data_manager.data_tables[csv_idx]
                                if df is not None and signal_name in df.columns:
                                    time_col = time_columns.get(str(csv_idx), "Time") if time_columns else "Time"
                                    if time_col in df.columns:
                                        times = df[time_col].values
                                        if len(times) > 0:
                                            all_times.extend(times.tolist() if hasattr(times, 'tolist') else times)
                        except:
                            pass
                    
                    if all_times:
                        # PERFORMANCE: Use numpy for min/max on large arrays
                        times_array = np.array(all_times)
                        x_min, x_max = float(np.min(times_array)), float(np.max(times_array))
            
            try:
                # Slider value is now actual time value (not 0-100)
                x_val = float(slider_value)
                # Clamp to valid range
                x_val = max(x_min, min(x_max, x_val))
                # Return both cursor data and text display
                cursor_display = f"T: {x_val:.3f}"
                return {"x": x_val, "initialized": True}, cursor_display
            except:
                pass
            
            return dash.no_update, dash.no_update
        
        # Show/hide slider and update its range based on assignments (not figure to avoid cycle)
        # Cursor animation controls (play/stop) - supports both bottom and header buttons
        @self.app.callback(
            [
                Output("interval-cursor-animation", "disabled"),
                Output("btn-cursor-play", "style"),
                Output("btn-cursor-stop", "style"),
            ],
            [
                Input("btn-cursor-play", "n_clicks"),
                Input("btn-cursor-stop", "n_clicks"),
                Input("btn-cursor-play-header", "n_clicks"),
                Input("btn-cursor-stop-header", "n_clicks"),
            ],
            [
                State("interval-cursor-animation", "disabled"),
                State("time-cursor-check", "value"),
            ],
            prevent_initial_call=True,
        )
        def control_cursor_animation(play_clicks, stop_clicks, play_header, stop_header, interval_disabled, cursor_enabled):
            """Control cursor animation playback"""
            if not cursor_enabled:
                return True, {"opacity": 0.5}, {"opacity": 0.5}
            
            ctx = callback_context
            if not ctx.triggered:
                return dash.no_update, dash.no_update, dash.no_update
            
            trigger = ctx.triggered[0]["prop_id"]
            
            if "btn-cursor-play" in trigger:
                # Start animation
                print("[CURSOR] Starting cursor animation")
                return False, {"opacity": 0.5}, {"opacity": 1.0}
            elif "btn-cursor-stop" in trigger:
                # Stop animation
                print("[CURSOR] Stopping cursor animation")
                return True, {"opacity": 1.0}, {"opacity": 0.5}
            
            return dash.no_update, dash.no_update, dash.no_update
        
        # Show/hide header cursor controls based on cursor enabled
        @self.app.callback(
            Output("cursor-controls-header", "style"),
            Input("time-cursor-check", "value"),
            prevent_initial_call=True,
        )
        def toggle_header_cursor_controls(cursor_enabled):
            if cursor_enabled:
                return {"display": "flex", "minWidth": "150px"}
            return {"display": "none"}
        
        # Sync header slider to main slider (value, min, max)
        @self.app.callback(
            [
                Output("cursor-slider-header", "value"),
                Output("cursor-slider-header", "min"),
                Output("cursor-slider-header", "max"),
            ],
            [
                Input("cursor-slider", "value"),
                Input("cursor-slider", "min"),
                Input("cursor-slider", "max"),
            ],
            prevent_initial_call=True,
        )
        def sync_header_slider(value, min_val, max_val):
            return value, min_val, max_val
        
        # Update cursor from header slider
        @self.app.callback(
            Output("cursor-slider", "value", allow_duplicate=True),
            Input("cursor-slider-header", "value"),
            prevent_initial_call=True,
        )
        def update_from_header_slider(value):
            return value
        
        # Animate cursor position when play is active
        @self.app.callback(
            Output("cursor-slider", "value", allow_duplicate=True),
            Input("interval-cursor-animation", "n_intervals"),
            [
                State("cursor-slider", "value"),
                State("cursor-slider", "min"),
                State("cursor-slider", "max"),
                State("interval-cursor-animation", "disabled"),
            ],
            prevent_initial_call=True,
        )
        def animate_cursor(n_intervals, current_value, min_val, max_val, is_disabled):
            """Animate cursor by incrementing slider value"""
            if is_disabled or current_value is None:
                return dash.no_update
            
            # Increment by 0.5% per interval (100ms), so full animation takes ~20 seconds
            step = (max_val - min_val) * 0.005
            new_value = current_value + step
            
            # Loop back to start when reaching end
            if new_value > max_val:
                new_value = min_val
            
            return new_value

        @self.app.callback(
            [
                Output("cursor-slider-container", "style"),
                Output("cursor-slider", "disabled"),
                Output("cursor-slider", "min"),
                Output("cursor-slider", "max"),
                Output("cursor-slider", "value"),
                Output("cursor-slider-value", "children"),
            ],
            [
                Input("time-cursor-check", "value"),
                Input("store-cursor-x", "data"),
                Input("store-assignments", "data"),  # Use assignments to compute range, not figure
            ],
            [
                State("store-selected-tab", "data"),
                State("store-selected-subplot", "data"),
                State("store-time-columns", "data"),
            ],
            prevent_initial_call=True,
        )
        def update_cursor_slider(cursor_enabled, cursor_data, assignments, sel_tab, sel_subplot, time_columns):
            if not cursor_enabled:
                return {"display": "none"}, True, 0, 100, 50, "--"
            
            # Compute time range from actual data, not figure (avoids circular dependency)
            x_min, x_max = 0, 100
            cursor_x = None
            slider_value = 50
            
            # Get time range from assigned signals
            if assignments:
                tab_key = str(sel_tab or 0)
                subplot_key = str(sel_subplot or 0)
                sp_assignment = assignments.get(tab_key, {}).get(subplot_key, [])
                
                if isinstance(sp_assignment, list) and sp_assignment:
                    all_times = []
                    for sig in sp_assignment:
                        csv_idx = sig.get("csv_idx", -1)
                        signal_name = sig.get("signal", "")
                        
                        try:
                            if csv_idx == -1 and signal_name in self.derived_signals:
                                ds = self.derived_signals[signal_name]
                                times = ds.get("time", [])
                                if times:
                                    all_times.extend(times)
                            elif csv_idx >= 0 and csv_idx < len(self.data_manager.data_tables):
                                df = self.data_manager.data_tables[csv_idx]
                                if df is not None and signal_name in df.columns:
                                    time_col = time_columns.get(str(csv_idx), "Time") if time_columns else "Time"
                                    if time_col in df.columns:
                                        times = df[time_col].values
                                        if len(times) > 0:
                                            all_times.extend(times.tolist() if hasattr(times, 'tolist') else times)
                        except:
                            pass
                    
                    if all_times:
                        # PERFORMANCE: Use numpy for min/max on large arrays
                        times_array = np.array(all_times)
                        x_min, x_max = float(np.min(times_array)), float(np.max(times_array))
            
            if cursor_data and cursor_data.get("initialized"):
                cursor_x = cursor_data.get("x")
                if cursor_x is not None and x_max > x_min:
                    # Use actual time value for slider (not percentage)
                    slider_value = max(x_min, min(x_max, cursor_x))
                    cursor_display = f"T: {cursor_x:.3f}"
                else:
                    slider_value = x_min
                    cursor_display = f"T: {x_min:.3f}"
            else:
                slider_value = x_min
                cursor_display = f"T: {x_min:.3f}"
            
            # Calculate step size based on range (aim for ~1000 steps)
            time_range = x_max - x_min
            step = time_range / 1000 if time_range > 0 else 0.1
            
            return (
                {"display": "flex", "flexDirection": "column", "padding": "5px 10px", "backgroundColor": "rgba(0,0,0,0.05)", "borderRadius": "4px", "marginTop": "5px"},
                False,  # Enabled
                x_min,
                x_max,
                slider_value,
                cursor_display,
            )

        # =================================================================
        # Browser-style tabs with + and x buttons
        # =================================================================
        @self.app.callback(
            Output("tabs-container", "children"),
            [
                Input("tabs", "children"),
                Input("tabs", "value"),
            ],
            prevent_initial_call=False,
        )
        def render_browser_tabs(tabs_children, active_tab):
            """Render browser-style tabs with close buttons."""
            tabs_children = tabs_children or []
            active_tab = active_tab or "tab-0"

            tab_buttons = []
            for i, tab in enumerate(tabs_children):
                tab_value = f"tab-{i}"
                is_active = tab_value == active_tab

                # Extract label
                if isinstance(tab, dict):
                    label = tab.get("props", {}).get("label", f"Tab {i+1}")
                else:
                    label = f"Tab {i+1}"

                tab_buttons.append(
                    html.Div(
                        [
                            html.Span(
                                label,
                                id={"type": "tab-btn", "idx": i},
                                className="me-1",
                                style={"cursor": "pointer"},
                            ),
                            (
                                html.Span(
                                    "×",
                                    id={"type": "tab-close", "idx": i},
                                    className="text-danger",
                                    style={
                                        "cursor": "pointer",
                                        "fontSize": "14px",
                                        "fontWeight": "bold",
                                    },
                                    title="Close tab (or middle-click)",
                                )
                                if len(tabs_children) > 1
                                else None
                            ),
                        ],
                        className=f"px-2 py-1 me-1 rounded {'bg-primary text-white' if is_active else 'bg-secondary bg-opacity-25'}",
                        style={
                            "display": "inline-flex",
                            "alignItems": "center",
                            "fontSize": "11px",
                            "cursor": "pointer",
                        },
                        id={"type": "tab-container", "idx": i},
                    )
                )

            return tab_buttons

        # Handle tab button clicks (select tab)
        @self.app.callback(
            Output("tabs", "value", allow_duplicate=True),
            Input({"type": "tab-btn", "idx": ALL}, "n_clicks"),
            State({"type": "tab-btn", "idx": ALL}, "id"),
            prevent_initial_call=True,
        )
        def handle_tab_click(n_clicks, ids):
            ctx = callback_context
            if not ctx.triggered or not any(n_clicks):
                return dash.no_update

            try:
                trigger = ctx.triggered[0]["prop_id"]
                clicked_id = safe_json_parse(trigger.split(".")[0])
                if clicked_id:
                    return f"tab-{clicked_id['idx']}"
            except:
                pass
            return dash.no_update

        # Handle tab close button clicks
        @self.app.callback(
            [
                Output("tabs", "children", allow_duplicate=True),
                Output("store-num-tabs", "data", allow_duplicate=True),
                Output("store-layouts", "data", allow_duplicate=True),
                Output("tabs", "value", allow_duplicate=True),
                Output("store-assignments", "data", allow_duplicate=True),
            ],
            Input({"type": "tab-close", "idx": ALL}, "n_clicks"),
            [
                State({"type": "tab-close", "idx": ALL}, "id"),
                State("tabs", "children"),
                State("tabs", "value"),
                State("store-layouts", "data"),
                State("store-assignments", "data"),
            ],
            prevent_initial_call=True,
        )
        def handle_tab_close(n_clicks, ids, tabs, current, layouts, assignments):
            if not any(n_clicks):
                return (
                    dash.no_update,
                    dash.no_update,
                    dash.no_update,
                    dash.no_update,
                    dash.no_update,
                )

            ctx = callback_context
            if not ctx.triggered:
                return (
                    dash.no_update,
                    dash.no_update,
                    dash.no_update,
                    dash.no_update,
                    dash.no_update,
                )

            try:
                trigger = ctx.triggered[0]["prop_id"]
                clicked_id = safe_json_parse(trigger.split(".")[0])
                if not clicked_id:
                    return (
                        dash.no_update,
                        dash.no_update,
                        dash.no_update,
                        dash.no_update,
                        dash.no_update,
                    )

                del_idx = clicked_id["idx"]
                tabs = tabs or []
                layouts = layouts or {}
                assignments = assignments or {}

                if len(tabs) <= 1:
                    return (
                        dash.no_update,
                        dash.no_update,
                        dash.no_update,
                        dash.no_update,
                        dash.no_update,
                    )

                # Remove the tab
                tabs = [t for i, t in enumerate(tabs) if i != del_idx]

                # Rebuild layouts and assignments with new indices
                # FIX: Deep copy and sanitize assignments to prevent corruption, especially for tab "0"
                import copy
                new_layouts = {}
                new_assignments = {}
                new_i = 0
                for old_i in range(len(tabs) + 1):
                    if old_i == del_idx:
                        continue
                    old_key = str(old_i)
                    new_key = str(new_i)
                    
                    if old_key in layouts:
                        new_layouts[new_key] = copy.deepcopy(layouts[old_key])
                    
                    if old_key in assignments:
                        # Deep copy and sanitize the tab's assignment structure
                        old_tab_assignments = assignments[old_key]
                        new_tab_assignments = {}
                        
                        # Ensure it's a dict
                        if isinstance(old_tab_assignments, dict):
                            for sp_key, sp_data in old_tab_assignments.items():
                                # Sanitize: ensure each subplot assignment is a list
                                if isinstance(sp_data, list):
                                    # Validate each item in the list is a proper dict
                                    sanitized_list = [
                                        s for s in sp_data
                                        if isinstance(s, dict) and 'csv_idx' in s and 'signal' in s
                                    ]
                                    new_tab_assignments[sp_key] = sanitized_list
                                elif isinstance(sp_data, dict):
                                    # Old dict format - convert to empty list
                                    new_tab_assignments[sp_key] = []
                                else:
                                    # Invalid format - use empty list
                                    new_tab_assignments[sp_key] = []
                        else:
                            # Invalid tab structure - create empty dict
                            new_tab_assignments = {}
                        
                        new_assignments[new_key] = new_tab_assignments
                    else:
                        # Tab doesn't exist in assignments - initialize with empty structure
                        new_assignments[new_key] = {}
                    
                    new_i += 1

                # Rename tabs
                for i, t in enumerate(tabs):
                    if isinstance(t, dict):
                        t["props"]["label"] = f"Tab {i+1}"
                        t["props"]["value"] = f"tab-{i}"

                # Set new active tab
                new_active_idx = max(0, del_idx - 1) if del_idx > 0 else 0
                new_active = f"tab-{new_active_idx}"

                return tabs, len(tabs), new_layouts, new_active, new_assignments
            except Exception as e:
                logger.exception(f"Error closing tab: {e}")
                return (
                    dash.no_update,
                    dash.no_update,
                    dash.no_update,
                    dash.no_update,
                    dash.no_update,
                )

        # =================================================================
        # Template Save/Load
        # =================================================================
        # Template save/load disabled - use session save/load instead
        # Can be repurposed for CSV path remapping if needed
        @self.app.callback(
            [
                Output("download-template", "data"),
                Output("status-text", "children", allow_duplicate=True),
            ],
            Input("btn-save-template", "n_clicks"),
            [
                State("store-assignments", "data"),
                State("store-layouts", "data"),
                State("store-num-tabs", "data"),
                State("store-subplot-modes", "data"),
                State("store-signal-props", "data"),
                State("store-x-axis-signal", "data"),
                State("store-subplot-metadata", "data"),
            ],
            prevent_initial_call=True,
        )
        def save_template(
            n,
            assign,
            layouts,
            num_tabs,
            subplot_modes,
            props,
            x_axis_signals,
            subplot_metadata,
        ):
            # Template save/load disabled - use session save/load instead
            # Can be repurposed for CSV path remapping if needed
            return dash.no_update, "Template feature disabled - use Session Save/Load"
            
            # OLD CODE (disabled):
            if False and not n:
                return dash.no_update, dash.no_update
            try:
                # Extract only signal names from assignments (not CSV indices)
                # This allows template to work with different CSVs that have same signal names
                template_assignments = {}
                for tab_key, subplots in (assign or {}).items():
                    template_assignments[tab_key] = {}
                    for sp_key, signals in subplots.items():
                        # Assignments are always a list
                        if isinstance(signals, list):
                            template_assignments[tab_key][sp_key] = [
                                {"signal": s.get("signal", "")} for s in signals
                            ]

                template_data = {
                    "type": "template",
                    "version": "3.0",
                    "assignments": template_assignments,
                    "layouts": layouts,
                    "num_tabs": num_tabs or 1,
                    "subplot_modes": subplot_modes or {},
                    "props": props or {},
                    "x_axis_signals": x_axis_signals or {},
                    "subplot_metadata": subplot_metadata or {},
                }

                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"signal_viewer_template_{timestamp}.json"
                return (
                    dict(
                        content=json.dumps(template_data, indent=2), filename=filename
                    ),
                    "[OK] Template saved",
                )
            except Exception as e:
                return dash.no_update, f"[ERROR] {e}"

        @self.app.callback(
            [
                Output("store-assignments", "data", allow_duplicate=True),
                Output("store-layouts", "data", allow_duplicate=True),
                Output("store-num-tabs", "data", allow_duplicate=True),
                Output("tabs", "children", allow_duplicate=True),
                Output("store-subplot-modes", "data", allow_duplicate=True),
                Output("store-signal-props", "data", allow_duplicate=True),
                Output("store-x-axis-signal", "data", allow_duplicate=True),
                Output("store-subplot-metadata", "data", allow_duplicate=True),
                Output("status-text", "children", allow_duplicate=True),
            ],
            Input("upload-template", "contents"),
            [
                State("upload-template", "filename"),
                State("store-csv-files", "data"),
            ],
            prevent_initial_call=True,
        )
        def load_template(contents, filename, csv_files):
            # Template save/load disabled - use session save/load instead
            return [dash.no_update] * 9
            try:
                content_type, content_string = contents.split(",")
                decoded = base64.b64decode(content_string).decode("utf-8")
                d = json.loads(decoded)

                if d.get("type") != "template":
                    return [dash.no_update] * 8 + ["[ERROR] Not a template file"]

                # Match template signal names to available signals in loaded CSVs
                template_assignments = d.get("assignments", {})
                matched_assignments = {}

                # Build a map of signal_name -> csv_idx
                signal_to_csv = {}
                for csv_idx, fp in enumerate(csv_files or []):
                    if csv_idx < len(self.data_manager.data_tables):
                        df = self.data_manager.data_tables[csv_idx]
                        if df is not None:
                            for col in df.columns:
                                if col.lower() != "time":
                                    if col not in signal_to_csv:
                                        signal_to_csv[col] = csv_idx

                # Convert template assignments to full assignments
                for tab_key, subplots in template_assignments.items():
                    matched_assignments[tab_key] = {}
                    for sp_key, signals in subplots.items():
                        # Assignments are always a list
                        matched_signals = []
                        if isinstance(signals, list):
                            for sig in signals:
                                sig_name = sig.get("signal", "")
                                if sig_name in signal_to_csv:
                                    matched_signals.append(
                                        {
                                            "csv_idx": signal_to_csv[sig_name],
                                            "signal": sig_name,
                                        }
                                    )
                        matched_assignments[tab_key][sp_key] = matched_signals

                num_tabs = d.get("num_tabs", 1)
                tabs_children = [
                    dcc.Tab(label=f"Tab {i+1}", value=f"tab-{i}")
                    for i in range(num_tabs)
                ]

                return (
                    matched_assignments,
                    d.get("layouts", {}),
                    num_tabs,
                    tabs_children,
                    d.get("subplot_modes", {}),
                    d.get("props", {}),
                    d.get("x_axis_signals", {}),
                    d.get("subplot_metadata", {}),
                    f"[OK] Template loaded: {filename}",
                )
            except Exception as e:
                return [dash.no_update] * 8 + [f"[ERROR] {e}"]

        # =================================================================
        # Time Column Selection Modal
        # =================================================================
        @self.app.callback(
            [
                Output("modal-time-cols", "is_open"),
                Output("time-column-selectors", "children"),
                Output("time-offset-inputs", "children"),
                Output("csv-header-settings", "children"),
            ],
            [
                Input("btn-time-cols", "n_clicks"),
                Input("btn-close-time-cols", "n_clicks"),
                Input("btn-apply-time-cols", "n_clicks"),
            ],
            [
                State("modal-time-cols", "is_open"),
                State("store-csv-files", "data"),
                State("store-time-columns", "data"),
                State("store-time-offsets", "data"),
                State("store-csv-settings", "data"),
            ],
            prevent_initial_call=True,
        )
        def toggle_time_cols_modal(
            open_click, close_click, apply_click, is_open, csv_files, time_cols, time_offsets, csv_settings
        ):
            ctx = callback_context
            trigger = ctx.triggered[0]["prop_id"] if ctx.triggered else ""

            if "btn-time-cols" in trigger:
                # Build selectors for each CSV
                selectors = []
                offset_inputs = []
                header_settings = []
                time_offsets = time_offsets or {}
                csv_settings = csv_settings or {}
                
                for csv_idx, fp in enumerate(csv_files or []):
                    if csv_idx < len(self.data_manager.data_tables):
                        df = self.data_manager.data_tables[csv_idx]
                        fname = os.path.basename(fp)
                        csv_setting = csv_settings.get(str(csv_idx), {})
                        current_header_row = csv_setting.get("header_row", 0)
                        
                        if df is not None:
                            options = [
                                {"label": col, "value": col} for col in df.columns
                            ]
                            current_time_col = (time_cols or {}).get(
                                str(csv_idx), df.columns[0]
                            )
                            current_offset = time_offsets.get(str(csv_idx), 0)

                            selectors.append(
                                html.Div(
                                    [
                                        html.Small(
                                            fname,
                                            className="text-info fw-bold",
                                        ),
                                        dbc.Select(
                                            id={
                                                "type": "time-col-select",
                                                "csv": csv_idx,
                                            },
                                            options=options,
                                            value=current_time_col,
                                            size="sm",
                                            className="mt-1",
                                        ),
                                    ],
                                    className="mb-2",
                                )
                            )
                            
                            # Time offset input
                            offset_inputs.append(
                                html.Div(
                                    [
                                        dbc.Row(
                                            [
                                                dbc.Col(
                                                    html.Small(
                                                        fname,
                                                        className="text-info",
                                                    ),
                                                    width=6,
                                                ),
                                                dbc.Col(
                                                    dbc.Input(
                                                        id={
                                                            "type": "time-offset-input",
                                                            "csv": csv_idx,
                                                        },
                                                        type="number",
                                                        value=current_offset,
                                                        step=0.001,
                                                        size="sm",
                                                        placeholder="0.0",
                                                        style={"width": "100px"},
                                                    ),
                                                    width=4,
                                                ),
                                                dbc.Col(
                                                    html.Small("sec", className="text-muted"),
                                                    width=2,
                                                ),
                                            ],
                                            className="align-items-center",
                                        ),
                                    ],
                                    className="mb-2",
                                )
                            )
                        
                        # Header settings (always show, even if df is None)
                        header_settings.append(
                            html.Div(
                                [
                                    dbc.Row(
                                        [
                                            dbc.Col(
                                                html.Small(fname, className="text-info"),
                                                width=5,
                                            ),
                                            dbc.Col(
                                                dbc.Input(
                                                    id={"type": "csv-header-row", "csv": csv_idx},
                                                    type="number",
                                                    value=current_header_row if current_header_row is not None else "",
                                                    min=0,
                                                    step=1,
                                                    size="sm",
                                                    placeholder="0 (first row)",
                                                    style={"width": "80px"},
                                                ),
                                                width=3,
                                            ),
                                            dbc.Col(
                                                dbc.Checkbox(
                                                    id={"type": "csv-no-header", "csv": csv_idx},
                                                    value=(current_header_row is None),
                                                    label="No header",
                                                    className="small",
                                                ),
                                                width=4,
                                            ),
                                        ],
                                        className="align-items-center",
                                    ),
                                ],
                                className="mb-2",
                            )
                        )

                if not selectors:
                    selectors = [html.P("No CSV files loaded", className="text-muted")]
                    offset_inputs = []
                    header_settings = [html.P("Load CSV files first", className="text-muted")]

                return True, selectors, offset_inputs, header_settings

            return False, dash.no_update, dash.no_update, dash.no_update

        @self.app.callback(
            [
                Output("store-time-columns", "data"),
                Output("store-time-offsets", "data"),
                Output("store-csv-settings", "data"),
                Output("store-refresh-trigger", "data", allow_duplicate=True),
                Output("status-text", "children", allow_duplicate=True),
            ],
            Input("btn-apply-time-cols", "n_clicks"),
            [
                State({"type": "time-col-select", "csv": ALL}, "value"),
                State({"type": "time-col-select", "csv": ALL}, "id"),
                State({"type": "time-offset-input", "csv": ALL}, "value"),
                State({"type": "time-offset-input", "csv": ALL}, "id"),
                State({"type": "csv-header-row", "csv": ALL}, "value"),
                State({"type": "csv-header-row", "csv": ALL}, "id"),
                State({"type": "csv-no-header", "csv": ALL}, "value"),
                State({"type": "csv-no-header", "csv": ALL}, "id"),
                State("store-time-columns", "data"),
                State("store-time-offsets", "data"),
                State("store-csv-settings", "data"),
                State("store-refresh-trigger", "data"),
            ],
            prevent_initial_call=True,
        )
        def apply_time_columns(n, col_values, col_ids, offset_values, offset_ids, 
                               header_values, header_ids, no_header_values, no_header_ids,
                               current_time_cols, current_offsets, current_csv_settings, refresh_trigger):
            if not n:
                return dash.no_update, dash.no_update, dash.no_update, dash.no_update, dash.no_update

            time_cols = current_time_cols or {}
            time_offsets = current_offsets or {}
            csv_settings = current_csv_settings or {}
            needs_reload = False

            # Apply time columns
            if col_values and col_ids:
                for val, id_dict in zip(col_values, col_ids):
                    csv_idx = id_dict.get("csv")
                    if csv_idx is not None and val:
                        time_cols[str(csv_idx)] = val

            # Apply time offsets
            if offset_values and offset_ids:
                for val, id_dict in zip(offset_values, offset_ids):
                    csv_idx = id_dict.get("csv")
                    if csv_idx is not None:
                        offset_val = float(val) if val else 0.0
                        if offset_val != 0:
                            time_offsets[str(csv_idx)] = offset_val
                        elif str(csv_idx) in time_offsets:
                            del time_offsets[str(csv_idx)]

            # Apply CSV header settings
            no_header_map = {}
            if no_header_values and no_header_ids:
                for val, id_dict in zip(no_header_values, no_header_ids):
                    csv_idx = id_dict.get("csv")
                    if csv_idx is not None:
                        no_header_map[csv_idx] = val
            
            if header_values and header_ids:
                for val, id_dict in zip(header_values, header_ids):
                    csv_idx = id_dict.get("csv")
                    if csv_idx is not None:
                        old_setting = csv_settings.get(str(csv_idx), {})
                        old_header = old_setting.get("header_row", 0)
                        
                        # Check if "no header" checkbox is checked
                        no_header = no_header_map.get(csv_idx, False)
                        
                        if no_header:
                            new_header = None
                        else:
                            new_header = int(val) if val is not None and val != "" else 0
                        
                        # Check if header setting changed
                        if old_header != new_header:
                            needs_reload = True
                            csv_settings[str(csv_idx)] = {"header_row": new_header}
                            
                            # Reload this CSV with new settings
                            try:
                                self.data_manager.read_initial_data(csv_idx, csv_settings[str(csv_idx)])
                                self.invalidate_caches()
                            except Exception as e:
                                logger.error(f"Error reloading CSV {csv_idx}: {e}")

            # Trigger refresh if needed
            new_refresh = (refresh_trigger or 0) + 1 if needs_reload else dash.no_update

            status = "[OK] CSV settings updated" + (" (reloaded)" if needs_reload else "")
            return time_cols, time_offsets, csv_settings, new_refresh, status

        # =================================================================
        # Export CSV Modal
        # =================================================================
        @self.app.callback(
            Output("modal-export-csv", "is_open"),
            [
                Input("btn-export-csv", "n_clicks"),
                Input("btn-close-export-csv", "n_clicks"),
                Input("btn-do-export-csv", "n_clicks"),
            ],
            State("modal-export-csv", "is_open"),
            prevent_initial_call=True,
        )
        def toggle_csv_export_modal(open_click, close_click, export_click, is_open):
            ctx = callback_context
            trigger = ctx.triggered[0]["prop_id"] if ctx.triggered else ""

            if "btn-export-csv" in trigger:
                return True
            return False

        @self.app.callback(
            [
                Output("download-csv-export", "data"),
                Output("status-text", "children", allow_duplicate=True),
            ],
            Input("btn-do-export-csv", "n_clicks"),
            [
                State("export-csv-scope", "value"),
                State("export-csv-include-time", "value"),
                State("store-assignments", "data"),
                State("store-selected-tab", "data"),
                State("store-selected-subplot", "data"),
                State("store-time-columns", "data"),
            ],
            prevent_initial_call=True,
        )
        def do_csv_export(
            n, scope, include_time, assignments, sel_tab, sel_subplot, time_cols
        ):
            if not n:
                return dash.no_update, dash.no_update

            try:
                import io
                import zipfile

                with self._perf_timer("CSV Export"):
                    # Collect signals based on scope
                    signals_to_export = []

                    if scope == "subplot":
                        tab_key = str(sel_tab or 0)
                        sp_key = str(sel_subplot or 0)
                        sp_signals = (assignments or {}).get(tab_key, {}).get(sp_key, [])
                        if isinstance(sp_signals, list):
                            signals_to_export = sp_signals
                    elif scope == "tab":
                        tab_key = str(sel_tab or 0)
                        for sp_key, sp_signals in (
                            (assignments or {}).get(tab_key, {}).items()
                        ):
                            if isinstance(sp_signals, list):
                                signals_to_export.extend(sp_signals)
                    else:  # all
                        for tab_key, subplots in (assignments or {}).items():
                            for sp_key, sp_signals in subplots.items():
                                if isinstance(sp_signals, list):
                                    signals_to_export.extend(sp_signals)

                    if not signals_to_export:
                        return dash.no_update, "⚠️ No signals to export. Assign signals to subplots first."

                    # Group signals by their time length to handle different lengths
                    signals_by_length = {}  # {length: [(time_data, sig_name, sig_data), ...]}
                    
                    for sig in signals_to_export:
                        csv_idx = sig.get("csv_idx", -1)
                        sig_name = sig.get("signal", "")

                        if csv_idx >= 0 and csv_idx < len(self.data_manager.data_tables):
                            df = self.data_manager.data_tables[csv_idx]
                            if df is not None and sig_name in df.columns:
                                sig_data = df[sig_name].values
                                data_len = len(sig_data)
                                
                                # Get time column for this signal
                                time_col_name = (time_cols or {}).get(str(csv_idx), "Time")
                                time_data = df[time_col_name].values if time_col_name in df.columns else None
                                
                                # Get CSV name for labeling
                                csv_name = (
                                    get_csv_short_name(self.data_manager.csv_file_paths[csv_idx])
                                    if csv_idx < len(self.data_manager.csv_file_paths)
                                    else f"C{csv_idx}"
                                )
                                col_name = f"{sig_name}_{csv_name}"
                                
                                if data_len not in signals_by_length:
                                    signals_by_length[data_len] = {
                                        "time": time_data,
                                        "signals": []
                                    }
                                signals_by_length[data_len]["signals"].append((col_name, sig_data))

                    if not signals_by_length:
                        return dash.no_update, "[ERROR] No data to export"

                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    
                    # If all signals have same length, export single CSV
                    if len(signals_by_length) == 1:
                        length, group = list(signals_by_length.items())[0]
                        export_data = {}
                        
                        if include_time and group["time"] is not None:
                            export_data["Time"] = group["time"]
                        
                        for col_name, sig_data in group["signals"]:
                            export_data[col_name] = sig_data
                        
                        export_df = pd.DataFrame(export_data)
                        csv_string = export_df.to_csv(index=False)
                        filename = f"signals_export_{timestamp}.csv"
                        
                        return (
                            dict(content=csv_string, filename=filename),
                            f"✅ Exported {len(signals_to_export)} signals ({length} samples)",
                        )
                    
                    # Multiple lengths - create ZIP with separate CSVs
                    zip_buffer = io.BytesIO()
                    csv_files_created = []
                    
                    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
                        for length, group in sorted(signals_by_length.items()):
                            export_data = {}
                            
                            if include_time and group["time"] is not None:
                                export_data["Time"] = group["time"]
                            
                            for col_name, sig_data in group["signals"]:
                                export_data[col_name] = sig_data
                            
                            export_df = pd.DataFrame(export_data)
                            csv_content = export_df.to_csv(index=False)
                            
                            csv_filename = f"signals_{length}samples.csv"
                            zf.writestr(csv_filename, csv_content)
                            csv_files_created.append(f"{csv_filename} ({len(group['signals'])} signals)")
                    
                    zip_buffer.seek(0)
                    zip_content = base64.b64encode(zip_buffer.read()).decode('utf-8')
                    
                    filename = f"signals_export_{timestamp}.zip"
                    
                    # Format summary
                    summary = f"✅ Exported to ZIP with {len(signals_by_length)} CSVs (different time lengths):\n"
                    for f in csv_files_created:
                        summary += f"  • {f}\n"
                    
                    return (
                        dict(
                            content=zip_content,
                            filename=filename,
                            base64=True,
                            type='application/zip'
                        ),
                        summary.strip(),
                    )

            except Exception as e:
                import traceback
                logger.error(f"CSV export error: {e}\n{traceback.format_exc()}")
                return dash.no_update, f"[ERROR] {e}"

        # =================================================================
        # Export PDF Modal (placeholder - requires additional dependencies)
        # =================================================================
        @self.app.callback(
            Output("modal-export-pdf", "is_open"),
            [
                Input("btn-export-pdf", "n_clicks"),
                Input("btn-close-export-pdf", "n_clicks"),
            ],
            State("modal-export-pdf", "is_open"),
            prevent_initial_call=True,
        )
        def toggle_pdf_export_modal(open_click, close_click, is_open):
            ctx = callback_context
            trigger = ctx.triggered[0]["prop_id"] if ctx.triggered else ""

            if "btn-export-pdf" in trigger:
                return True
            return False

        @self.app.callback(
            [
                Output("download-csv-export", "data", allow_duplicate=True),
                Output("status-text", "children", allow_duplicate=True),
            ],
            Input("btn-do-export-pdf", "n_clicks"),
            [
                State("export-pdf-scope", "value"),
                State("export-pdf-title", "value"),
                State("export-pdf-intro", "value"),
                State("export-pdf-conclusion", "value"),
                State("plot", "figure"),
                State("store-subplot-metadata", "data"),
                State("store-selected-tab", "data"),
                State("store-selected-subplot", "data"),
                State("store-assignments", "data"),
                State("store-layouts", "data"),
                State("store-num-tabs", "data"),
                State("store-subplot-modes", "data"),
                State("store-x-axis-signal", "data"),
                State("store-time-columns", "data"),
                State("theme-switch", "value"),
            ],
            prevent_initial_call=True,
        )
        def do_pdf_export(
            n,
            scope,
            title,
            intro,
            conclusion,
            figure,
            metadata,
            sel_tab,
            sel_subplot,
            assignments,
            layouts,
            num_tabs,
            subplot_modes,
            x_axis_signals,
            time_columns,
            is_dark,
        ):
            if not n:
                return dash.no_update, dash.no_update

            try:
                from datetime import datetime

                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                report_title = title or "Signal Analysis Report"
                metadata = metadata or {}
                assignments = assignments or {}
                layouts = layouts or {}
                num_tabs = num_tabs or 1
                theme = "dark" if is_dark else "light"

                # Helper to generate figure for a tab (without highlight)
                def generate_tab_figure(tab_idx):
                    t_key = str(tab_idx)
                    layout = layouts.get(t_key, {"rows": 1, "cols": 1})
                    rows = layout.get("rows", 1)
                    cols = layout.get("cols", 1)
                    tab_modes = (subplot_modes or {}).get(t_key, {})
                    tab_x_signals = (x_axis_signals or {}).get(t_key, {})
                    tab_metadata = (metadata or {}).get(t_key, {})

                    # Create figure
                    fig = self.create_figure(
                        rows,
                        cols,
                        theme,
                        -1,
                        assignments,
                        t_key,
                        False,
                        False,
                        None,
                        tab_modes,
                        tab_x_signals,
                        time_columns,
                        tab_metadata,
                        {},  # time_offsets
                    )

                    # CRITICAL FIX: Re-apply legend configurations
                    # using plotly.graph_objects for proper serialization
                    import plotly.graph_objects as go

                    layout_update = {}
                    for sp_idx in range(rows * cols):
                        r = sp_idx // cols
                        c = sp_idx % cols

                        x_end = (c + 1) / cols
                        y_end = 1 - r / rows

                        legend_name = f"legend{sp_idx + 1}" if sp_idx > 0 else "legend"

                        layout_update[legend_name] = go.layout.Legend(
                            xref="paper",
                            yref="paper",
                            x=x_end - 0.02,
                            y=y_end - 0.02,
                            xanchor="right",
                            yanchor="top",
                            orientation="v",
                            borderwidth=1,
                        )

                    fig.update_layout(layout_update)
                    return fig

                # Build tab sections (plot + descriptions grouped together)
                tab_sections = []  # List of (tab_title, plot_html, descriptions_html)
                fig_num = 1
                total_figs = 0

                def build_figure_descriptions(tab_key, tab_label=""):
                    """Build HTML for figure descriptions for a specific tab"""
                    nonlocal fig_num
                    descriptions = []
                    for sp_idx in range(16):
                        sp_meta = metadata.get(tab_key, {}).get(str(sp_idx), {})
                        sp_title = sp_meta.get("title", "")
                        sp_caption = sp_meta.get("caption", "")
                        sp_description = sp_meta.get("description", "")

                        if sp_title or sp_caption or sp_description:
                            # determine direction based on available text (caption > description > title)
                            text_for_dir = sp_caption or sp_description or sp_title
                            dir_val = get_text_direction_attr(text_for_dir)
                            section = f"<div class='fig-desc' dir='{dir_val}'>"
                            label_text = f"<b>Fig {fig_num}:</b> "
                            if sp_caption:
                                label_text += sp_caption
                            elif sp_title:
                                label_text += sp_title
                            section += (
                                f"<p class='fig-label' dir='{dir_val}'>{label_text}</p>"
                            )
                            if sp_description:
                                desc_html = sp_description.replace("\n", "<br>")
                                section += f"<p class='fig-text' dir='{dir_val}'>{desc_html}</p>"
                            section += "</div>"
                            descriptions.append(section)
                            fig_num += 1
                    return "".join(descriptions)

                scope_info = ""

                # OFFLINE FIX: First plot includes Plotly.js, rest don't (to avoid duplication)
                first_plot = True
                
                if scope == "subplot":
                    # Single subplot
                    fig = generate_tab_figure(sel_tab or 0)
                    plot_html = fig.to_html(
                        include_plotlyjs=True,  # OFFLINE: Embed Plotly.js
                        full_html=False,
                        config={
                            "displayModeBar": True,
                            "toImageButtonOptions": {
                                "format": "png",
                                "height": 800,
                                "width": 1200,
                            },
                        },
                    )
                    first_plot = False

                    tab_key = str(sel_tab or 0)
                    sp_meta = metadata.get(tab_key, {}).get(str(sel_subplot or 0), {})
                    sp_title = sp_meta.get("title", "")
                    sp_caption = sp_meta.get("caption", "")
                    sp_description = sp_meta.get("description", "")

                    desc_html = ""
                    if sp_title or sp_caption or sp_description:
                        # add dir attribute based on available text
                        text_for_dir = sp_caption or sp_description or sp_title
                        dir_val = get_text_direction_attr(text_for_dir)
                        desc_html = (
                            f"<div class='fig-desc' dir='{dir_val}'><b>Fig 1:</b> "
                        )
                        desc_html += sp_caption if sp_caption else sp_title
                        if sp_description:
                            desc_html += f"<p class='fig-text' dir='{dir_val}'>{sp_description.replace(chr(10), '<br>')}</p>"
                        desc_html += "</div>"
                        total_figs = 1

                    tab_sections.append(("", plot_html, desc_html))
                    scope_info = (
                        f"Subplot {(sel_subplot or 0) + 1} of Tab {int(tab_key) + 1}"
                    )

                elif scope == "tab":
                    # Single tab with all subplots
                    fig = generate_tab_figure(sel_tab or 0)
                    plot_html = fig.to_html(
                        include_plotlyjs=True,  # OFFLINE: Embed Plotly.js
                        full_html=False,
                        config={
                            "displayModeBar": True,
                            "toImageButtonOptions": {
                                "format": "png",
                                "height": 800,
                                "width": 1200,
                            },
                        },
                    )
                    first_plot = False

                    tab_key = str(sel_tab or 0)
                    descriptions_html = build_figure_descriptions(tab_key)
                    total_figs = fig_num - 1

                    tab_sections.append(("", plot_html, descriptions_html))
                    scope_info = f"Tab {int(tab_key) + 1}"

                else:  # all tabs
                    # Each tab with its descriptions grouped together
                    for tab_idx in range(num_tabs):
                        fig = generate_tab_figure(tab_idx)
                        plot_html = fig.to_html(
                            # OFFLINE: First plot embeds Plotly.js, rest don't
                            include_plotlyjs=first_plot,
                            full_html=False,
                            config={
                                "displayModeBar": True,
                                "toImageButtonOptions": {
                                    "format": "png",
                                    "height": 800,
                                    "width": 1200,
                                },
                            },
                        )
                        first_plot = False

                        t_key = str(tab_idx)
                        descriptions_html = build_figure_descriptions(
                            t_key, f"Tab {tab_idx + 1}"
                        )

                        tab_sections.append(
                            (f"Tab {tab_idx + 1}", plot_html, descriptions_html)
                        )

                    total_figs = fig_num - 1
                    scope_info = f"All {num_tabs} Tab(s)"

                # Build combined content: each tab's plot followed by its descriptions
                content_section = ""
                for tab_title, plot_html, descriptions_html in tab_sections:
                    content_section += f"""
                    <div class="tab-section">
                        {"<h3 class='tab-header'>" + tab_title + "</h3>" if tab_title else ""}
                        <div class="plot-container">
                            {plot_html}
                        </div>
                        {"<div class='descriptions-section'>" + descriptions_html + "</div>" if descriptions_html else ""}
                    </div>
                    """

                # Prepare intro/conclusion sections with proper RTL support
                intro_section = ""
                if intro:
                    intro_dir = get_text_direction_attr(intro)
                    intro_style = get_text_direction_style(intro)
                    intro_escaped = intro.replace("<", "&lt;").replace(">", "&gt;")
                    intro_section = f"<div class='section' dir='{intro_dir}'><h2>Introduction</h2><pre dir='{intro_dir}' style='white-space: pre-wrap; line-height: 1.6; font-family: inherit; background: #f9f9f9; padding: 10px; border-radius: 5px; {intro_style}'>{intro_escaped}</pre></div>"

                conclusion_section = ""
                if conclusion:
                    concl_dir = get_text_direction_attr(conclusion)
                    concl_style = get_text_direction_style(conclusion)
                    concl_escaped = conclusion.replace("<", "&lt;").replace(">", "&gt;")
                    conclusion_section = f"<div class='section' dir='{concl_dir}'><h2>Conclusion</h2><pre dir='{concl_dir}' style='white-space: pre-wrap; line-height: 1.6; font-family: inherit; background: #f9f9f9; padding: 10px; border-radius: 5px; {concl_style}'>{concl_escaped}</pre></div>"

                # Build full HTML document
                # OFFLINE FIX: Plotly.js is now embedded in the plot HTML, no CDN needed
                html_content = f"""<!DOCTYPE html>
                <html>
                <head>
                <meta charset="utf-8">
                <meta http-equiv="Content-Type" content="text/html; charset=utf-8">
                <title>{report_title}</title>
                <!-- Plotly.js is embedded in the plot below for offline use -->
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 0; padding: 20px; background: #fff; color: #333; }}
                    .container {{ max-width: 1200px; margin: 0 auto; }}
                    h1 {{ text-align: center; border-bottom: 2px solid #333; padding-bottom: 10px; }}
                    h2 {{ color: #333; margin-top: 30px; }}
                    h3 {{ color: #555; }}
                    .meta {{ text-align: center; color: #666; font-size: 12px; margin-bottom: 20px; }}
                    .section {{ margin: 20px 0; }}
                    .tab-section {{ margin: 30px 0; padding-bottom: 20px; border-bottom: 1px dashed #ccc; }}
                    .tab-header {{ color: #2E86AB; margin-bottom: 15px; font-size: 18px; }}
                    .plot-container {{ margin: 20px 0; border: 1px solid #ddd; border-radius: 5px; padding: 10px; background: #fafafa; }}
                    .descriptions-section {{ margin-top: 15px; }}
                    .fig-desc {{ margin: 10px 0; padding: 10px; background: #f0f0f0; border-radius: 5px; border-left: 3px solid #4ea8de; }}
                    .fig-label {{ font-style: italic; margin: 5px 0; }}
                    .fig-text {{ margin: 10px 0; line-height: 1.6; }}
                    pre[style*='direction: rtl'] {{ direction: rtl !important; text-align: right !important; unicode-bidi: embed !important; }}
                    pre[style*='direction: ltr'] {{ direction: ltr !important; text-align: left !important; unicode-bidi: embed !important; }}
                    @media print {{
                        .tab-section {{ page-break-inside: avoid; }}
                        body {{ background: #fff; }}
                    }}
                </style>
            </head>
            <body>
            <div class="container">
        <h1>{report_title}</h1>
        <p class="meta">Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")} | {scope_info}</p>
        
        {intro_section}
        
        {content_section}
        
        {conclusion_section}
    </div>
</body>
</html>"""

                filename = f"signal_report_{timestamp}.html"

                return (
                    dict(
                        content=html_content,
                        filename=filename,
                    ),
                    f"[OK] Exported {len(tab_sections)} tab(s), {total_figs} figure(s): {filename}",
                )

            except Exception as e:
                logger.exception(f"Export error: {e}")
                return dash.no_update, f"[ERROR] Export failed: {str(e)}"

        # =================================================================
        # Export Word Document
        # =================================================================
        @self.app.callback(
            Output("modal-export-word", "is_open"),
            [
                Input("btn-export-word", "n_clicks"),
                Input("btn-close-export-word", "n_clicks"),
                Input("btn-do-export-word", "n_clicks"),
            ],
            State("modal-export-word", "is_open"),
            prevent_initial_call=True,
        )
        def toggle_word_export_modal(open_click, close_click, export_click, is_open):
            ctx = callback_context
            trigger = ctx.triggered[0]["prop_id"] if ctx.triggered else ""

            if "btn-export-word" in trigger:
                return True
            return False

        @self.app.callback(
            [
                Output("download-word-export", "data"),
                Output("status-text", "children", allow_duplicate=True),
            ],
            Input("btn-do-export-word", "n_clicks"),
            [
                State("export-word-scope", "value"),
                State("export-word-title", "value"),
                State("export-word-intro", "value"),
                State("export-word-conclusion", "value"),
                State("store-subplot-metadata", "data"),
                State("store-selected-tab", "data"),
                State("store-selected-subplot", "data"),
                State("store-assignments", "data"),
                State("store-layouts", "data"),
                State("store-num-tabs", "data"),
                State("store-subplot-modes", "data"),
                State("store-x-axis-signal", "data"),
                State("store-time-columns", "data"),
                State("theme-switch", "value"),
            ],
            prevent_initial_call=True,
        )
        def do_word_export(
            n,
            scope,
            title,
            intro,
            conclusion,
            metadata,
            sel_tab,
            sel_subplot,
            assignments,
            layouts,
            num_tabs,
            subplot_modes,
            x_axis_signals,
            time_columns,
            is_dark,
        ):
            if not n:
                return dash.no_update, dash.no_update

            try:
                from docx import Document
                from docx.shared import Inches, Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                import io
                import base64

                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                report_title = title or "Signal Analysis Report"
                metadata = metadata or {}
                assignments = assignments or {}
                layouts = layouts or {}
                num_tabs = num_tabs or 1
                theme = "dark" if is_dark else "light"

                # Create Word document
                doc = Document()
                
                # Title
                title_para = doc.add_heading(report_title, 0)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Metadata
                meta_para = doc.add_paragraph()
                meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                meta_run = meta_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
                meta_run.font.size = Pt(10)
                meta_run.font.italic = True
                
                # Introduction
                if intro:
                    doc.add_heading("Introduction", level=1)
                    doc.add_paragraph(intro)
                
                # Helper to generate figure for a tab
                def generate_tab_figure(tab_idx):
                    t_key = str(tab_idx)
                    layout = layouts.get(t_key, {"rows": 1, "cols": 1})
                    rows = layout.get("rows", 1)
                    cols = layout.get("cols", 1)
                    tab_modes = (subplot_modes or {}).get(t_key, {})
                    tab_x_signals = (x_axis_signals or {}).get(t_key, {})
                    tab_metadata = (metadata or {}).get(t_key, {})

                    fig = self.create_figure(
                        rows, cols, theme, -1, assignments, t_key,
                        False, False, None, tab_modes, tab_x_signals,
                        time_columns, tab_metadata, {}
                    )
                    return fig

                fig_num = 1
                total_figs = 0

                # Determine tabs to export
                if scope == "subplot":
                    tabs_to_export = [sel_tab or 0]
                elif scope == "tab":
                    tabs_to_export = [sel_tab or 0]
                else:
                    tabs_to_export = list(range(num_tabs))

                for tab_idx in tabs_to_export:
                    t_key = str(tab_idx)
                    
                    if len(tabs_to_export) > 1:
                        doc.add_heading(f"Tab {tab_idx + 1}", level=1)
                    
                    # Generate and add the figure
                    fig = generate_tab_figure(tab_idx)
                    
                    # Convert figure to PNG image
                    img_bytes = fig.to_image(format="png", width=1200, height=800, scale=2)
                    img_stream = io.BytesIO(img_bytes)
                    
                    # Add image to document
                    doc.add_picture(img_stream, width=Inches(6.5))
                    last_para = doc.paragraphs[-1]
                    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add figure descriptions
                    tab_metadata = metadata.get(t_key, {})
                    for sp_idx in range(16):
                        sp_meta = tab_metadata.get(str(sp_idx), {})
                        sp_title = sp_meta.get("title", "")
                        sp_caption = sp_meta.get("caption", "")
                        sp_description = sp_meta.get("description", "")
                        
                        if sp_title or sp_caption or sp_description:
                            caption_text = f"Figure {fig_num}: "
                            caption_text += sp_caption if sp_caption else sp_title
                            
                            caption_para = doc.add_paragraph()
                            caption_run = caption_para.add_run(caption_text)
                            caption_run.font.italic = True
                            caption_run.font.size = Pt(10)
                            
                            if sp_description:
                                doc.add_paragraph(sp_description)
                            
                            fig_num += 1
                            total_figs += 1
                    
                    doc.add_paragraph()  # Add spacing
                
                # Conclusion
                if conclusion:
                    doc.add_heading("Conclusion", level=1)
                    doc.add_paragraph(conclusion)
                
                # Save to bytes
                doc_stream = io.BytesIO()
                doc.save(doc_stream)
                doc_stream.seek(0)
                
                # Encode for download
                doc_content = base64.b64encode(doc_stream.read()).decode('utf-8')
                
                filename = f"signal_report_{timestamp}.docx"
                
                return (
                    dict(
                        content=doc_content,
                        filename=filename,
                        base64=True,
                    ),
                    f"[OK] Exported to Word: {filename}",
                )

            except ImportError as e:
                if "docx" in str(e):
                    return dash.no_update, "⚠️ Word export requires python-docx. Install: pip install python-docx"
                elif "kaleido" in str(e).lower():
                    return dash.no_update, "⚠️ Image export requires kaleido. Install: pip install kaleido"
                return dash.no_update, f"⚠️ Missing package: {str(e)}"
            except ValueError as e:
                if "kaleido" in str(e).lower() or "orca" in str(e).lower():
                    return dash.no_update, "⚠️ Image export requires kaleido. Install: pip install kaleido"
                return dash.no_update, f"⚠️ Export failed: {str(e)}"
            except Exception as e:
                logger.exception(f"Word export error: {e}")
                return dash.no_update, f"[ERROR] Export failed: {str(e)}"

        # =================================================================
        # Statistics Panel & Display Options
        # =================================================================
        @self.app.callback(
            [
                Output("stats-panel-collapse", "is_open"),
                Output("stats-panel-content", "children"),
                Output("store-display-options", "data"),
            ],
            [
                Input("signal-display-options", "value"),
            ],
            [
                State("store-assignments", "data"),
                State("store-selected-tab", "data"),
                State("store-selected-subplot", "data"),
                State("store-display-options", "data"),
            ],
            prevent_initial_call=True,
        )
        def update_display_options(options, assignments, sel_tab, sel_subplot, display_opts):
            options = options or []
            show_stats = "stats" in options
            show_markers = "markers" in options
            normalize = "normalize" in options
            
            # Store display options per tab/subplot
            display_opts = display_opts or {}
            tab_key = str(sel_tab or 0)
            subplot_key = str(sel_subplot or 0)
            
            if tab_key not in display_opts:
                display_opts[tab_key] = {}
            display_opts[tab_key][subplot_key] = {
                "markers": show_markers,
                "normalize": normalize,
            }
            
            # Build statistics content
            stats_content = []
            if show_stats:
                assignments = assignments or {}
                assigned = assignments.get(tab_key, {}).get(subplot_key, [])
                
                if isinstance(assigned, list) and len(assigned) > 0:
                    for sig in assigned:
                        csv_idx = sig.get("csv_idx", -1)
                        sig_name = sig.get("signal", "")
                        
                        try:
                            if csv_idx == -1:
                                # Derived signal
                                continue
                            
                            data = self.data_manager.get_signal_data(csv_idx, sig_name)
                            if data is not None and len(data) > 0:
                                min_val = float(np.nanmin(data))
                                max_val = float(np.nanmax(data))
                                mean_val = float(np.nanmean(data))
                                std_val = float(np.nanstd(data))
                                
                                stats_content.append(
                                    html.Div(
                                        [
                                            html.Strong(f"{sig_name[:15]}:", className="text-info"),
                                            html.Span(
                                                f" Min: {min_val:.3g} | Max: {max_val:.3g} | "
                                                f"Mean: {mean_val:.3g} | Std: {std_val:.3g}",
                                                style={"fontSize": "9px"},
                                            ),
                                        ],
                                        className="mb-1",
                                    )
                                )
                        except Exception:
                            pass
                    
                    if not stats_content:
                        stats_content = [html.Span("No data available", className="text-muted")]
                else:
                    stats_content = [html.Span("No signals assigned", className="text-muted")]
            
            return show_stats, stats_content, display_opts

        # =================================================================
        # Annotation Modal & Management
        # =================================================================
        @self.app.callback(
            [
                Output("modal-annotation", "is_open"),
                Output("annotation-x", "value"),
                Output("annotation-y", "value"),
            ],
            [
                Input("btn-open-annotation", "n_clicks"),
                Input("btn-close-annotation", "n_clicks"),
                Input("btn-add-annotation", "n_clicks"),
            ],
            [
                State("modal-annotation", "is_open"),
                State("store-cursor-x", "data"),
            ],
            prevent_initial_call=True,
        )
        def toggle_annotation_modal(open_click, close_click, add_click, is_open, cursor_data):
            ctx = callback_context
            trigger = ctx.triggered[0]["prop_id"] if ctx.triggered else ""
            
            if "btn-open-annotation" in trigger:
                # Pre-fill with cursor position if available
                x_val = cursor_data.get("x") if cursor_data else None
                return True, x_val, None
            elif "btn-add-annotation" in trigger or "btn-close-annotation" in trigger:
                return False, None, None
            
            return is_open, None, None

        @self.app.callback(
            [
                Output("store-annotations", "data"),
                Output("store-refresh-trigger", "data", allow_duplicate=True),
            ],
            [
                Input("btn-add-annotation", "n_clicks"),
                Input("btn-clear-annotations", "n_clicks"),
            ],
            [
                State("annotation-x", "value"),
                State("annotation-y", "value"),
                State("annotation-text", "value"),
                State("annotation-color", "value"),
                State("annotation-fontsize", "value"),
                State("annotation-arrow", "value"),
                State("store-annotations", "data"),
                State("store-selected-tab", "data"),
                State("store-selected-subplot", "data"),
                State("store-refresh-trigger", "data"),
            ],
            prevent_initial_call=True,
        )
        def manage_annotations(
            add_click, clear_click, x, y, text, color, fontsize, arrow,
            annotations, sel_tab, sel_subplot, refresh
        ):
            ctx = callback_context
            trigger = ctx.triggered[0]["prop_id"] if ctx.triggered else ""
            
            annotations = annotations or {}
            tab_key = str(sel_tab or 0)
            subplot_key = str(sel_subplot or 0)
            
            if "btn-clear-annotations" in trigger:
                if tab_key in annotations:
                    annotations[tab_key][subplot_key] = []
                return annotations, (refresh or 0) + 1
            
            if "btn-add-annotation" in trigger:
                if x is not None and text:
                    if tab_key not in annotations:
                        annotations[tab_key] = {}
                    if subplot_key not in annotations[tab_key]:
                        annotations[tab_key][subplot_key] = []
                    
                    annotations[tab_key][subplot_key].append({
                        "x": x,
                        "y": y,
                        "text": text,
                        "color": color or "#ffcc00",
                        "fontsize": fontsize or 12,
                        "arrow": bool(arrow),
                    })
                    return annotations, (refresh or 0) + 1
            
            return no_update, no_update

        # =================================================================
        # Subplot Metadata (title, caption) - save and sync
        # =================================================================
        @self.app.callback(
            Output("store-subplot-metadata", "data"),
            [
                Input("subplot-title-input", "value"),
                Input("subplot-caption-input", "value"),
                Input("subplot-description-input", "value"),
            ],
            [
                State("store-subplot-metadata", "data"),
                State("store-selected-tab", "data"),
                State("store-selected-subplot", "data"),
            ],
            prevent_initial_call=True,
        )
        def save_subplot_metadata(
            title, caption, description, metadata, sel_tab, sel_subplot
        ):
            metadata = metadata or {}
            tab_key = str(sel_tab or 0)
            subplot_key = str(sel_subplot or 0)

            if tab_key not in metadata:
                metadata[tab_key] = {}
            if subplot_key not in metadata[tab_key]:
                metadata[tab_key][subplot_key] = {}

            metadata[tab_key][subplot_key]["title"] = title or ""
            metadata[tab_key][subplot_key]["caption"] = caption or ""
            metadata[tab_key][subplot_key]["description"] = description or ""

            return metadata

        # Sync subplot metadata when switching subplots
        @self.app.callback(
            [
                Output("subplot-title-input", "value"),
                Output("subplot-caption-input", "value"),
                Output("subplot-description-input", "value"),
            ],
            [
                Input("store-selected-subplot", "data"),
                Input("tabs", "value"),
            ],
            [State("store-subplot-metadata", "data")],
            prevent_initial_call=True,
        )
        def sync_subplot_metadata(sel_subplot, active_tab, metadata):
            metadata = metadata or {}
            tab_idx = (
                int(active_tab.split("-")[1]) if active_tab and "-" in active_tab else 0
            )
            tab_key = str(tab_idx)
            subplot_key = str(sel_subplot or 0)

            sp_meta = metadata.get(tab_key, {}).get(subplot_key, {})
            return (
                sp_meta.get("title", ""),
                sp_meta.get("caption", ""),
                sp_meta.get("description", ""),
            )

        # =================================================================
        # Document Text (intro, conclusion) - save and sync
        # =================================================================
        @self.app.callback(
            [
                Output("store-document-text", "data"),
                Output("status-text", "children", allow_duplicate=True),
            ],
            Input("btn-save-doc-text", "n_clicks"),
            [
                State("export-pdf-intro", "value"),
                State("export-pdf-conclusion", "value"),
                State("store-document-text", "data"),
            ],
            prevent_initial_call=True,
        )
        def save_document_text(n, intro, conclusion, current_doc):
            if not n:
                return dash.no_update, dash.no_update

            doc_text = current_doc or {}
            doc_text["introduction"] = intro or ""
            doc_text["conclusion"] = conclusion or ""

            return doc_text, "[OK] Document text saved"

        # Load document text when modal opens
        @self.app.callback(
            [
                Output("export-pdf-intro", "value"),
                Output("export-pdf-conclusion", "value"),
            ],
            Input("modal-export-pdf", "is_open"),
            State("store-document-text", "data"),
            prevent_initial_call=True,
        )
        def load_document_text(is_open, doc_text):
            if is_open:
                doc_text = doc_text or {}
                return doc_text.get("introduction", ""), doc_text.get("conclusion", "")
            return dash.no_update, dash.no_update


        # =================================================================
        # CSV Loading Callbacks (Flexible CSV Loader)
        # =================================================================
        
        @self.app.callback(
            [
                Output("modal-csv-loader", "is_open"),
                Output("store-pending-csv-path", "data"),
                Output("csv-raw-preview-text", "children"),
                Output("csv-file-info-display", "children"),
            ],
            [
                Input("upload-csv", "contents"),
                Input("btn-csv-confirm", "n_clicks"),
                Input("btn-csv-cancel", "n_clicks"),
            ],
            [
                State("upload-csv", "filename"),
                State("store-pending-csv-path", "data"),
            ],
            prevent_initial_call=True,
        )
        def toggle_csv_modal_and_preview(contents, confirm, cancel, filename, pending_path):
            """Open CSV loading modal with preview"""
            ctx = callback_context
            trigger = ctx.triggered[0]["prop_id"] if ctx.triggered else ""
            
            if "upload-csv" in trigger and contents:
                try:
                    import base64
                    import tempfile
                    
                    content_type, content_string = contents.split(',')
                    decoded = base64.b64decode(content_string)
                    
                    # Create temp file
                    temp_path = os.path.join(tempfile.gettempdir(), filename)
                    with open(temp_path, 'wb') as f:
                        f.write(decoded)
                    
                    # Get preview
                    preview = self.data_manager.get_csv_preview(temp_path, max_lines=15)
                    
                    # File info
                    file_size = os.path.getsize(temp_path)
                    size_mb = file_size / (1024 * 1024)
                    file_info = html.Div([
                        html.Strong(f"📁 {filename}"),
                        html.Div(f"Size: {size_mb:.2f} MB", style={'color': '#888', 'fontSize': '12px'}),
                    ])
                    
                    return True, temp_path, preview, file_info
                    
                except Exception as e:
                    return True, None, f"Error: {str(e)}", html.Div("Error reading file", style={'color': '#ff6b6b'})
            
            elif "btn-csv-confirm" in trigger or "btn-csv-cancel" in trigger:
                return False, None, "", ""
            
            return no_update, no_update, no_update, no_update
        
        @self.app.callback(
            [
                Output("csv-delimiter-dropdown", "value"),
                Output("csv-header-row-input", "value"),
            ],
            Input("btn-csv-auto-detect", "n_clicks"),
            State("store-pending-csv-path", "data"),
            prevent_initial_call=True,
        )
        def auto_detect_csv_settings(n_clicks, filepath):
            """Auto-detect CSV format"""
            if not filepath or not os.path.exists(filepath):
                return no_update, no_update
            
            try:
                format_info = self.data_manager.detect_csv_format(filepath)
                delimiter = format_info.get('delimiter', ',')
                header_row = format_info.get('header_row', 0)
                
                return delimiter, header_row
            except Exception as e:
                logger.error(f"Auto-detect error: {e}")
                return no_update, no_update
        
        @self.app.callback(
            Output("csv-parsed-preview-table", "children"),
            [
                Input("csv-delimiter-dropdown", "value"),
                Input("csv-header-row-input", "value"),
            ],
            State("store-pending-csv-path", "data"),
            prevent_initial_call=True,
        )
        def update_csv_parsed_preview(delimiter, header_row, filepath):
            """Update parsed preview when settings change"""
            if not filepath or not os.path.exists(filepath):
                return html.Div("No file loaded", style={'color': '#888'})
            
            try:
                # Convert 'auto' to None
                if delimiter == 'auto':
                    delimiter = None
                
                # Load preview
                df = self.data_manager.flexible_loader.load_csv(
                    filepath,
                    auto_detect=(delimiter is None),
                    delimiter=delimiter,
                    header_row=header_row,
                    skiprows=0,
                    preview_mode=True,
                )
                
                if df is None or df.empty:
                    return html.Div("Could not parse CSV", style={'color': '#ff6b6b'})
                
                # Create preview table
                from dash import dash_table
                preview_df = df.head(8)
                
                return dash_table.DataTable(
                    data=preview_df.to_dict('records'),
                    columns=[{'name': col, 'id': col} for col in preview_df.columns],
                    style_table={'overflowX': 'auto', 'maxHeight': '250px'},
                    style_cell={
                        'textAlign': 'left',
                        'padding': '6px',
                        'backgroundColor': '#16213e',
                        'color': '#e8e8e8',
                        'fontSize': '10px',
                        'fontFamily': 'monospace',
                    },
                    style_header={
                        'backgroundColor': '#0f3460',
                        'fontWeight': 'bold',
                    },
                )
                
            except Exception as e:
                return html.Div(f"Parse error: {str(e)}", style={'color': '#ff6b6b', 'fontSize': '12px'})
        
        @self.app.callback(
            [
                Output("store-csv-files", "data", allow_duplicate=True),
                Output("csv-list", "children", allow_duplicate=True),
                Output("status-text", "children", allow_duplicate=True),
                Output("store-refresh-trigger", "data", allow_duplicate=True),
            ],
            Input("btn-csv-confirm", "n_clicks"),
            [
                State("store-pending-csv-path", "data"),
                State("csv-delimiter-dropdown", "value"),
                State("csv-header-row-input", "value"),
                State("store-csv-files", "data"),
                State("store-refresh-trigger", "data"),
            ],
            prevent_initial_call=True,
        )
        def load_csv_with_flexible_loader(n_clicks, filepath, delimiter, header_row, current_csvs, refresh_counter):
            """Load CSV with confirmed settings"""
            if not filepath or not os.path.exists(filepath):
                return no_update, no_update, no_update, no_update
            
            refresh_counter = refresh_counter or 0
            
            try:
                # Convert 'auto' to None
                if delimiter == 'auto':
                    delimiter = None
                
                # Load CSV
                success = self.data_manager.load_csv_with_settings(
                    filepath,
                    csv_idx=None,
                    delimiter=delimiter,
                    header_row=header_row,
                    skip_rows=0,
                    auto_detect=(delimiter is None),
                )
                
                if success:
                    current_csvs = current_csvs or []
                    if filepath not in current_csvs:
                        current_csvs.append(filepath)
                    
                    csv_idx = len(self.data_manager.data_tables) - 1
                    df = self.data_manager.data_tables[csv_idx]
                    
                    filename = os.path.basename(filepath)
                    
                    # Build CSV list UI
                    items = []
                    for i, f in enumerate(current_csvs):
                        items.append(
                            html.Div(
                                [
                                    html.I(className="fas fa-file me-1", style={"color": "#f4a261"}),
                                    html.Span(
                                        get_csv_display_name(f, current_csvs), 
                                        style={"fontSize": "10px"},
                                        title=f
                                    ),
                                    html.A(
                                        "×",
                                        id={"type": "del-csv", "idx": i},
                                        className="float-end text-danger",
                                        style={"cursor": "pointer"},
                                    ),
                                ],
                                className="py-1",
                            )
                        )
                    
                    return current_csvs, items, f"[OK] Loaded {filename} ({len(df)} rows, {len(df.columns)} cols)", refresh_counter + 1
                else:
                    return no_update, no_update, "[ERROR] Failed to load CSV", no_update
                    
            except Exception as e:
                return no_update, no_update, f"[ERROR] Error: {str(e)}", no_update


        # =================================================================
        # Refresh CSV - Complete Reload
        # =================================================================
        
        @self.app.callback(
            [
                Output("store-refresh-trigger", "data", allow_duplicate=True),
                Output("status-text", "children", allow_duplicate=True),
                Output("store-assignments", "data", allow_duplicate=True),  # Update assignments
            ],
            Input("btn-refresh-csv", "n_clicks"),
            [
                State("store-csv-files", "data"),
                State("store-assignments", "data"),  # Get current assignments
            ],
            prevent_initial_call=True,
        )
        def refresh_all_csvs(n_clicks, csv_files, assignments):
            """Refresh all CSVs - re-read data and update everything"""
            if not csv_files:
                return no_update, "No CSVs loaded"
            
            try:
                print("\n" + "="*80)
                print("[REFRESH] REFRESH BUTTON CLICKED")
                print("="*80)
                logger.info("=== REFRESH: Starting CSV refresh ===")
                
                # Show current state with full paths and file info
                print(f"CSV files to refresh: {len(csv_files)}")
                for i, fp in enumerate(csv_files):
                    # Check file on disk
                    if os.path.exists(fp):
                        file_size = os.path.getsize(fp)
                        file_mtime = datetime.fromtimestamp(os.path.getmtime(fp))
                        file_info = f"({file_size:,} bytes, modified: {file_mtime})"
                    else:
                        file_info = "[WARN] FILE NOT FOUND ON DISK!"
                    
                    if i < len(self.data_manager.data_tables) and self.data_manager.data_tables[i] is not None:
                        rows = len(self.data_manager.data_tables[i])
                        print(f"  [{i}] {fp}")
                        print(f"       In memory: {rows} rows | On disk: {file_info}")
                    else:
                        print(f"  [{i}] {fp}")
                        print(f"       In memory: No data | On disk: {file_info}")
                
                # Clear ALL caches including disk cache (critical for refresh!)
                self.invalidate_caches()
                self.data_manager.invalidate_cache(clear_disk_cache=True)
                print("[OK] All caches invalidated (including disk cache)")
                
                # FORCE reload by clearing tables and mod times
                for i in range(len(csv_files)):
                    if i < len(self.data_manager.data_tables):
                        old_rows = len(self.data_manager.data_tables[i]) if self.data_manager.data_tables[i] is not None else 0
                        self.data_manager.data_tables[i] = None
                        print(f"  Cleared table {i} (was {old_rows} rows)")
                    if i < len(self.data_manager.last_read_rows):
                        self.data_manager.last_read_rows[i] = 0
                    if i < len(self.data_manager.last_file_mod_times):
                        self.data_manager.last_file_mod_times[i] = 0
                
                logger.info(f"REFRESH: Cleared {len(csv_files)} table(s)")
                
                # Store old signal lists before refresh to detect changes
                old_signal_lists = {}
                for i in range(len(csv_files)):
                    if i < len(self.data_manager.data_tables) and self.data_manager.data_tables[i] is not None:
                        df = self.data_manager.data_tables[i]
                        old_signal_lists[i] = set([c for c in df.columns if c.lower() != 'time'])
                    else:
                        old_signal_lists[i] = set()
                
                # Now reload all CSVs from ORIGINAL paths (not uploads cache)
                print("\n[RELOAD] Reloading CSVs from original paths...")
                
                for i, uploads_path in enumerate(csv_files):
                    # Get original path if available, otherwise use uploads path
                    original_path = self.original_file_paths.get(uploads_path, uploads_path)
                    
                    if os.path.exists(original_path):
                        try:
                            # Read from ORIGINAL path
                            df = pd.read_csv(original_path, low_memory=False)
                            
                            # Ensure Time column
                            if "Time" not in df.columns and len(df.columns) > 0:
                                df.rename(columns={df.columns[0]: "Time"}, inplace=True)
                            
                            self.data_manager.data_tables[i] = df
                            self.data_manager.last_read_rows[i] = len(df)
                            self.data_manager.last_file_mod_times[i] = os.path.getmtime(original_path)
                            
                            # If original != uploads, copy updated file to uploads for caching
                            if original_path != uploads_path:
                                try:
                                    import shutil
                                    shutil.copy2(original_path, uploads_path)
                                    print(f"  [OK] [{i}] {os.path.basename(original_path)}: {len(df)} rows (synced to cache)")
                                except:
                                    print(f"  [OK] [{i}] {os.path.basename(original_path)}: {len(df)} rows")
                            else:
                                print(f"  [OK] [{i}] {os.path.basename(original_path)}: {len(df)} rows")
                            
                            # Show first and last row to verify it's fresh data
                            if len(df) > 0:
                                print(f"       First row Time: {df['Time'].iloc[0]}")
                                print(f"       Last row Time:  {df['Time'].iloc[-1]}")
                        except Exception as e:
                            print(f"  [ERROR] [{i}] {os.path.basename(original_path)}: Error - {e}")
                            self.data_manager.data_tables[i] = None
                    else:
                        print(f"  [WARN] [{i}] Original not found: {original_path}")
                        # Try uploads path as fallback
                        if os.path.exists(uploads_path):
                            try:
                                df = pd.read_csv(uploads_path, low_memory=False)
                                if "Time" not in df.columns and len(df.columns) > 0:
                                    df.rename(columns={df.columns[0]: "Time"}, inplace=True)
                                self.data_manager.data_tables[i] = df
                                print(f"  [WARN] [{i}] Using cached: {os.path.basename(uploads_path)} ({len(df)} rows)")
                            except Exception as e:
                                print(f"  [ERROR] [{i}] Cache also failed: {e}")
                                self.data_manager.data_tables[i] = None
                        else:
                            self.data_manager.data_tables[i] = None
                
                # Update signal names
                self.data_manager.update_signal_names()
                
                # Get new signal lists after refresh
                new_signal_lists = {}
                for i in range(len(csv_files)):
                    if i < len(self.data_manager.data_tables) and self.data_manager.data_tables[i] is not None:
                        df = self.data_manager.data_tables[i]
                        new_signal_lists[i] = set([c for c in df.columns if c.lower() != 'time'])
                    else:
                        new_signal_lists[i] = set()
                
                # Clean up assignments: remove signals that no longer exist, keep those that do
                assignments = assignments or {}
                assignments_updated = False
                removed_count = 0
                kept_count = 0
                
                for tab_key in assignments:
                    for sp_key in assignments[tab_key]:
                        sp_assignment = assignments[tab_key][sp_key]
                        if isinstance(sp_assignment, list):
                            # Filter out signals that no longer exist
                            valid_signals = []
                            for sig_info in sp_assignment:
                                csv_idx = sig_info.get("csv_idx", -1)
                                signal_name = sig_info.get("signal", "")
                                
                                if csv_idx == -1:
                                    # Derived signal - keep it (will be recalculated if needed)
                                    valid_signals.append(sig_info)
                                    kept_count += 1
                                elif csv_idx in new_signal_lists:
                                    # Check if signal still exists in CSV
                                    if signal_name in new_signal_lists[csv_idx]:
                                        valid_signals.append(sig_info)
                                        kept_count += 1
                                    else:
                                        removed_count += 1
                                        print(f"  [REMOVED] Signal '{signal_name}' from CSV {csv_idx} (no longer exists)")
                                else:
                                    # CSV index out of range
                                    removed_count += 1
                            
                            if len(valid_signals) != len(sp_assignment):
                                assignments[tab_key][sp_key] = valid_signals
                                assignments_updated = True
                
                if assignments_updated:
                    print(f"\n[ASSIGNMENTS] Updated: kept {kept_count}, removed {removed_count} invalid signals")
                
                # Show new state
                print("\n📊 After refresh:")
                
                # Clear derived signals that depend on removed signals
                derived_to_remove = []
                for derived_name, derived_info in self.derived_signals.items():
                    source_signal_key = derived_info.get("source_signal_key")
                    source_csv_idx = derived_info.get("source_csv_idx")
                    source_signals = derived_info.get("source_signals", [])
                    
                    should_remove = False
                    if source_csv_idx is not None and source_csv_idx >= 0:
                        # Check if source signal still exists
                        if source_csv_idx not in new_signal_lists:
                            should_remove = True
                        elif source_signal_key:
                            csv_idx, sig_name = source_signal_key.split(":", 1)
                            csv_idx = int(csv_idx)
                            if csv_idx >= 0 and sig_name not in new_signal_lists.get(csv_idx, set()):
                                should_remove = True
                    
                    # Check source_signals for multi-op signals
                    if not should_remove and source_signals:
                        for sig_key in source_signals:
                            if ":" in sig_key:
                                csv_idx, sig_name = sig_key.split(":", 1)
                                csv_idx = int(csv_idx)
                                if csv_idx >= 0 and sig_name not in new_signal_lists.get(csv_idx, set()):
                                    should_remove = True
                                    break
                    
                    if should_remove:
                        derived_to_remove.append(derived_name)
                
                for name in derived_to_remove:
                    del self.derived_signals[name]
                    print(f"  [REMOVED] Derived signal '{name}' (source signal removed)")
                
                # Note: Valid derived signals will be recalculated with fresh data on next plot update
                # This ensures they reflect the refreshed CSV data
                valid_derived_count = len(self.derived_signals)
                if valid_derived_count > 0:
                    print(f"  [INFO] {valid_derived_count} derived signal(s) will be recalculated on next plot update")
                
                # Trigger full UI refresh by updating trigger
                refresh_trigger = datetime.now().timestamp()
                
                # Count results
                total_rows = sum(len(df) for df in self.data_manager.data_tables if df is not None)
                num_csvs = len([df for df in self.data_manager.data_tables if df is not None])
                
                logger.info(f"REFRESH: Complete - {num_csvs} CSV(s), {total_rows:,} rows")
                
                print(f"\n[OK] REFRESH COMPLETE: {num_csvs} CSV(s), {total_rows:,} rows")
                print("="*80 + "\n")
                
                return refresh_trigger, f"[OK] Refreshed {num_csvs} CSV(s), {total_rows:,} rows", assignments
                
            except Exception as e:
                logger.error(f"REFRESH ERROR: {e}")
                import traceback
                traceback.print_exc()
                print(f"\n[ERROR] REFRESH FAILED: {e}")
                print("="*80 + "\n")
                return no_update, f"[ERROR] Refresh failed: {str(e)}", dash.no_update
        
        # =================================================================
        # CSV Streaming Feature
        # =================================================================
        
        @self.app.callback(
            [
                Output("store-streaming-active", "data"),
                Output("btn-stream-csv", "children"),
                Output("btn-stream-csv", "color"),
                Output("interval-streaming", "disabled"),
                Output("status-text", "children", allow_duplicate=True),
            ],
            Input("btn-stream-csv", "n_clicks"),
            State("store-streaming-active", "data"),
            prevent_initial_call=True,
        )
        def toggle_csv_streaming(n_clicks, is_streaming):
            """Toggle CSV streaming on/off"""
            if is_streaming:
                # Stop streaming
                self.data_manager.stop_streaming_all()
                print("[STOP] Streaming stopped by user")
                return False, "▶ Stream", "success", True, "[STOP] Streaming stopped"
            else:
                # Start streaming
                if not self.data_manager.csv_file_paths:
                    print("[ERROR] Cannot start streaming - no CSVs loaded")
                    return False, "▶ Stream", "success", True, "[ERROR] No CSVs loaded"
                
                # Enable streaming with 0.2s update rate
                self.data_manager.update_rate = 0.2
                self.data_manager.timeout_duration = 1.0  # Stop after 1s of no updates
                self.data_manager.start_streaming_all()
                
                print(f"[START] Streaming started for {len(self.data_manager.csv_file_paths)} CSV(s)")
                print(f"   Update rate: {self.data_manager.update_rate}s")
                print(f"   Timeout: {self.data_manager.timeout_duration}s")
                
                return True, "⏹ Stop", "danger", False, "[START] Streaming active (updates every 0.2s)"
        
        @self.app.callback(
            [
                Output("store-refresh-trigger", "data", allow_duplicate=True),
                Output("store-streaming-active", "data", allow_duplicate=True),
                Output("btn-stream-csv", "children", allow_duplicate=True),
                Output("btn-stream-csv", "color", allow_duplicate=True),
                Output("interval-streaming", "disabled", allow_duplicate=True),
                Output("status-text", "children", allow_duplicate=True),
            ],
            Input("interval-streaming", "n_intervals"),
            State("store-streaming-active", "data"),
            prevent_initial_call=True,
        )
        def update_from_streaming(n_intervals, is_streaming):
            """Update plots when streaming detects changes, with auto-stop on timeout"""
            if not is_streaming:
                return no_update, no_update, no_update, no_update, no_update, no_update
            
            # Check for CSV updates and read new data
            result = self.data_manager.check_and_update_streaming()
            
            # Check if we should auto-stop due to timeout
            if result.get('should_stop', False):
                self.data_manager.stop_streaming_all()
                print(f"[STOP] Auto-stopped streaming (timeout)")
                return (
                    datetime.now().timestamp(),  # Refresh trigger
                    False,  # streaming active
                    "▶ Stream",  # button text
                    "success",  # button color
                    True,  # interval disabled
                    result.get('status_text', '[STOP] Streaming stopped (timeout)')
                )
            
            if result.get('updated', False):
                print(f"   [OK] {result.get('status_text', 'Changes detected!')}")
                # CRITICAL: Clear app's caches when new data arrives
                self.invalidate_caches()
                # Trigger plot refresh and update status
                return (
                    datetime.now().timestamp(),
                    no_update,  # keep streaming active
                    no_update,  # keep button
                    no_update,  # keep color
                    no_update,  # keep interval enabled
                    f"[START] {result.get('status_text', 'Streaming...')}"
                )
            else:
                # No changes - just update status to show we're still watching
                return (
                    no_update,  # no refresh needed
                    no_update,
                    no_update,
                    no_update,
                    no_update,
                    f"[START] {result.get('status_text', 'Streaming...')}"
                )

        # =================================================================
        # Load CSV from Path (for streaming from original location)
        # =================================================================
        
        @self.app.callback(
            [
                Output("store-csv-files", "data", allow_duplicate=True),
                Output("store-refresh-trigger", "data", allow_duplicate=True),
                Output("status-text", "children", allow_duplicate=True),
                Output("input-csv-path", "value"),
            ],
            Input("btn-load-path", "n_clicks"),
            [
                State("input-csv-path", "value"),
                State("store-csv-files", "data"),
            ],
            prevent_initial_call=True,
        )
        def load_csv_from_path(n_clicks, original_path, current_files):
            """Load CSV directly from original path - NO COPY to uploads.
            Cache goes to uploads/.cache/, streaming reads from original."""
            if not original_path or not original_path.strip():
                return no_update, no_update, "[ERROR] Enter a file path", no_update
            
            original_path = original_path.strip().strip('"').strip("'")
            
            if not os.path.exists(original_path):
                return no_update, no_update, f"[ERROR] Not found: {original_path}", original_path
            
            if not original_path.lower().endswith('.csv'):
                return no_update, no_update, "[ERROR] Must be .csv file", original_path
            
            current_files = current_files or []
            fname = os.path.basename(original_path)
            
            # Check if already loaded (by original path)
            if original_path in current_files or original_path in self.original_file_paths.values():
                return no_update, no_update, f"[WARN] Already loaded: {fname}", ""
            
            # Use ORIGINAL path directly - no copy!
            current_files.append(original_path)
            
            # Store original path mapping (path → itself for path loads)
            self.original_file_paths[original_path] = original_path
            
            # Update data_manager
            self.data_manager.csv_file_paths = current_files
            idx = len(current_files) - 1
            self.data_manager.original_source_paths[idx] = original_path
            
            while len(self.data_manager.data_tables) < len(current_files):
                self.data_manager.data_tables.append(None)
                self.data_manager.last_read_rows.append(0)
                self.data_manager.last_file_mod_times.append(0)
            
            try:
                self.data_manager.read_initial_data(idx)
                rows = len(self.data_manager.data_tables[idx]) if self.data_manager.data_tables[idx] is not None else 0
                parent_folder = os.path.basename(os.path.dirname(original_path))
                print(f"[OK] Loaded: {parent_folder}/{fname} ({rows} rows) - streaming from original")
                return current_files, datetime.now().timestamp(), f"[OK] {parent_folder}/{fname} ({rows:,} rows)", ""
            except Exception as e:
                print(f"[ERROR] Load failed: {e}")
                return current_files, datetime.now().timestamp(), f"[ERROR] {str(e)}", original_path
        
        # =================================================================
        # CSV Comparison Features
        # =================================================================
        
        @self.app.callback(
            [
                Output("modal-compare", "is_open"),
                Output("compare-csv1", "options"),
                Output("compare-csv2", "options"),
                Output("compare-signal1", "options"),
                Output("compare-signal2", "options"),
            ],
            [
                Input("btn-compare-csvs", "n_clicks"),
                Input("btn-close-compare", "n_clicks"),
            ],
            [
                State("modal-compare", "is_open"),
                State("store-csv-files", "data"),
            ],
            prevent_initial_call=True,
        )
        def toggle_compare_modal(open_click, close_click, is_open, csv_files):
            """Open/close the compare modal and populate dropdowns"""
            ctx = callback_context
            if not ctx.triggered:
                return no_update, no_update, no_update, no_update, no_update
            
            trigger = ctx.triggered[0]["prop_id"]
            
            if "btn-close-compare" in trigger:
                return False, no_update, no_update, no_update, no_update
            
            if "btn-compare-csvs" in trigger:
                # Build CSV options
                csv_options = []
                signal_options = []
                
                for i, path in enumerate(csv_files or []):
                    display_name = get_csv_display_name(path, csv_files)
                    csv_options.append({"label": display_name, "value": str(i)})
                    
                    # Add signals from this CSV
                    if i < len(self.data_manager.data_tables) and self.data_manager.data_tables[i] is not None:
                        for col in self.data_manager.data_tables[i].columns:
                            if col != "Time":
                                signal_options.append({
                                    "label": f"{col} ({display_name})",
                                    "value": f"{i}:{col}"
                                })
                
                return True, csv_options, csv_options, signal_options, signal_options
            
            return no_update, no_update, no_update, no_update, no_update
        
        @self.app.callback(
            [
                Output("compare-signal1", "options", allow_duplicate=True),
                Output("compare-signal2", "options", allow_duplicate=True),
            ],
            [
                Input("compare-csv1", "value"),
                Input("compare-csv2", "value"),
            ],
            State("store-csv-files", "data"),
            prevent_initial_call=True,
        )
        def update_signal_dropdowns(csv1_idx, csv2_list, csv_files):
            """Update signal dropdowns based on selected CSVs"""
            signal1_options = []
            signal2_options = []
            
            if csv1_idx is not None:
                idx = int(csv1_idx)
                if idx < len(self.data_manager.data_tables) and self.data_manager.data_tables[idx] is not None:
                    df = self.data_manager.data_tables[idx]
                    display_name = get_csv_display_name(csv_files[idx], csv_files) if idx < len(csv_files) else f"CSV {idx}"
                    for col in df.columns:
                        if col != "Time":
                            signal1_options.append({"label": col, "value": f"{idx}:{col}"})
            
            # csv2_list is a list from Checklist - add signals from all selected CSVs
            if csv2_list:
                for csv_idx_str in csv2_list:
                    idx = int(csv_idx_str)
                    if idx < len(self.data_manager.data_tables) and self.data_manager.data_tables[idx] is not None:
                        df = self.data_manager.data_tables[idx]
                        display_name = get_csv_display_name(csv_files[idx], csv_files) if idx < len(csv_files) else f"CSV {idx}"
                        for col in df.columns:
                            if col != "Time":
                                signal2_options.append({"label": f"{col} ({display_name})", "value": f"{idx}:{col}"})
            
            return signal1_options, signal2_options
        
        @self.app.callback(
            Output("compare-results", "children"),
            [
                Input("btn-do-compare-csvs", "n_clicks"),
                Input("btn-do-compare-signals", "n_clicks"),
                Input("btn-plot-diff", "n_clicks"),
            ],
            [
                State("compare-csv1", "value"),
                State("compare-csv2", "value"),
                State("compare-signal1", "value"),
                State("compare-signal2", "value"),
                State("store-csv-files", "data"),
            ],
            prevent_initial_call=True,
        )
        def do_comparison(csv_click, signal_click, diff_click, csv1_idx, csv2_list, sig1_key, sig2_key, csv_files):
            """Perform CSV or signal comparison (supports multiple CSVs)"""
            from helpers import compare_signals, compare_csv_signals
            
            ctx = callback_context
            if not ctx.triggered:
                return no_update
            
            trigger = ctx.triggered[0]["prop_id"]
            
            if "btn-do-compare-csvs" in trigger:
                if csv1_idx is None or not csv2_list:
                    return html.Div("Select reference CSV and at least one to compare", className="text-warning")
                
                ref_idx = int(csv1_idx)
                compare_indices = [int(x) for x in csv2_list if int(x) != ref_idx]
                
                if not compare_indices:
                    return html.Div("Select at least one different CSV to compare", className="text-warning")
                
                df_ref = self.data_manager.data_tables[ref_idx] if ref_idx < len(self.data_manager.data_tables) else None
                
                if df_ref is None:
                    return html.Div("Reference CSV not loaded", className="text-danger")
                
                # Compare reference to each selected CSV
                all_results = []
                for comp_idx in compare_indices:
                    df_comp = self.data_manager.data_tables[comp_idx] if comp_idx < len(self.data_manager.data_tables) else None
                    if df_comp is None:
                        continue
                    
                    results = compare_csv_signals(df_ref, df_comp)
                    comp_name = get_csv_display_name(csv_files[comp_idx], csv_files)
                    all_results.append((comp_name, results))
                
                if not all_results:
                    return html.Div("No valid CSVs to compare", className="text-danger")
                
                # Build results - if multiple CSVs, show average distance from mean
                if len(all_results) == 1:
                    # Single comparison - same as before
                    comp_name, results = all_results[0]
                    summary = results.get("_summary", {})
                    rows = []
                    for sig_name, metrics in results.items():
                        if sig_name == "_summary":
                            continue
                        if "error" in metrics:
                            rows.append(html.Tr([html.Td(sig_name), html.Td(f"Error: {metrics['error']}", colSpan=4, className="text-danger")]))
                        else:
                            corr_class = "text-success" if metrics["correlation"] > 0.99 else "text-warning" if metrics["correlation"] > 0.9 else "text-danger"
                            rows.append(html.Tr([
                                html.Td(sig_name), html.Td(f"{metrics['correlation']:.4f}", className=corr_class),
                                html.Td(f"{metrics['rmse']:.6f}"), html.Td(f"{metrics['max_diff']:.6f}"),
                                html.Td(f"{metrics['match_rate']:.1f}%"),
                            ]))
                    
                    return html.Div([
                        html.H6(f"📊 Comparison with {comp_name}"),
                        html.P(f"Signals compared: {summary.get('compared_count', 0)}, Avg corr: {summary.get('avg_correlation', 0):.4f}", className="small"),
                        dbc.Table([
                            html.Thead(html.Tr([html.Th("Signal"), html.Th("Corr"), html.Th("RMSE"), html.Th("MaxDiff"), html.Th("Match%")])),
                            html.Tbody(rows)
                        ], striped=True, hover=True, size="sm")
                    ])
                else:
                    # Multiple CSVs - show average metrics across all comparisons
                    ref_name = get_csv_display_name(csv_files[ref_idx], csv_files)
                    signal_metrics = {}  # {signal: [list of metrics from each comparison]}
                    
                    for comp_name, results in all_results:
                        for sig_name, metrics in results.items():
                            if sig_name == "_summary" or "error" in metrics:
                                continue
                            if sig_name not in signal_metrics:
                                signal_metrics[sig_name] = []
                            signal_metrics[sig_name].append(metrics)
                    
                    rows = []
                    for sig_name, metrics_list in signal_metrics.items():
                        avg_corr = np.mean([m["correlation"] for m in metrics_list])
                        avg_rmse = np.mean([m["rmse"] for m in metrics_list])
                        max_diff = max([m["max_diff"] for m in metrics_list])
                        avg_match = np.mean([m["match_rate"] for m in metrics_list])
                        
                        corr_class = "text-success" if avg_corr > 0.99 else "text-warning" if avg_corr > 0.9 else "text-danger"
                        rows.append(html.Tr([
                            html.Td(sig_name), html.Td(f"{avg_corr:.4f}", className=corr_class),
                            html.Td(f"{avg_rmse:.6f}"), html.Td(f"{max_diff:.6f}"),
                            html.Td(f"{avg_match:.1f}%"),
                        ]))
                    
                    return html.Div([
                        html.H6(f"📊 Multi-CSV Comparison ({len(all_results)} CSVs vs {ref_name})"),
                        html.P(f"Showing average metrics across all comparisons", className="small text-muted"),
                        dbc.Table([
                            html.Thead(html.Tr([html.Th("Signal"), html.Th("Avg Corr"), html.Th("Avg RMSE"), html.Th("Max Diff"), html.Th("Avg Match%")])),
                            html.Tbody(rows)
                        ], striped=True, hover=True, size="sm")
                    ])
            
            if "btn-do-compare-signals" in trigger:
                if not sig1_key or not sig2_key:
                    return html.Div("Please select two signals to compare", className="text-warning")
                
                # Parse signal keys
                csv1_idx, sig1_name = sig1_key.split(":", 1)
                csv2_idx, sig2_name = sig2_key.split(":", 1)
                csv1_idx, csv2_idx = int(csv1_idx), int(csv2_idx)
                
                df1 = self.data_manager.data_tables[csv1_idx]
                df2 = self.data_manager.data_tables[csv2_idx]
                
                if df1 is None or df2 is None:
                    return html.Div("CSV data not loaded", className="text-danger")
                
                time1 = df1["Time"].values if "Time" in df1.columns else np.arange(len(df1))
                time2 = df2["Time"].values if "Time" in df2.columns else np.arange(len(df2))
                data1 = df1[sig1_name].values
                data2 = df2[sig2_name].values
                
                metrics = compare_signals(time1, data1, time2, data2)
                
                if "error" in metrics:
                    return html.Div(f"Error: {metrics['error']}", className="text-danger")
                
                corr_class = "text-success" if metrics["correlation"] > 0.99 else "text-warning" if metrics["correlation"] > 0.9 else "text-danger"
                
                return html.Div([
                    html.H6(f"📊 Signal Comparison"),
                    html.P(f"Comparing: {sig1_name} vs {sig2_name}", className="small text-muted"),
                    dbc.Table([
                        html.Tbody([
                            html.Tr([html.Td("Correlation"), html.Td(f"{metrics['correlation']:.4f}", className=corr_class)]),
                            html.Tr([html.Td("RMSE"), html.Td(f"{metrics['rmse']:.6f}")]),
                            html.Tr([html.Td("MAE"), html.Td(f"{metrics['mae']:.6f}")]),
                            html.Tr([html.Td("Max Difference"), html.Td(f"{metrics['max_diff']:.6f}")]),
                            html.Tr([html.Td("Mean Difference (Bias)"), html.Td(f"{metrics['mean_diff']:.6f}")]),
                            html.Tr([html.Td("Percent Difference"), html.Td(f"{metrics['percent_diff']:.2f}%")]),
                            html.Tr([html.Td("Match Rate (±1%)"), html.Td(f"{metrics['match_rate']:.1f}%")]),
                            html.Tr([html.Td("Points Compared"), html.Td(f"{metrics['num_points']:,}")]),
                        ])
                    ], bordered=True, size="sm")
                ])
            
            if "btn-plot-diff" in trigger:
                # Plot difference between two signals
                if not sig1_key or not sig2_key:
                    return html.Div("Please select two signals to plot difference", className="text-warning")
                
                # Parse signal keys
                csv1_idx, sig1_name = sig1_key.split(":", 1)
                csv2_idx, sig2_name = sig2_key.split(":", 1)
                csv1_idx, csv2_idx = int(csv1_idx), int(csv2_idx)
                
                df1 = self.data_manager.data_tables[csv1_idx]
                df2 = self.data_manager.data_tables[csv2_idx]
                
                if df1 is None or df2 is None:
                    return html.Div("CSV data not loaded", className="text-danger")
                
                time1 = df1["Time"].values if "Time" in df1.columns else np.arange(len(df1))
                time2 = df2["Time"].values if "Time" in df2.columns else np.arange(len(df2))
                data1 = df1[sig1_name].values if sig1_name in df1.columns else None
                data2 = df2[sig2_name].values if sig2_name in df2.columns else None
                
                if data1 is None or data2 is None:
                    return html.Div("Signal not found in CSV", className="text-danger")
                
                # Interpolate to common time base
                try:
                    from scipy import interpolate as sp_interp
                    has_scipy = True
                except ImportError:
                    has_scipy = False
                
                common_time = np.union1d(time1, time2)
                
                # Only use overlapping range
                t_min = max(time1.min(), time2.min())
                t_max = min(time1.max(), time2.max())
                common_time = common_time[(common_time >= t_min) & (common_time <= t_max)]
                
                if len(common_time) == 0:
                    return html.Div("No overlapping time range", className="text-warning")
                
                # Interpolate both signals
                if has_scipy:
                    f1 = sp_interp.interp1d(time1, data1, bounds_error=False, fill_value=np.nan)
                    f2 = sp_interp.interp1d(time2, data2, bounds_error=False, fill_value=np.nan)
                    interp1 = f1(common_time)
                    interp2 = f2(common_time)
                else:
                    # Fallback to numpy interpolation
                    interp1 = np.interp(common_time, time1, data1)
                    interp2 = np.interp(common_time, time2, data2)
                
                diff = interp1 - interp2
                
                # Create 2 subplots: signals on top, difference on bottom
                from plotly.subplots import make_subplots
                fig = make_subplots(rows=2, cols=1, shared_xaxes=True,
                                    subplot_titles=[f"Signals: {sig1_name} vs {sig2_name}", "Difference"],
                                    vertical_spacing=0.12)
                
                # Top: Both signals
                csv1_name = get_csv_display_name(csv_files[csv1_idx], csv_files) if csv_files else f"CSV{csv1_idx}"
                csv2_name = get_csv_display_name(csv_files[csv2_idx], csv_files) if csv_files else f"CSV{csv2_idx}"
                
                # Use Scattergl for WebGL acceleration (faster rendering)
                fig.add_trace(go.Scattergl(x=common_time, y=interp1, mode="lines",
                                         name=f"{sig1_name} ({csv1_name})",
                                         line=dict(color="#4ea8de", width=1)), row=1, col=1)
                fig.add_trace(go.Scattergl(x=common_time, y=interp2, mode="lines",
                                         name=f"{sig2_name} ({csv2_name})",
                                         line=dict(color="#f4a261", width=1)), row=1, col=1)

                # Bottom: Difference
                fig.add_trace(go.Scattergl(x=common_time, y=diff, mode="lines", name="Difference",
                                         line=dict(color="#ff6b6b", width=1)), row=2, col=1)
                fig.add_hline(y=0, line_dash="dash", line_color="gray", row=2, col=1)
                
                fig.update_layout(
                    height=400,
                    margin=dict(l=50, r=20, t=40, b=30),
                    paper_bgcolor="#1a1a2e",
                    plot_bgcolor="#16213e",
                    font=dict(color="#e8e8e8", size=10),
                    showlegend=True,
                    legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="center", x=0.5),
                )
                fig.update_xaxes(gridcolor="#2a2a4e")
                fig.update_yaxes(gridcolor="#2a2a4e")
                
                # Stats
                valid_diff = diff[~np.isnan(diff)]
                stats_text = f"Mean diff: {np.mean(valid_diff):.4f}, Std: {np.std(valid_diff):.4f}, Max abs: {np.max(np.abs(valid_diff)):.4f}"
                
                return html.Div([
                    html.P(stats_text, className="small text-muted mb-1"),
                    dcc.Graph(figure=fig, config={"displayModeBar": True, "displaylogo": False})
                ])
            
            return no_update

        # =================================================================
        # PRESET TEMPLATES CALLBACKS
        # =================================================================
        
        @self.app.callback(
            Output("modal-preset", "is_open"),
            [
                Input("btn-save-preset", "n_clicks"),
                Input("btn-close-preset-modal", "n_clicks"),
                Input("btn-do-save-preset", "n_clicks"),
            ],
            State("modal-preset", "is_open"),
            prevent_initial_call=True,
        )
        def toggle_preset_modal(save_click, close_click, do_save_click, is_open):
            ctx = callback_context
            trigger = ctx.triggered[0]["prop_id"] if ctx.triggered else ""
            if "btn-save-preset" in trigger:
                return True
            return False

        @self.app.callback(
            [
                Output("store-presets", "data", allow_duplicate=True),
                Output("preset-name-input", "value"),
            ],
            Input("btn-do-save-preset", "n_clicks"),
            [
                State("preset-name-input", "value"),
                State("preset-description-input", "value"),
                State("preset-include-signals", "value"),
                State("preset-include-layout", "value"),
                State("store-assignments", "data"),
                State("store-layouts", "data"),
                State("rows-input", "value"),
                State("cols-input", "value"),
                State("store-presets", "data"),
            ],
            prevent_initial_call=True,
        )
        def save_preset(n_clicks, name, description, include_signals, include_layout, 
                       assignments, layouts, rows, cols, presets):
            if not name:
                return dash.no_update, dash.no_update
            
            import time
            presets = presets or {}
            
            # Create preset object
            preset_id = f"preset_{int(time.time())}"
            preset = {
                "id": preset_id,
                "name": name,
                "description": description or "",
                "created": time.time(),
            }
            
            if include_signals:
                preset["assignments"] = assignments
            if include_layout:
                preset["layouts"] = layouts
                preset["rows"] = rows
                preset["cols"] = cols
            
            presets[preset_id] = preset
            print(f"[PRESET] Saved preset: {name}")
            
            # Toast notification via JS
            return presets, ""

        @self.app.callback(
            Output("dropdown-presets", "children"),
            Input("store-presets", "data"),
            prevent_initial_call=True,
        )
        def update_preset_dropdown(presets):
            """Update preset dropdown menu with saved presets."""
            items = [
                dbc.DropdownMenuItem("📋 Save as Preset", id="btn-save-preset"),
                dbc.DropdownMenuItem(divider=True),
            ]
            
            if not presets:
                items.append(dbc.DropdownMenuItem("No saved presets", disabled=True))
            else:
                for preset_id, preset in presets.items():
                    items.append(
                        dbc.DropdownMenuItem(
                            [
                                html.Span(preset.get("name", "Unnamed"), className="me-2"),
                                html.Small(preset.get("description", "")[:20], className="text-muted"),
                            ],
                            id={"type": "preset-item", "id": preset_id},
                        )
                    )
                items.append(dbc.DropdownMenuItem(divider=True))
                items.append(dbc.DropdownMenuItem("🗑️ Clear All Presets", id="btn-clear-presets", className="text-danger"))
            
            return items

        @self.app.callback(
            [
                Output("store-assignments", "data", allow_duplicate=True),
                Output("store-layouts", "data", allow_duplicate=True),
                Output("rows-input", "value", allow_duplicate=True),
                Output("cols-input", "value", allow_duplicate=True),
            ],
            Input({"type": "preset-item", "id": ALL}, "n_clicks"),
            State("store-presets", "data"),
            prevent_initial_call=True,
        )
        def load_preset(n_clicks, presets):
            """Load a preset when clicked."""
            ctx = callback_context
            if not ctx.triggered or not any(n_clicks):
                return dash.no_update, dash.no_update, dash.no_update, dash.no_update
            
            trigger = ctx.triggered[0]["prop_id"]
            # Extract preset ID from trigger
            import json
            try:
                trigger_id = json.loads(trigger.split(".")[0])
                preset_id = trigger_id.get("id")
            except:
                return dash.no_update, dash.no_update, dash.no_update, dash.no_update
            
            if preset_id and presets and preset_id in presets:
                preset = presets[preset_id]
                print(f"[PRESET] Loading preset: {preset.get('name')}")
                
                return (
                    preset.get("assignments", dash.no_update),
                    preset.get("layouts", dash.no_update),
                    preset.get("rows", dash.no_update),
                    preset.get("cols", dash.no_update),
                )
            
            return dash.no_update, dash.no_update, dash.no_update, dash.no_update

        @self.app.callback(
            Output("store-presets", "data", allow_duplicate=True),
            Input("btn-clear-presets", "n_clicks"),
            prevent_initial_call=True,
        )
        def clear_presets(n_clicks):
            if n_clicks:
                print("[PRESET] Cleared all presets")
                return {}
            return dash.no_update

        # =================================================================
        # COMPARE MODE CALLBACKS
        # =================================================================
        
        @self.app.callback(
            [
                Output("store-compare-mode", "data"),
                Output("compare-mode-panel", "style"),
                Output("split-plot", "style"),
                Output("btn-compare-mode", "color"),
            ],
            [
                Input("btn-compare-mode", "n_clicks"),
                Input("btn-exit-compare", "n_clicks"),
            ],
            State("store-compare-mode", "data"),
            prevent_initial_call=True,
        )
        def toggle_compare_mode(compare_click, exit_click, is_compare):
            ctx = callback_context
            trigger = ctx.triggered[0]["prop_id"] if ctx.triggered else ""
            
            if "btn-compare-mode" in trigger:
                new_state = not is_compare
            elif "btn-exit-compare" in trigger:
                new_state = False
            else:
                return dash.no_update, dash.no_update, dash.no_update, dash.no_update
            
            if new_state:
                # Enter compare mode
                return (
                    True,
                    {"display": "block", "height": "100%"},
                    {"display": "none"},
                    "primary",
                )
            else:
                # Exit compare mode
                return (
                    False,
                    {"display": "none"},
                    {"height": "100%", "overflow": "hidden"},
                    "secondary",
                )

        @self.app.callback(
            [
                Output("compare-tab-left", "options"),
                Output("compare-tab-right", "options"),
            ],
            Input("store-compare-mode", "data"),
            State("tabs", "children"),
            prevent_initial_call=True,
        )
        def update_compare_tab_options(is_compare, tabs_children):
            """Populate tab options for compare mode."""
            if not is_compare or not tabs_children:
                return [], []
            
            options = []
            for i, child in enumerate(tabs_children):
                if hasattr(child, 'props') and 'label' in child.props:
                    options.append({"label": child.props['label'], "value": str(i)})
                else:
                    options.append({"label": f"Tab {i+1}", "value": str(i)})
            
            return options, options

        @self.app.callback(
            [
                Output("compare-plot-left", "figure"),
                Output("compare-plot-right", "figure"),
            ],
            [
                Input("compare-tab-left", "value"),
                Input("compare-tab-right", "value"),
            ],
            [
                State("store-assignments", "data"),
                State("store-layouts", "data"),
                State("theme-switch", "value"),
                State("store-time-columns", "data"),
            ],
            prevent_initial_call=True,
        )
        def update_compare_plots(left_tab, right_tab, assignments, layouts, is_dark, time_columns):
            """Update compare mode plots based on selected tabs."""
            if not left_tab and not right_tab:
                return dash.no_update, dash.no_update
            
            theme = "plotly_dark" if is_dark else "plotly_white"
            left_fig = dash.no_update
            right_fig = dash.no_update
            
            if left_tab:
                layout = layouts.get(left_tab, {"rows": 1, "cols": 1}) if layouts else {"rows": 1, "cols": 1}
                left_fig = self.create_figure(
                    rows=layout.get("rows", 1),
                    cols=layout.get("cols", 1),
                    theme=theme,
                    selected_subplot=0,
                    assignments=assignments.get(left_tab, {}),
                    tab_key=left_tab,
                    time_columns=time_columns,
                )
            
            if right_tab:
                layout = layouts.get(right_tab, {"rows": 1, "cols": 1}) if layouts else {"rows": 1, "cols": 1}
                right_fig = self.create_figure(
                    rows=layout.get("rows", 1),
                    cols=layout.get("cols", 1),
                    theme=theme,
                    selected_subplot=0,
                    assignments=assignments.get(right_tab, {}),
                    tab_key=right_tab,
                    time_columns=time_columns,
                )
            
            return left_fig, right_fig

        # =================================================================
        # SIGNAL REORDERING CALLBACK (Drag & Drop)
        # =================================================================
        
        # Clientside callback to handle drag & drop events from JavaScript
        self.app.clientside_callback(
            """
            function(id) {
                // Listen for signal reorder events from JavaScript
                if (!window._signalReorderSetup) {
                    window._signalReorderSetup = true;
                    
                    document.addEventListener('signalReorder', function(e) {
                        var detail = e.detail;
                        // Store in a hidden element or trigger callback
                        var store = document.getElementById('store-signal-reorder');
                        if (store) {
                            // Update the store data via custom approach
                            window._pendingReorder = detail;
                        }
                    });
                    
                    console.log('[SignalViewer] Signal reorder listener setup');
                }
                return window.dash_clientside.no_update;
            }
            """,
            Output("store-signal-reorder", "data", allow_duplicate=True),
            Input("app-container", "id"),
            prevent_initial_call=True,
        )

        logger.info("All callbacks registered successfully (including CSV loading, refresh, streaming, comparison, presets, and compare mode)")


def main():
    """Main entry point for Signal Viewer Pro."""
    import webbrowser
    import threading
    import time

    try:
        print("=" * 50)
        print(f"  {APP_TITLE} - {APP_HOST}:{APP_PORT}")
        print("=" * 50)

        logger.info("Creating SignalViewerApp...")
        app = SignalViewerApp()
        logger.info(f"App created with {len(app.app.callback_map)} callbacks")

        # Open browser after slight delay
        threading.Thread(
            target=lambda: (
                time.sleep(1.5),
                webbrowser.open(f"http://{APP_HOST}:{APP_PORT}"),
            ),
            daemon=True,
        ).start()

        # use_reloader=False prevents double-startup and double browser windows
        app.app.run(debug=True, port=APP_PORT, host=APP_HOST, use_reloader=False)

    except Exception as e:
        logger.exception(f"Application error: {e}")
        input("Press Enter to exit...")


if __name__ == "__main__":
    main()